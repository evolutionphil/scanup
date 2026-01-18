from fastapi import FastAPI, APIRouter, HTTPException, Depends, Request, Response
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Dict, Any, Tuple
import uuid
from datetime import datetime, timezone, timedelta
import httpx
from passlib.context import CryptContext
from jose import jwt, JWTError
import base64
from io import BytesIO
from PIL import Image, ImageEnhance, ImageFilter, ImageDraw
import re
import cv2
import numpy as np
# Email service
from email_service import (
    send_welcome_email,
    send_verification_email,
    send_password_reset_email,
    send_purchase_confirmation_email,
    send_premium_welcome_email,
    generate_verification_code,
    generate_reset_code
)
# Note: emergentintegrations was removed for Railway deployment compatibility
# Using pytesseract (Tesseract OCR) as a public alternative
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
import boto3
from botocore.exceptions import ClientError
import certifi

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection with SSL certificate for Atlas
mongo_url = os.environ.get('MONGO_URL')
if not mongo_url:
    raise ValueError("MONGO_URL environment variable is required")
client = AsyncIOMotorClient(mongo_url, tlsCAFile=certifi.where() if 'mongodb+srv' in mongo_url else None)
db = client[os.environ.get('DB_NAME', 'scanup')]  # Default to 'scanup' if not set

# Emergent LLM key for OCR
EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY", "")

# AWS S3 Configuration
AWS_ACCESS_KEY_ID = os.environ.get("AWS_ACCESS_KEY_ID", "")
AWS_SECRET_ACCESS_KEY = os.environ.get("AWS_SECRET_ACCESS_KEY", "")
AWS_S3_BUCKET_NAME = os.environ.get("AWS_S3_BUCKET_NAME", os.environ.get("AWS_BUCKET_NAME", "scanup-documents"))
# Support both AWS_REGION and AWS_S3_REGION for flexibility
AWS_REGION = os.environ.get("AWS_REGION", os.environ.get("AWS_S3_REGION", "us-east-1"))

# S3 client will be initialized after logger
s3_client = None

# Create the main app without a prefix
app = FastAPI()

# Add validation error handler to log 422 errors
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    logger.error(f"[422 Validation Error] URL: {request.url}")
    logger.error(f"[422 Validation Error] Errors: {exc.errors()}")
    # Log body if small enough (don't log huge base64)
    try:
        body = await request.body()
        body_str = body.decode('utf-8')[:500] if body else 'empty'
        logger.error(f"[422 Validation Error] Body preview: {body_str}...")
    except:
        pass
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors()}
    )

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
JWT_SECRET = os.environ.get("JWT_SECRET", "your-super-secret-key-change-in-production")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_DAYS = 7

security = HTTPBearer(auto_error=False)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize S3 client after logger is configured
if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY:
    s3_client = boto3.client(
        's3',
        aws_access_key_id=AWS_ACCESS_KEY_ID,
        aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        region_name=AWS_REGION
    )
    logger.info(f"✅ AWS S3 initialized: bucket={AWS_S3_BUCKET_NAME}, region={AWS_REGION}")
else:
    logger.warning("⚠️ AWS S3 not configured - images will be stored in MongoDB")

# ==================== MODELS ====================

# User Models
class UserBase(BaseModel):
    email: EmailStr
    name: str
    picture: Optional[str] = None

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    photo_url: Optional[str] = None
    avatar_url: Optional[str] = None
    subscription_type: str = "free"  # free, premium, trial
    subscription_expires_at: Optional[datetime] = None
    ocr_usage_today: int = 0
    ocr_usage_date: Optional[str] = None
    # Scan limits
    scans_today: int = 0
    scans_this_month: int = 0
    last_scan_date: Optional[str] = None
    scan_month: Optional[str] = None  # Format: "2025-01"
    # Trial period
    trial_start_date: Optional[datetime] = None
    trial_used: bool = False
    # Email verification
    email_verified: bool = False
    # Timestamps
    created_at: datetime
    updated_at: Optional[datetime] = None

# Free tier limits
FREE_SCANS_PER_DAY = 10
FREE_SCANS_PER_MONTH = 100
FREE_OCR_PER_DAY = 3
TRIAL_DURATION_DAYS = 7

class UserResponse(BaseModel):
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    photo_url: Optional[str] = None
    avatar_url: Optional[str] = None
    subscription_type: str
    is_premium: bool
    is_trial: bool = False
    trial_days_remaining: int = 0
    ocr_remaining_today: int
    scans_remaining_today: int
    scans_remaining_month: int

class AuthResponse(BaseModel):
    user: UserResponse
    token: str

# Document Models
class PageData(BaseModel):
    page_id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    image_base64: Optional[str] = None  # Base64 image (MongoDB storage)
    image_url: Optional[str] = None      # S3 URL (cloud storage)
    thumbnail_base64: Optional[str] = None
    thumbnail_url: Optional[str] = None  # S3 thumbnail URL
    original_image_base64: Optional[str] = None  # For non-destructive editing
    ocr_text: Optional[str] = None
    filter_applied: str = "original"  # original, grayscale, bw, enhanced
    rotation: int = 0
    order: int = 0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class DocumentCreate(BaseModel):
    name: str
    folder_id: Optional[str] = None
    tags: List[str] = []
    pages: List[PageData] = []

class DocumentUpdate(BaseModel):
    name: Optional[str] = None
    folder_id: Optional[str] = None
    tags: Optional[List[str]] = None
    pages: Optional[List[PageData]] = None
    is_password_protected: Optional[bool] = None
    password_hash: Optional[str] = None

class Document(BaseModel):
    document_id: str
    user_id: str
    name: str
    folder_id: Optional[str] = None
    tags: List[str] = []
    pages: List[PageData] = []
    ocr_full_text: Optional[str] = None
    is_password_protected: bool = False
    storage_type: Optional[str] = None  # 's3' or 'mongodb'
    has_watermark: Optional[bool] = None
    created_at: datetime
    updated_at: datetime

# Folder Models
class FolderCreate(BaseModel):
    name: str
    color: str = "#3B82F6"
    parent_id: Optional[str] = None

class Folder(BaseModel):
    folder_id: str
    user_id: str
    name: str
    color: str
    parent_id: Optional[str] = None
    is_protected: bool = False
    password_hash: Optional[str] = None
    created_at: datetime

# OCR Models
class OCRRequest(BaseModel):
    image_base64: str
    language: str = "en"

class OCRResponse(BaseModel):
    text: str
    confidence: Optional[float] = None

# Image Processing Models
class ImageProcessRequest(BaseModel):
    image_base64: str
    operation: str  # crop, rotate, filter
    params: Dict[str, Any] = {}

class ImageProcessResponse(BaseModel):
    processed_image_base64: str

# Subscription Models
class SubscriptionUpdate(BaseModel):
    subscription_type: str  # free, premium
    duration_days: Optional[int] = 30

# Session Models for OAuth
class SessionDataResponse(BaseModel):
    id: str
    email: str
    name: str
    picture: Optional[str] = None
    session_token: str

# ==================== HELPER FUNCTIONS ====================

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_jwt_token(user_id: str) -> str:
    expiration = datetime.now(timezone.utc) + timedelta(days=JWT_EXPIRATION_DAYS)
    payload = {
        "sub": user_id,
        "exp": expiration,
        "iat": datetime.now(timezone.utc)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def decode_jwt_token(token: str) -> Optional[str]:
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload.get("sub")
    except JWTError:
        return None

async def get_current_user(request: Request, credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)) -> User:
    token = None
    
    # Check Authorization header first
    if credentials:
        token = credentials.credentials
    
    # Check cookies as fallback
    if not token:
        token = request.cookies.get("session_token")
    
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    # First try JWT token
    user_id = decode_jwt_token(token)
    
    if user_id:
        user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if user_doc:
            return User(**user_doc)
    
    # Then try session token (OAuth)
    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if session:
        expires_at = session.get("expires_at")
        if expires_at:
            if expires_at.tzinfo is None:
                expires_at = expires_at.replace(tzinfo=timezone.utc)
            if expires_at > datetime.now(timezone.utc):
                user_doc = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
                if user_doc:
                    return User(**user_doc)
    
    raise HTTPException(status_code=401, detail="Invalid or expired token")

def user_to_response(user: User) -> UserResponse:
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    current_month = datetime.now(timezone.utc).strftime("%Y-%m")
    
    # Check if user is in trial period
    is_trial = False
    trial_days_remaining = 0
    if user.subscription_type == "trial" and user.trial_start_date:
        trial_end = user.trial_start_date + timedelta(days=TRIAL_DURATION_DAYS)
        if trial_end.tzinfo is None:
            trial_end = trial_end.replace(tzinfo=timezone.utc)
        if datetime.now(timezone.utc) < trial_end:
            is_trial = True
            trial_days_remaining = (trial_end - datetime.now(timezone.utc)).days
    
    # Check premium status
    is_premium = user.subscription_type == "premium" or is_trial
    if user.subscription_expires_at and user.subscription_type == "premium":
        expires_at = user.subscription_expires_at
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        is_premium = expires_at > datetime.now(timezone.utc)
    
    # Calculate OCR remaining
    ocr_remaining = FREE_OCR_PER_DAY
    if is_premium:
        ocr_remaining = 9999  # Unlimited for premium
    elif user.ocr_usage_date == today:
        ocr_remaining = max(0, FREE_OCR_PER_DAY - user.ocr_usage_today)
    
    # Calculate scans remaining today
    scans_today = 0
    if user.last_scan_date == today:
        scans_today = user.scans_today
    scans_remaining_today = FREE_SCANS_PER_DAY - scans_today if not is_premium else 9999
    
    # Calculate scans remaining this month
    scans_month = 0
    if user.scan_month == current_month:
        scans_month = user.scans_this_month
    scans_remaining_month = FREE_SCANS_PER_MONTH - scans_month if not is_premium else 9999
    
    return UserResponse(
        user_id=user.user_id,
        email=user.email,
        name=user.name,
        picture=user.picture,
        photo_url=getattr(user, 'photo_url', None) or getattr(user, 'avatar_url', None) or user.picture,
        avatar_url=getattr(user, 'avatar_url', None) or getattr(user, 'photo_url', None) or user.picture,
        subscription_type=user.subscription_type,
        is_premium=is_premium,
        is_trial=is_trial,
        trial_days_remaining=trial_days_remaining,
        ocr_remaining_today=ocr_remaining,
        scans_remaining_today=max(0, scans_remaining_today),
        scans_remaining_month=max(0, scans_remaining_month)
    )

def add_watermark(image_base64: str, watermark_text: str = "ScanUp") -> str:
    """Add a single watermark to image center for free users"""
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        image = Image.open(BytesIO(image_data)).convert('RGBA')
        
        # Create watermark layer
        watermark = Image.new('RGBA', image.size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(watermark)
        
        # Calculate font size based on image size (reasonable size, not too big)
        font_size = max(30, min(image.width, image.height) // 15)
        
        # Try to use a built-in font, fallback to default
        try:
            from PIL import ImageFont
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", font_size)
        except:
            font = ImageFont.load_default()
        
        # Get text size
        text = watermark_text
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        # Single watermark in the center of the image
        x = (image.width - text_width) // 2
        y = (image.height - text_height) // 2
        
        # Draw shadow/outline first for visibility on any background
        shadow_offset = 2
        draw.text((x + shadow_offset, y + shadow_offset), text, font=font, fill=(0, 0, 0, 80))
        
        # Draw main watermark text (semi-transparent gray)
        draw.text((x, y), text, font=font, fill=(128, 128, 128, 120))
        
        # Composite
        watermarked = Image.alpha_composite(image, watermark)
        
        # Convert to RGB for JPEG
        if watermarked.mode == 'RGBA':
            background = Image.new('RGB', watermarked.size, (255, 255, 255))
            background.paste(watermarked, mask=watermarked.split()[3])
            watermarked = background
        
        # Encode
        buffer = BytesIO()
        watermarked.save(buffer, format='JPEG', quality=95)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"Watermark error: {str(e)}")
        return image_base64  # Return original if watermarking fails

# ==================== AWS S3 FUNCTIONS ====================

def upload_to_s3(image_base64: str, user_id: str, document_id: str, page_id: str, image_type: str = "page") -> Optional[str]:
    """
    Upload image to S3 and return the URL.
    
    Args:
        image_base64: Base64 encoded image
        user_id: User ID
        document_id: Document ID
        page_id: Page ID
        image_type: "page" or "thumbnail"
    
    Returns:
        S3 URL or None if upload fails
    """
    if not s3_client:
        logger.warning("S3 not configured, returning None")
        return None
    
    try:
        # Clean base64 string
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        # Decode image
        image_data = base64.b64decode(image_base64)
        
        # Generate S3 key
        s3_key = f"users/{user_id}/documents/{document_id}/{image_type}_{page_id}.jpg"
        
        # Upload to S3
        s3_client.put_object(
            Bucket=AWS_S3_BUCKET_NAME,
            Key=s3_key,
            Body=image_data,
            ContentType='image/jpeg'
        )
        
        # Generate URL
        s3_url = f"https://{AWS_S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
        logger.info(f"✅ Uploaded to S3: {s3_key}")
        
        return s3_url
    except ClientError as e:
        logger.error(f"❌ S3 upload error: {e}")
        return None
    except Exception as e:
        logger.error(f"❌ S3 upload error: {e}")
        return None


def delete_from_s3(user_id: str, document_id: str, page_id: Optional[str] = None) -> bool:
    """
    Delete image(s) from S3.
    
    Args:
        user_id: User ID
        document_id: Document ID
        page_id: Page ID (if None, deletes all pages for document)
    
    Returns:
        True if successful
    """
    if not s3_client:
        return False
    
    try:
        if page_id:
            # Delete specific page
            s3_key = f"users/{user_id}/documents/{document_id}/page_{page_id}.jpg"
            s3_client.delete_object(Bucket=AWS_S3_BUCKET_NAME, Key=s3_key)
            # Also delete thumbnail
            thumb_key = f"users/{user_id}/documents/{document_id}/thumbnail_{page_id}.jpg"
            s3_client.delete_object(Bucket=AWS_S3_BUCKET_NAME, Key=thumb_key)
        else:
            # Delete all objects for document
            prefix = f"users/{user_id}/documents/{document_id}/"
            response = s3_client.list_objects_v2(Bucket=AWS_S3_BUCKET_NAME, Prefix=prefix)
            
            if 'Contents' in response:
                for obj in response['Contents']:
                    s3_client.delete_object(Bucket=AWS_S3_BUCKET_NAME, Key=obj['Key'])
        
        logger.info(f"✅ Deleted from S3: {user_id}/{document_id}/{page_id or 'all'}")
        return True
    except Exception as e:
        logger.error(f"❌ S3 delete error: {e}")
        return False


def create_thumbnail(image_base64: str, max_size: int = 300) -> str:
    """Create a thumbnail from image base64"""
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        img = Image.open(BytesIO(image_data))
        
        # Calculate thumbnail size maintaining aspect ratio
        ratio = min(max_size / img.width, max_size / img.height)
        new_size = (int(img.width * ratio), int(img.height * ratio))
        
        img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        # Convert to RGB if necessary
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        
        buffer = BytesIO()
        img.save(buffer, format='JPEG', quality=70)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        logger.error(f"Thumbnail creation error: {e}")
        return image_base64[:1000]  # Return truncated original as fallback


async def check_scan_limits(user: User) -> Tuple[bool, str]:
    """Check if user can scan. Returns (can_scan, message)"""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    current_month = datetime.now(timezone.utc).strftime("%Y-%m")
    
    # Premium users have no limits
    is_premium = user.subscription_type in ["premium", "trial"]
    if is_premium:
        if user.subscription_type == "trial" and user.trial_start_date:
            trial_end = user.trial_start_date + timedelta(days=TRIAL_DURATION_DAYS)
            if datetime.now(timezone.utc) >= trial_end.replace(tzinfo=timezone.utc):
                is_premium = False
    
    if is_premium:
        return True, ""
    
    # Check daily limit
    scans_today = user.scans_today if user.last_scan_date == today else 0
    if scans_today >= FREE_SCANS_PER_DAY:
        return False, f"Daily scan limit reached ({FREE_SCANS_PER_DAY} scans/day). Upgrade to Premium for unlimited scans."
    
    # Check monthly limit
    scans_month = user.scans_this_month if user.scan_month == current_month else 0
    if scans_month >= FREE_SCANS_PER_MONTH:
        return False, f"Monthly scan limit reached ({FREE_SCANS_PER_MONTH} scans/month). Upgrade to Premium for unlimited scans."
    
    return True, ""

async def increment_scan_count(user_id: str):
    """Increment user's scan count"""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    current_month = datetime.now(timezone.utc).strftime("%Y-%m")
    
    user = await db.users.find_one({"user_id": user_id})
    if not user:
        return
    
    # Reset daily count if new day
    scans_today = user.get("scans_today", 0)
    if user.get("last_scan_date") != today:
        scans_today = 0
    
    # Reset monthly count if new month
    scans_month = user.get("scans_this_month", 0)
    if user.get("scan_month") != current_month:
        scans_month = 0
    
    await db.users.update_one(
        {"user_id": user_id},
        {
            "$set": {
                "scans_today": scans_today + 1,
                "last_scan_date": today,
                "scans_this_month": scans_month + 1,
                "scan_month": current_month,
                "updated_at": datetime.now(timezone.utc)
            }
        }
    )

def create_thumbnail(image_base64: str, max_size: int = 200) -> str:
    """Create a thumbnail from base64 image"""
    try:
        # Decode base64
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        image = Image.open(BytesIO(image_data))
        
        # Convert RGBA to RGB if needed
        if image.mode == 'RGBA':
            background = Image.new('RGB', image.size, (255, 255, 255))
            background.paste(image, mask=image.split()[3])
            image = background
        
        # Create thumbnail
        image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        
        # Convert back to base64
        buffer = BytesIO()
        image.save(buffer, format="JPEG", quality=70)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        logger.error(f"Error creating thumbnail: {e}")
        return image_base64

def convert_to_rgb(image: Image.Image) -> Image.Image:
    """Convert image to RGB mode"""
    if image.mode == 'RGBA':
        background = Image.new('RGB', image.size, (255, 255, 255))
        background.paste(image, mask=image.split()[3])
        return background
    elif image.mode != 'RGB':
        return image.convert('RGB')
    return image

def apply_image_filter(
    image_base64: str, 
    filter_type: str, 
    brightness: int = 0, 
    contrast: int = 0, 
    saturation: int = 0
) -> str:
    """Apply filter and adjustments to image
    
    Args:
        image_base64: Base64 encoded image
        filter_type: Filter preset ('original', 'grayscale', 'bw', 'enhanced', 'document')
        brightness: Adjustment from -50 to +50
        contrast: Adjustment from -50 to +50
        saturation: Adjustment from -50 to +50
    """
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        image = Image.open(BytesIO(image_data))
        image = convert_to_rgb(image)
        
        # Apply filter preset first
        if filter_type == "grayscale":
            image = image.convert("L").convert("RGB")
        elif filter_type == "bw":
            image = image.convert("L")
            image = image.point(lambda x: 0 if x < 128 else 255, '1')
            image = image.convert("RGB")
        elif filter_type == "enhanced":
            # Enhance contrast and sharpness
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.3)
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.5)
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(1.1)
        elif filter_type == "document":
            # Document mode - high contrast, sharpen
            image = image.convert("L")  # Grayscale
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            image = image.filter(ImageFilter.SHARPEN)
            image = image.convert("RGB")
        
        # Apply manual adjustments (brightness, contrast, saturation)
        # Convert adjustment values (-50 to +50) to enhancement factors (0.5 to 1.5)
        
        if brightness != 0:
            # Brightness: -50 -> 0.5, 0 -> 1.0, +50 -> 1.5
            brightness_factor = 1.0 + (brightness / 100.0)
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(brightness_factor)
        
        if contrast != 0:
            # Contrast: -50 -> 0.5, 0 -> 1.0, +50 -> 1.5
            contrast_factor = 1.0 + (contrast / 100.0)
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(contrast_factor)
        
        if saturation != 0:
            # Saturation: -50 -> 0.5, 0 -> 1.0, +50 -> 1.5
            saturation_factor = 1.0 + (saturation / 100.0)
            enhancer = ImageEnhance.Color(image)
            image = enhancer.enhance(saturation_factor)
        
        buffer = BytesIO()
        image.save(buffer, format="JPEG", quality=90)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        logger.error(f"Error applying filter: {e}")
        return image_base64

def rotate_image(image_base64: str, degrees: int) -> str:
    """Rotate image by degrees"""
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        image = Image.open(BytesIO(image_data))
        image = convert_to_rgb(image)
        
        # Rotate image (negative because PIL rotates counter-clockwise)
        image = image.rotate(-degrees, expand=True, fillcolor='white')
        
        buffer = BytesIO()
        image.save(buffer, format="JPEG", quality=90)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        logger.error(f"Error rotating image: {e}")
        return image_base64

def crop_image(image_base64: str, x: int, y: int, width: int, height: int) -> str:
    """Crop image"""
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        image = Image.open(BytesIO(image_data))
        image = convert_to_rgb(image)
        
        # Crop image
        image = image.crop((x, y, x + width, y + height))
        
        buffer = BytesIO()
        image.save(buffer, format="JPEG", quality=90)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        logger.error(f"Error cropping image: {e}")
        return image_base64

def detect_document_edges(image_base64: str, document_type: str = "document") -> Dict[str, Any]:
    """Detect document edges using OpenCV with improved multi-approach detection
    
    This function tries multiple detection strategies to find document edges:
    1. Canny edge detection with various thresholds
    2. Adaptive thresholding
    3. Color-based detection (for white documents on darker backgrounds)
    4. Morphological operations
    """
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        nparr = np.frombuffer(image_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return {"detected": False, "corners": None, "message": "Could not decode image"}
        
        height, width = img.shape[:2]
        
        # Resize for faster processing
        max_dim = 1000  # Increased for better accuracy
        scale = min(1.0, max_dim / max(width, height))
        if scale < 1.0:
            img_small = cv2.resize(img, None, fx=scale, fy=scale)
        else:
            img_small = img.copy()
            scale = 1.0
        
        small_h, small_w = img_small.shape[:2]
        
        # Minimum area threshold - lowered to 3% for better detection
        min_area = small_h * small_w * 0.03
        # Maximum area - no more than 99% of image
        max_area = small_h * small_w * 0.99
        
        best_contour = None
        best_score = 0
        
        def score_contour(contour, approx):
            """Score a contour based on multiple criteria"""
            area = cv2.contourArea(contour)
            if area < min_area or area > max_area:
                return 0
            
            # Must be quadrilateral
            if len(approx) != 4:
                return 0
            
            # Check if it's convex
            if not cv2.isContourConvex(approx):
                return 0
            
            # Score based on area (larger is better, but not too large)
            area_score = area / (small_h * small_w)
            
            # Score based on how rectangular it is (aspect ratio)
            rect = cv2.minAreaRect(approx)
            rect_area = rect[1][0] * rect[1][1]
            if rect_area > 0:
                rectangularity = area / rect_area
            else:
                rectangularity = 0
            
            # Combined score
            return area_score * 0.6 + rectangularity * 0.4
        
        # Convert to different color spaces
        gray = cv2.cvtColor(img_small, cv2.COLOR_BGR2GRAY)
        hsv = cv2.cvtColor(img_small, cv2.COLOR_BGR2HSV)
        
        # APPROACH 1: Multiple Canny edge detection with different preprocessing
        preprocessed_images = [
            cv2.GaussianBlur(gray, (5, 5), 0),
            cv2.bilateralFilter(gray, 9, 75, 75),
            cv2.medianBlur(gray, 5),
        ]
        
        canny_params = [(20, 80), (30, 100), (50, 150), (75, 200), (100, 250)]
        
        for prep_img in preprocessed_images:
            for low_thresh, high_thresh in canny_params:
                edges = cv2.Canny(prep_img, low_thresh, high_thresh)
                
                # Dilate to connect broken lines
                kernel = np.ones((3, 3), np.uint8)
                edges = cv2.dilate(edges, kernel, iterations=2)
                edges = cv2.erode(edges, kernel, iterations=1)
                
                contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                for contour in contours:
                    for epsilon_factor in [0.01, 0.02, 0.03, 0.04]:
                        epsilon = epsilon_factor * cv2.arcLength(contour, True)
                        approx = cv2.approxPolyDP(contour, epsilon, True)
                        
                        score = score_contour(contour, approx)
                        if score > best_score:
                            best_score = score
                            best_contour = approx
        
        # APPROACH 2: Adaptive thresholding
        if best_score < 0.3:
            for block_size in [11, 15, 21]:
                for c in [2, 5, 10]:
                    thresh = cv2.adaptiveThreshold(
                        cv2.GaussianBlur(gray, (5, 5), 0), 
                        255, 
                        cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                        cv2.THRESH_BINARY, 
                        block_size, 
                        c
                    )
                    
                    # Morphological operations
                    kernel = np.ones((5, 5), np.uint8)
                    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
                    thresh = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)
                    
                    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                    
                    for contour in contours:
                        for epsilon_factor in [0.01, 0.02, 0.03]:
                            epsilon = epsilon_factor * cv2.arcLength(contour, True)
                            approx = cv2.approxPolyDP(contour, epsilon, True)
                            
                            score = score_contour(contour, approx)
                            if score > best_score:
                                best_score = score
                                best_contour = approx
        
        # APPROACH 3: Color-based detection for white/light documents
        if best_score < 0.3:
            # Detect bright regions (documents are usually white/light)
            _, bright_mask = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY)
            
            # Also try detecting by saturation (documents have low saturation)
            s_channel = hsv[:, :, 1]
            _, low_sat_mask = cv2.threshold(s_channel, 50, 255, cv2.THRESH_BINARY_INV)
            
            for mask in [bright_mask, low_sat_mask]:
                kernel = np.ones((7, 7), np.uint8)
                mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel)
                mask = cv2.morphologyEx(mask, cv2.MORPH_OPEN, kernel)
                
                contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                for contour in contours:
                    for epsilon_factor in [0.01, 0.02, 0.03]:
                        epsilon = epsilon_factor * cv2.arcLength(contour, True)
                        approx = cv2.approxPolyDP(contour, epsilon, True)
                        
                        score = score_contour(contour, approx)
                        if score > best_score:
                            best_score = score
                            best_contour = approx
        
        if best_contour is not None and best_score > 0.05:
            # Scale corners back to original size
            corners = []
            for point in best_contour:
                x, y = point[0]
                corners.append({
                    "x": int(x / scale),
                    "y": int(y / scale)
                })
            
            # Order corners: top-left, top-right, bottom-right, bottom-left
            corners = order_corners(corners)
            
            return {
                "detected": True,
                "corners": corners,
                "width": width,
                "height": height,
                "confidence": round(best_score, 2)
            }
        
        # Return default corners if no detection - full frame with slight padding
        padding = 0.05
        default_corners = [
            {"x": int(width * padding), "y": int(height * padding)},
            {"x": int(width * (1 - padding)), "y": int(height * padding)},
            {"x": int(width * (1 - padding)), "y": int(height * (1 - padding))},
            {"x": int(width * padding), "y": int(height * (1 - padding))}
        ]
        
        return {
            "detected": False, 
            "corners": default_corners,
            "width": width,
            "height": height,
            "message": "Could not detect document edges. Default crop area provided."
        }
    except Exception as e:
        logger.error(f"Error detecting edges: {e}")
        return {"detected": False, "corners": None, "message": str(e)}

def fix_image_orientation_pil(image_bytes):
    """
    Fix image orientation based on EXIF data using PIL.
    Returns corrected image bytes with proper orientation.
    
    Args:
        image_bytes: Original image bytes
    Returns:
        Tuple of (corrected_bytes, was_rotated, original_orientation)
    """
    try:
        from PIL import Image as PILImage
        from PIL import ExifTags
        from io import BytesIO
        
        pil_img = PILImage.open(BytesIO(image_bytes))
        original_size = pil_img.size
        orientation = 1  # Default: normal
        
        # Try to get EXIF orientation
        try:
            exif = pil_img._getexif()
            if exif:
                for tag, value in exif.items():
                    if ExifTags.TAGS.get(tag) == 'Orientation':
                        orientation = value
                        logger.info(f"EXIF Orientation found: {value} (size before: {original_size})")
                        
                        if value == 3:
                            pil_img = pil_img.rotate(180, expand=True)
                            logger.info("Applied 180 degree rotation")
                        elif value == 6:
                            pil_img = pil_img.rotate(-90, expand=True)  # 90 CW = -90 in PIL
                            logger.info("Applied 90 degree clockwise rotation")
                        elif value == 8:
                            pil_img = pil_img.rotate(90, expand=True)  # 90 CCW = 90 in PIL
                            logger.info("Applied 90 degree counter-clockwise rotation")
                        
                        logger.info(f"Size after rotation: {pil_img.size}")
                        break
        except (AttributeError, KeyError, IndexError) as e:
            logger.debug(f"No EXIF orientation found: {e}")
        
        # Convert back to bytes
        buffer = BytesIO()
        # Save as JPEG without EXIF to avoid re-rotation
        if pil_img.mode == 'RGBA':
            pil_img = pil_img.convert('RGB')
        pil_img.save(buffer, format='JPEG', quality=95)
        corrected_bytes = buffer.getvalue()
        
        was_rotated = orientation in [3, 6, 8]
        return corrected_bytes, was_rotated, orientation
        
    except Exception as e:
        logger.error(f"Error fixing orientation: {e}")
        return image_bytes, False, 1


def fix_image_orientation(img, original_bytes=None):
    """
    Fix image orientation based on EXIF data.
    OpenCV doesn't handle EXIF orientation automatically, so we need to do it manually.
    
    Args:
        img: OpenCV image (numpy array)
        original_bytes: Original image bytes (to read EXIF from)
    """
    try:
        from PIL import Image as PILImage
        from PIL import ExifTags
        from io import BytesIO
        
        # Try to get EXIF from original bytes first (more reliable)
        pil_img = None
        if original_bytes is not None:
            try:
                pil_img = PILImage.open(BytesIO(original_bytes))
            except:
                pass
        
        # Fallback: Convert OpenCV image to PIL (may lose EXIF)
        if pil_img is None:
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pil_img = PILImage.fromarray(img_rgb)
        
        # Try to get EXIF orientation
        try:
            exif = pil_img._getexif()
            if exif:
                for tag, value in exif.items():
                    if ExifTags.TAGS.get(tag) == 'Orientation':
                        logger.info(f"EXIF Orientation tag found: {value}")
                        if value == 3:
                            img = cv2.rotate(img, cv2.ROTATE_180)
                            logger.info("Applied 180 degree rotation")
                        elif value == 6:
                            img = cv2.rotate(img, cv2.ROTATE_90_CLOCKWISE)
                            logger.info("Applied 90 degree clockwise rotation")
                        elif value == 8:
                            img = cv2.rotate(img, cv2.ROTATE_90_COUNTERCLOCKWISE)
                            logger.info("Applied 90 degree counter-clockwise rotation")
                        break
        except (AttributeError, KeyError, IndexError) as e:
            logger.debug(f"No EXIF orientation found: {e}")
    except Exception as e:
        logger.debug(f"Could not process EXIF: {e}")
    
    return img

def order_corners(corners: List[Dict]) -> List[Dict]:
    """
    Order corners in consistent order: TL, TR, BR, BL
    This ensures perspective transform works correctly regardless of input order.
    """
    if len(corners) != 4:
        return corners
    
    # Convert to numpy for easier manipulation
    pts = np.array([[c['x'], c['y']] for c in corners], dtype=np.float32)
    
    # Sort by y-coordinate to get top and bottom pairs
    sorted_by_y = pts[np.argsort(pts[:, 1])]
    top_pts = sorted_by_y[:2]
    bottom_pts = sorted_by_y[2:]
    
    # Sort top points by x to get TL, TR
    top_sorted = top_pts[np.argsort(top_pts[:, 0])]
    tl, tr = top_sorted[0], top_sorted[1]
    
    # Sort bottom points by x to get BL, BR
    bottom_sorted = bottom_pts[np.argsort(bottom_pts[:, 0])]
    bl, br = bottom_sorted[0], bottom_sorted[1]
    
    return [
        {'x': float(tl[0]), 'y': float(tl[1])},  # Top-Left
        {'x': float(tr[0]), 'y': float(tr[1])},  # Top-Right
        {'x': float(br[0]), 'y': float(br[1])},  # Bottom-Right
        {'x': float(bl[0]), 'y': float(bl[1])}   # Bottom-Left
    ]

def perspective_crop(image_base64: str, corners: List[Dict]) -> str:
    """
    Apply perspective transform to crop document and make it front-facing.
    
    This function:
    1. Takes the 4 corner points of a document in the image
    2. Applies a perspective warp to make the document rectangular
    3. The output looks like it was photographed from directly above (front-facing)
    
    Args:
        image_base64: Base64 encoded image
        corners: List of 4 corner points in order [TL, TR, BR, BL]
                 Can be in pixel coordinates or will be auto-ordered
    
    Returns:
        Base64 encoded cropped and perspective-corrected image
    """
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        image_data = base64.b64decode(image_base64)
        nparr = np.frombuffer(image_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None or len(corners) != 4:
            return image_base64
        
        # Fix EXIF orientation BEFORE processing
        img = fix_image_orientation(img)
        
        # Ensure corners are in correct order: TL, TR, BR, BL
        ordered_corners = order_corners(corners)
        
        # Source points - use float32 for precision
        src_pts = np.float32([
            [ordered_corners[0]["x"], ordered_corners[0]["y"]],  # top-left
            [ordered_corners[1]["x"], ordered_corners[1]["y"]],  # top-right
            [ordered_corners[2]["x"], ordered_corners[2]["y"]],  # bottom-right
            [ordered_corners[3]["x"], ordered_corners[3]["y"]]   # bottom-left
        ])
        
        # Calculate the actual document dimensions from the corners
        # Top edge width
        width_top = np.sqrt(
            (ordered_corners[1]["x"] - ordered_corners[0]["x"])**2 + 
            (ordered_corners[1]["y"] - ordered_corners[0]["y"])**2
        )
        # Bottom edge width
        width_bottom = np.sqrt(
            (ordered_corners[2]["x"] - ordered_corners[3]["x"])**2 + 
            (ordered_corners[2]["y"] - ordered_corners[3]["y"])**2
        )
        # Left edge height
        height_left = np.sqrt(
            (ordered_corners[3]["x"] - ordered_corners[0]["x"])**2 + 
            (ordered_corners[3]["y"] - ordered_corners[0]["y"])**2
        )
        # Right edge height
        height_right = np.sqrt(
            (ordered_corners[2]["x"] - ordered_corners[1]["x"])**2 + 
            (ordered_corners[2]["y"] - ordered_corners[1]["y"])**2
        )
        
        # Use the maximum dimensions to preserve content
        output_width = max(int(round(width_top)), int(round(width_bottom)))
        output_height = max(int(round(height_left)), int(round(height_right)))
        
        # Ensure minimum dimensions
        output_width = max(output_width, 100)
        output_height = max(output_height, 100)
        
        # For proper document aspect ratio, adjust dimensions
        # Standard A4 is ~1:1.414, ID cards are ~1.586:1
        # We use the natural aspect ratio from the selection
        
        # Destination points - perfect rectangle (front-facing view)
        dst_pts = np.float32([
            [0, 0],                          # top-left
            [output_width - 1, 0],           # top-right  
            [output_width - 1, output_height - 1],  # bottom-right
            [0, output_height - 1]           # bottom-left
        ])
        
        # Get perspective transform matrix
        matrix = cv2.getPerspectiveTransform(src_pts, dst_pts)
        
        # Apply perspective transform with high-quality interpolation
        # This makes the document appear as if photographed from directly above
        warped = cv2.warpPerspective(
            img, 
            matrix, 
            (output_width, output_height),
            flags=cv2.INTER_CUBIC,  # High quality interpolation
            borderMode=cv2.BORDER_CONSTANT,
            borderValue=(255, 255, 255)  # White border for any empty areas
        )
        
        # Optional: Apply slight sharpening to improve text clarity after transform
        # kernel = np.array([[-0.5, -0.5, -0.5], [-0.5, 5, -0.5], [-0.5, -0.5, -0.5]])
        # warped = cv2.filter2D(warped, -1, kernel)
        
        logger.info(f"Perspective crop: {img.shape[:2]} -> {warped.shape[:2]}, output={output_width}x{output_height}")
        
        # Encode back to base64 with maximum quality
        _, buffer = cv2.imencode('.jpg', warped, [cv2.IMWRITE_JPEG_QUALITY, 95])
        return base64.b64encode(buffer).decode()
        
    except Exception as e:
        logger.error(f"Error in perspective crop: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return image_base64

async def perform_ocr_with_tesseract(image_base64: str) -> str:
    """Perform OCR using Tesseract OCR (public, free alternative)
    
    This function uses pytesseract which works with the Tesseract OCR engine.
    It's a robust, free, and widely-used OCR solution that works on any platform.
    """
    if not TESSERACT_AVAILABLE:
        return "OCR service not available. Tesseract OCR is not installed."
    
    try:
        if "," in image_base64:
            image_base64 = image_base64.split(",")[1]
        
        # Decode base64 image
        image_data = base64.b64decode(image_base64)
        image = Image.open(BytesIO(image_data))
        
        # Convert to RGB if necessary (Tesseract works better with RGB)
        if image.mode in ('RGBA', 'P'):
            background = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'RGBA':
                background.paste(image, mask=image.split()[3])
            else:
                background.paste(image)
            image = background
        elif image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Preprocess image for better OCR results
        # Convert to grayscale
        image_np = np.array(image)
        gray = cv2.cvtColor(image_np, cv2.COLOR_RGB2GRAY)
        
        # Apply adaptive thresholding for better text detection
        # This helps with varying lighting conditions
        processed = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )
        
        # Convert back to PIL Image for Tesseract
        processed_image = Image.fromarray(processed)
        
        # Perform OCR with Tesseract
        # Configure Tesseract for best document scanning results
        custom_config = r'--oem 3 --psm 6'  # OEM 3 = default, PSM 6 = uniform block of text
        
        extracted_text = pytesseract.image_to_string(processed_image, config=custom_config)
        
        # Clean up the extracted text
        extracted_text = extracted_text.strip()
        
        if not extracted_text:
            # Try again with original image (sometimes preprocessing hurts)
            extracted_text = pytesseract.image_to_string(image, config=custom_config)
            extracted_text = extracted_text.strip()
        
        return extracted_text if extracted_text else "No text detected"
    
    except Exception as e:
        logger.error(f"OCR error: {e}")
        return f"OCR error: {str(e)}"


# Alias for backward compatibility
async def perform_ocr_with_openai(image_base64: str) -> str:
    """Backward compatible OCR function - now uses Tesseract"""
    return await perform_ocr_with_tesseract(image_base64)

# ==================== AUTH ENDPOINTS ====================

@api_router.post("/auth/register", response_model=AuthResponse)
async def register(user_data: UserCreate, response: Response):
    """Register a new user with email/password"""
    # Check if user exists
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    
    # Generate verification code
    verification_code = generate_verification_code()
    
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "name": user_data.name,
        "password_hash": hash_password(user_data.password),
        "picture": None,
        "subscription_type": "free",
        "subscription_expires_at": None,
        "ocr_usage_today": 0,
        "ocr_usage_date": None,
        "scans_today": 0,
        "scans_this_month": 0,
        "last_scan_date": None,
        "scan_month": None,
        "trial_start_date": None,
        "trial_used": False,
        "email_verified": False,
        "verification_code": verification_code,
        "verification_code_expires": now + timedelta(hours=24),
        "created_at": now,
        "updated_at": now
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_jwt_token(user_id)
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=JWT_EXPIRATION_DAYS * 24 * 60 * 60,
        path="/"
    )
    
    # Send welcome email (async, don't block)
    try:
        await send_welcome_email(user_data.email, user_data.name)
        await send_verification_email(user_data.email, user_data.name, verification_code)
        logger.info(f"Welcome and verification emails sent to {user_data.email}")
    except Exception as e:
        logger.error(f"Failed to send emails to {user_data.email}: {e}")
    
    user = User(**{k: v for k, v in user_doc.items() if k not in ["password_hash", "verification_code", "verification_code_expires"]})
    return AuthResponse(user=user_to_response(user), token=token)

@api_router.post("/auth/login", response_model=AuthResponse)
async def login(credentials: UserLogin, response: Response):
    """Login with email/password"""
    user_doc = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    
    if not user_doc or not user_doc.get("password_hash"):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    if not verify_password(credentials.password, user_doc["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    token = create_jwt_token(user_doc["user_id"])
    
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=JWT_EXPIRATION_DAYS * 24 * 60 * 60,
        path="/"
    )
    
    user = User(**{k: v for k, v in user_doc.items() if k not in ["password_hash", "verification_code", "verification_code_expires", "reset_code", "reset_code_expires"]})
    return AuthResponse(user=user_to_response(user), token=token)


# ============ GOOGLE OAUTH ============
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET")

class GoogleAuthRequest(BaseModel):
    credential: str  # Google ID token from frontend

@api_router.post("/auth/google", response_model=AuthResponse)
async def google_auth(data: GoogleAuthRequest):
    """Authenticate with Google OAuth - verify token and create/login user"""
    try:
        # Verify the Google ID token
        async with httpx.AsyncClient() as client:
            # Verify token with Google
            response = await client.get(
                f"https://oauth2.googleapis.com/tokeninfo?id_token={data.credential}"
            )
            
            if response.status_code != 200:
                logger.error(f"Google token verification failed: {response.status_code}")
                raise HTTPException(status_code=401, detail="Invalid Google token")
            
            google_data = response.json()
            logger.info(f"Google token data - aud: {google_data.get('aud')}, email: {google_data.get('email')}")
            logger.info(f"Expected GOOGLE_CLIENT_ID: {GOOGLE_CLIENT_ID}")
            
            # Verify the audience (client ID) - check if it matches our client ID
            token_aud = google_data.get("aud")
            if not GOOGLE_CLIENT_ID:
                logger.warning("GOOGLE_CLIENT_ID not set, skipping audience verification")
            elif token_aud != GOOGLE_CLIENT_ID:
                # Also check if it's a valid Google client ID format (might be from different platform)
                logger.warning(f"Token audience mismatch - got: {token_aud}, expected: {GOOGLE_CLIENT_ID}")
                # For now, allow it if email is verified (less strict for testing)
                if not google_data.get("email_verified"):
                    raise HTTPException(status_code=401, detail="Token not issued for this application")
            
            email = google_data.get("email")
            name = google_data.get("name", "")
            picture = google_data.get("picture", "")
            google_id = google_data.get("sub")
            
            if not email:
                raise HTTPException(status_code=400, detail="Email not provided by Google")
            
            # Check if user exists
            existing_user = await db.users.find_one({"email": email})
            now = datetime.now(timezone.utc)
            
            if existing_user:
                # User exists - update Google info and login
                await db.users.update_one(
                    {"email": email},
                    {"$set": {
                        "google_id": google_id,
                        "profile_picture": picture or existing_user.get("profile_picture"),
                        "name": name or existing_user.get("name"),
                        "email_verified": True,
                        "last_login": now,
                        "updated_at": now
                    }}
                )
                user_id = existing_user["user_id"]
                is_new = False
            else:
                # Create new user
                user_id = str(uuid.uuid4())
                new_user = {
                    "user_id": user_id,
                    "email": email,
                    "name": name,
                    "password_hash": None,  # No password for OAuth users
                    "google_id": google_id,
                    "profile_picture": picture,
                    "email_verified": True,
                    "is_premium": False,
                    "subscription_type": "free",
                    "created_at": now,
                    "updated_at": now,
                    "last_login": now
                }
                await db.users.insert_one(new_user)
                is_new = True
                
                # Send welcome email
                try:
                    await send_welcome_email(email, name or "there")
                except Exception as e:
                    logger.error(f"Failed to send welcome email: {e}")
            
            # Generate JWT token
            token = create_access_token({"sub": user_id, "email": email})
            
            # Get updated user data
            user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
            
            return AuthResponse(
                access_token=token,
                token_type="bearer",
                user=UserResponse(**user_doc)
            )
            
    except httpx.RequestError as e:
        logger.error(f"Google auth request error: {e}")
        raise HTTPException(status_code=500, detail="Failed to verify Google token")
    except HTTPException:
        raise  # Re-raise HTTP exceptions as-is
    except Exception as e:
        import traceback
        logger.error(f"Google auth error: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Google auth failed: {str(e)}")


# Alternative: Google OAuth with authorization code flow (for web)
class GoogleCodeRequest(BaseModel):
    code: str
    redirect_uri: str

@api_router.post("/auth/google/code", response_model=AuthResponse)
async def google_auth_code(data: GoogleCodeRequest):
    """Authenticate with Google OAuth using authorization code"""
    try:
        async with httpx.AsyncClient() as client:
            # Exchange code for tokens
            token_response = await client.post(
                "https://oauth2.googleapis.com/token",
                data={
                    "code": data.code,
                    "client_id": GOOGLE_CLIENT_ID,
                    "client_secret": GOOGLE_CLIENT_SECRET,
                    "redirect_uri": data.redirect_uri,
                    "grant_type": "authorization_code"
                }
            )
            
            if token_response.status_code != 200:
                error_data = token_response.json()
                logger.error(f"Google token exchange failed: {error_data}")
                raise HTTPException(status_code=401, detail="Failed to exchange authorization code")
            
            tokens = token_response.json()
            id_token = tokens.get("id_token")
            
            if not id_token:
                raise HTTPException(status_code=401, detail="No ID token received")
            
            # Verify and process the ID token
            return await google_auth(GoogleAuthRequest(credential=id_token))
            
    except httpx.RequestError as e:
        logger.error(f"Google auth code request error: {e}")
        raise HTTPException(status_code=500, detail="Failed to exchange authorization code")


# Password Reset Request
class PasswordResetRequest(BaseModel):
    email: EmailStr

@api_router.post("/auth/forgot-password")
async def forgot_password(data: PasswordResetRequest):
    """Request a password reset code"""
    user_doc = await db.users.find_one({"email": data.email}, {"_id": 0})
    
    # Don't reveal if email exists or not
    if not user_doc:
        return {"message": "If this email exists, a reset code will be sent"}
    
    # Generate reset code
    reset_code = generate_reset_code()
    now = datetime.now(timezone.utc)
    
    # Store reset code in database
    await db.users.update_one(
        {"email": data.email},
        {"$set": {
            "reset_code": reset_code,
            "reset_code_expires": now + timedelta(hours=1),
            "updated_at": now
        }}
    )
    
    # Send reset email
    try:
        user_name = user_doc.get("name", "User")
        await send_password_reset_email(data.email, user_name, reset_code)
        logger.info(f"Password reset email sent to {data.email}")
    except Exception as e:
        logger.error(f"Failed to send reset email to {data.email}: {e}")
    
    return {"message": "If this email exists, a reset code will be sent"}


# Password Reset Confirm
class PasswordResetConfirm(BaseModel):
    email: EmailStr
    code: str
    new_password: str

@api_router.post("/auth/reset-password")
async def reset_password(data: PasswordResetConfirm):
    """Reset password using the code"""
    user_doc = await db.users.find_one({"email": data.email}, {"_id": 0})
    
    if not user_doc:
        raise HTTPException(status_code=400, detail="Invalid email or code")
    
    # Check code
    stored_code = user_doc.get("reset_code")
    expires = user_doc.get("reset_code_expires")
    
    if not stored_code or stored_code != data.code:
        raise HTTPException(status_code=400, detail="Invalid or expired code")
    
    if expires and expires < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Code has expired")
    
    # Update password
    now = datetime.now(timezone.utc)
    await db.users.update_one(
        {"email": data.email},
        {"$set": {
            "password_hash": hash_password(data.new_password),
            "reset_code": None,
            "reset_code_expires": None,
            "updated_at": now
        }}
    )
    
    return {"message": "Password reset successfully"}


# Email Verification
class VerifyEmailRequest(BaseModel):
    email: EmailStr
    code: str

@api_router.post("/auth/verify-email")
async def verify_email(data: VerifyEmailRequest):
    """Verify email using the code"""
    user_doc = await db.users.find_one({"email": data.email}, {"_id": 0})
    
    if not user_doc:
        raise HTTPException(status_code=400, detail="Invalid email or code")
    
    # Check code
    stored_code = user_doc.get("verification_code")
    expires = user_doc.get("verification_code_expires")
    
    if not stored_code or stored_code != data.code:
        raise HTTPException(status_code=400, detail="Invalid or expired code")
    
    if expires and expires < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Code has expired")
    
    # Mark as verified
    now = datetime.now(timezone.utc)
    await db.users.update_one(
        {"email": data.email},
        {"$set": {
            "email_verified": True,
            "verification_code": None,
            "verification_code_expires": None,
            "updated_at": now
        }}
    )
    
    return {"message": "Email verified successfully"}


# Resend Verification Email
class ResendVerificationRequest(BaseModel):
    email: EmailStr

@api_router.post("/auth/resend-verification")
async def resend_verification(data: ResendVerificationRequest):
    """Resend verification email"""
    user_doc = await db.users.find_one({"email": data.email}, {"_id": 0})
    
    if not user_doc:
        return {"message": "If this email exists, a verification code will be sent"}
    
    if user_doc.get("email_verified"):
        return {"message": "Email is already verified"}
    
    # Generate new code
    verification_code = generate_verification_code()
    now = datetime.now(timezone.utc)
    
    await db.users.update_one(
        {"email": data.email},
        {"$set": {
            "verification_code": verification_code,
            "verification_code_expires": now + timedelta(hours=24),
            "updated_at": now
        }}
    )
    
    # Send email
    try:
        user_name = user_doc.get("name", "User")
        await send_verification_email(data.email, user_name, verification_code)
        logger.info(f"Verification email resent to {data.email}")
    except Exception as e:
        logger.error(f"Failed to resend verification email to {data.email}: {e}")
    
    return {"message": "If this email exists, a verification code will be sent"}


# Purchase confirmation email endpoint
class PurchaseEmailRequest(BaseModel):
    product_id: str
    product_name: str
    price: str

@api_router.post("/email/purchase-confirmation")
async def send_purchase_email(data: PurchaseEmailRequest, user: User = Depends(get_current_user)):
    """Send purchase confirmation email"""
    try:
        await send_purchase_confirmation_email(
            user.email,
            user.name,
            data.product_name,
            data.price
        )
        
        # Also send premium welcome if it's a subscription
        if "premium" in data.product_id.lower() or "subscription" in data.product_id.lower():
            plan_name = "Premium Monthly" if "monthly" in data.product_id.lower() else "Premium Yearly"
            await send_premium_welcome_email(user.email, user.name, plan_name)
        
        return {"message": "Purchase confirmation email sent"}
    except Exception as e:
        logger.error(f"Failed to send purchase email: {e}")
        raise HTTPException(status_code=500, detail="Failed to send email")

@api_router.post("/auth/google/callback")
async def google_oauth_callback(request: Request, response: Response):
    """Process Google OAuth session_id from Emergent Auth"""
    body = await request.json()
    session_id = body.get("session_id")
    
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id required")
    
    # Exchange session_id for user data
    async with httpx.AsyncClient() as client:
        try:
            resp = await client.get(
                "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
                headers={"X-Session-ID": session_id}
            )
            if resp.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid session")
            
            user_data = resp.json()
        except Exception as e:
            logger.error(f"OAuth error: {e}")
            raise HTTPException(status_code=500, detail="OAuth service error")
    
    # Check if user exists
    existing = await db.users.find_one({"email": user_data["email"]}, {"_id": 0})
    
    if existing:
        user_id = existing["user_id"]
    else:
        # Create new user
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc)
        
        new_user = {
            "user_id": user_id,
            "email": user_data["email"],
            "name": user_data["name"],
            "picture": user_data.get("picture"),
            "subscription_type": "free",
            "subscription_expires_at": None,
            "ocr_usage_today": 0,
            "ocr_usage_date": None,
            "created_at": now,
            "updated_at": now
        }
        await db.users.insert_one(new_user)
    
    # Create session
    session_token = user_data.get("session_token", str(uuid.uuid4()))
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at,
        "created_at": datetime.now(timezone.utc)
    })
    
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=7 * 24 * 60 * 60,
        path="/"
    )
    
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    user = User(**{k: v for k, v in user_doc.items() if k != "password_hash"})
    
    return {"user": user_to_response(user), "token": session_token}

@api_router.post("/auth/google/native")
async def google_native_login(request: Request, response: Response):
    """Process native Google Sign-In from Android/iOS app"""
    body = await request.json()
    
    email = body.get("email")
    name = body.get("name")
    photo = body.get("photo")
    id_token = body.get("id_token")  # For future verification with Google
    
    if not email:
        raise HTTPException(status_code=400, detail="email required")
    
    # Check if user exists
    existing = await db.users.find_one({"email": email}, {"_id": 0})
    
    if existing:
        user_id = existing["user_id"]
        # Update photo if provided
        if photo and photo != existing.get("picture"):
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {"picture": photo, "updated_at": datetime.now(timezone.utc)}}
            )
    else:
        # Create new user
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc)
        
        new_user = {
            "user_id": user_id,
            "email": email,
            "name": name or email.split("@")[0],
            "picture": photo,
            "subscription_type": "free",
            "subscription_expires_at": None,
            "ocr_usage_today": 0,
            "ocr_usage_date": None,
            "created_at": now,
            "updated_at": now
        }
        await db.users.insert_one(new_user)
    
    # Create session
    session_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(days=30)
    
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at,
        "created_at": datetime.now(timezone.utc)
    })
    
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=30 * 24 * 60 * 60,
        path="/"
    )
    
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    user = User(**{k: v for k, v in user_doc.items() if k != "password_hash"})
    
    return {"user": user_to_response(user), "token": session_token}

@api_router.post("/auth/apple/native")
async def apple_native_login(request: Request, response: Response):
    """Process native Apple Sign-In from iOS app"""
    body = await request.json()
    
    identity_token = body.get("identity_token")  # JWT token from Apple
    user_id_apple = body.get("user_id")  # Apple's user ID
    email = body.get("email")  # May be null if user hides email
    full_name = body.get("full_name")  # May be null after first login
    
    if not user_id_apple:
        raise HTTPException(status_code=400, detail="user_id required")
    
    # Use Apple's user_id as a fallback identifier if email is hidden
    lookup_id = email if email else f"apple_{user_id_apple}"
    
    # Check if user exists by Apple ID or email
    existing = await db.users.find_one(
        {"$or": [{"apple_id": user_id_apple}, {"email": email}]} if email else {"apple_id": user_id_apple},
        {"_id": 0}
    )
    
    if existing:
        user_id = existing["user_id"]
        # Update Apple ID if not set
        if not existing.get("apple_id"):
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {"apple_id": user_id_apple, "updated_at": datetime.now(timezone.utc)}}
            )
    else:
        # Create new user
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc)
        
        # Generate placeholder email if hidden
        user_email = email or f"{user_id_apple[:8]}@privaterelay.appleid.com"
        
        new_user = {
            "user_id": user_id,
            "email": user_email,
            "name": full_name or "Apple User",
            "apple_id": user_id_apple,
            "picture": None,
            "subscription_type": "free",
            "subscription_expires_at": None,
            "ocr_usage_today": 0,
            "ocr_usage_date": None,
            "created_at": now,
            "updated_at": now
        }
        await db.users.insert_one(new_user)
    
    # Create session
    session_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(days=30)
    
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at,
        "created_at": datetime.now(timezone.utc)
    })
    
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=30 * 24 * 60 * 60,
        path="/"
    )
    
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    user = User(**{k: v for k, v in user_doc.items() if k not in ["password_hash", "apple_id"]})
    
    return {"user": user_to_response(user), "token": session_token}

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: User = Depends(get_current_user)):
    """Get current user profile"""
    return user_to_response(current_user)

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    """Logout user"""
    token = request.cookies.get("session_token")
    if token:
        await db.user_sessions.delete_one({"session_token": token})
    
    response.delete_cookie(key="session_token", path="/")
    return {"message": "Logged out successfully"}

# ==================== USER ENDPOINTS ====================

@api_router.put("/users/subscription", response_model=UserResponse)
async def update_subscription(
    subscription: SubscriptionUpdate,
    current_user: User = Depends(get_current_user)
):
    """Update user subscription (mock implementation)"""
    expires_at = None
    if subscription.subscription_type == "premium":
        expires_at = datetime.now(timezone.utc) + timedelta(days=subscription.duration_days or 30)
    
    await db.users.update_one(
        {"user_id": current_user.user_id},
        {"$set": {
            "subscription_type": subscription.subscription_type,
            "subscription_expires_at": expires_at,
            "updated_at": datetime.now(timezone.utc)
        }}
    )
    
    user_doc = await db.users.find_one({"user_id": current_user.user_id}, {"_id": 0})
    user = User(**{k: v for k, v in user_doc.items() if k != "password_hash"})
    return user_to_response(user)

@api_router.put("/users/profile", response_model=UserResponse)
async def update_profile(
    name: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Update user profile"""
    update_data = {"updated_at": datetime.now(timezone.utc)}
    if name:
        update_data["name"] = name
    
    await db.users.update_one(
        {"user_id": current_user.user_id},
        {"$set": update_data}
    )
    
    user_doc = await db.users.find_one({"user_id": current_user.user_id}, {"_id": 0})
    user = User(**{k: v for k, v in user_doc.items() if k != "password_hash"})
    return user_to_response(user)

@api_router.post("/users/start-trial", response_model=UserResponse)
async def start_trial(
    current_user: User = Depends(get_current_user)
):
    """Start 7-day free trial"""
    # Check if trial already used
    if current_user.trial_used:
        raise HTTPException(status_code=400, detail="Trial already used. Subscribe to premium for full access.")
    
    # Check if already premium
    if current_user.subscription_type == "premium":
        raise HTTPException(status_code=400, detail="Already a premium subscriber.")
    
    await db.users.update_one(
        {"user_id": current_user.user_id},
        {"$set": {
            "subscription_type": "trial",
            "trial_start_date": datetime.now(timezone.utc),
            "trial_used": True,
            "updated_at": datetime.now(timezone.utc)
        }}
    )
    
    user_doc = await db.users.find_one({"user_id": current_user.user_id}, {"_id": 0})
    user = User(**{k: v for k, v in user_doc.items() if k != "password_hash"})
    return user_to_response(user)

@api_router.get("/users/usage-stats")
async def get_usage_stats(
    current_user: User = Depends(get_current_user)
):
    """Get user's current usage statistics"""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    current_month = datetime.now(timezone.utc).strftime("%Y-%m")
    
    scans_today = current_user.scans_today if current_user.last_scan_date == today else 0
    scans_month = current_user.scans_this_month if current_user.scan_month == current_month else 0
    ocr_today = current_user.ocr_usage_today if current_user.ocr_usage_date == today else 0
    
    is_premium = current_user.subscription_type in ["premium", "trial"]
    
    return {
        "scans_today": scans_today,
        "scans_this_month": scans_month,
        "ocr_today": ocr_today,
        "limits": {
            "scans_per_day": FREE_SCANS_PER_DAY if not is_premium else None,
            "scans_per_month": FREE_SCANS_PER_MONTH if not is_premium else None,
            "ocr_per_day": FREE_OCR_PER_DAY if not is_premium else None,
        },
        "is_premium": is_premium
    }

class UpdatePremiumRequest(BaseModel):
    user_id: str
    is_premium: bool
    has_removed_ads: bool = False
    subscription_type: Optional[str] = None

@api_router.post("/user/update-premium")
async def update_premium_status(
    request: UpdatePremiumRequest,
    current_user: User = Depends(get_current_user)
):
    """Update user's premium status from app purchase"""
    # Ensure user can only update their own status
    if request.user_id != current_user.user_id:
        raise HTTPException(status_code=403, detail="Cannot update other user's status")
    
    update_data = {}
    
    if request.is_premium:
        # Map subscription type
        if request.subscription_type:
            if "yearly" in request.subscription_type:
                update_data["subscription_type"] = "premium"
                update_data["subscription_expires_at"] = datetime.now(timezone.utc) + timedelta(days=365)
            elif "monthly" in request.subscription_type:
                update_data["subscription_type"] = "premium"
                update_data["subscription_expires_at"] = datetime.now(timezone.utc) + timedelta(days=30)
            else:
                update_data["subscription_type"] = "premium"
    
    if request.has_removed_ads:
        # Store has_removed_ads in user document (add field if needed)
        update_data["has_removed_ads"] = True
    
    if update_data:
        result = await db.users.update_one(
            {"user_id": current_user.user_id},
            {"$set": update_data}
        )
        logger.info(f"Updated premium status for user {current_user.user_id}: {update_data}")
    
    return {"success": True, "updated": update_data}


class AvatarUploadRequest(BaseModel):
    avatar_base64: str


@api_router.post("/user/avatar")
async def upload_avatar(
    request: AvatarUploadRequest,
    current_user: User = Depends(get_current_user)
):
    """Upload user avatar image"""
    try:
        avatar_base64 = request.avatar_base64
        
        # Remove data URL prefix if present
        if avatar_base64.startswith('data:'):
            avatar_base64 = avatar_base64.split(',')[1]
        
        avatar_url = None
        
        # Try to upload to S3
        if s3_client:
            try:
                # Generate avatar key
                avatar_key = f"avatars/{current_user.user_id}/avatar.jpg"
                
                # Decode and upload
                image_data = base64.b64decode(avatar_base64)
                
                s3_client.put_object(
                    Bucket=AWS_S3_BUCKET_NAME,
                    Key=avatar_key,
                    Body=image_data,
                    ContentType='image/jpeg',
                )
                
                avatar_url = f"https://{AWS_S3_BUCKET_NAME}.s3.amazonaws.com/{avatar_key}?t={int(datetime.now().timestamp())}"
                logger.info(f"✅ Avatar uploaded to S3 for user {current_user.user_id}")
            except Exception as s3_err:
                logger.warning(f"S3 avatar upload failed: {s3_err}")
        
        # If S3 failed or not available, store as base64 in DB
        if not avatar_url:
            # Store small version in DB (limit to 100KB)
            if len(avatar_base64) > 100000:
                raise HTTPException(status_code=400, detail="Avatar image too large. Please use a smaller image.")
            avatar_url = f"data:image/jpeg;base64,{avatar_base64}"
        
        # Update user document
        await db.users.update_one(
            {"user_id": current_user.user_id},
            {"$set": {
                "avatar_url": avatar_url,
                "photo_url": avatar_url,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        return {"success": True, "avatar_url": avatar_url}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Avatar upload error: {e}")
        raise HTTPException(status_code=500, detail="Failed to upload avatar")


@api_router.post("/documents/remove-watermarks")
async def remove_watermarks_from_documents(
    current_user: User = Depends(get_current_user)
):
    """
    Remove watermarks from all user's documents after becoming premium.
    Replaces watermarked images with original (non-watermarked) versions.
    """
    # Check if user is premium
    is_premium = current_user.subscription_type in ["premium", "trial"]
    if current_user.subscription_type == "trial" and current_user.trial_start_date:
        trial_end = current_user.trial_start_date + timedelta(days=TRIAL_DURATION_DAYS)
        if datetime.now(timezone.utc) >= trial_end.replace(tzinfo=timezone.utc):
            is_premium = False
    
    # Also check has_removed_ads flag (one-time purchase)
    user_data = await db.users.find_one({"user_id": current_user.user_id})
    has_removed_ads = user_data.get("has_removed_ads", False) if user_data else False
    
    if not is_premium and not has_removed_ads:
        raise HTTPException(status_code=403, detail="Only premium users can remove watermarks")
    
    # Get all user's documents
    documents = await db.documents.find({"user_id": current_user.user_id}).to_list(length=1000)
    
    updated_count = 0
    pages_updated = 0
    
    for doc in documents:
        doc_updated = False
        updated_pages = []
        
        for page in doc.get("pages", []):
            page_updated = False
            
            # Check if page has watermark
            if page.get("has_watermark", False):
                # S3 storage: Replace image_url with original_image_url
                if page.get("original_image_url"):
                    page["image_url"] = page["original_image_url"]
                    # Generate new thumbnail from original
                    try:
                        # Download original from S3
                        original_url = page["original_image_url"]
                        if s3_client and original_url:
                            # Extract S3 key from URL
                            s3_key = original_url.split(f"{AWS_S3_BUCKET_NAME}/")[-1] if AWS_S3_BUCKET_NAME in original_url else None
                            if s3_key:
                                response = s3_client.get_object(Bucket=AWS_S3_BUCKET_NAME, Key=s3_key)
                                image_data = response['Body'].read()
                                image_base64 = base64.b64encode(image_data).decode('utf-8')
                                
                                # Create new thumbnail
                                new_thumbnail_base64 = create_thumbnail(image_base64)
                                
                                # Upload new thumbnail
                                new_thumbnail_url = upload_to_s3(
                                    new_thumbnail_base64, 
                                    current_user.user_id, 
                                    doc["document_id"], 
                                    page["page_id"], 
                                    "thumbnail"
                                )
                                if new_thumbnail_url:
                                    page["thumbnail_url"] = new_thumbnail_url
                    except Exception as e:
                        logger.warning(f"Failed to update thumbnail for page {page.get('page_id')}: {e}")
                    
                    # Clean up original_image_url field
                    page.pop("original_image_url", None)
                    page["has_watermark"] = False
                    page_updated = True
                    pages_updated += 1
                
                # Base64 storage: Replace image_base64 with original_image_base64
                elif page.get("original_image_base64"):
                    page["image_base64"] = page["original_image_base64"]
                    # Create new thumbnail
                    page["thumbnail_base64"] = create_thumbnail(page["original_image_base64"])
                    # Clean up
                    page.pop("original_image_base64", None)
                    page["has_watermark"] = False
                    page_updated = True
                    pages_updated += 1
            
            updated_pages.append(page)
            if page_updated:
                doc_updated = True
        
        # Update document in database if any page was updated
        if doc_updated:
            await db.documents.update_one(
                {"document_id": doc["document_id"]},
                {"$set": {"pages": updated_pages}}
            )
            updated_count += 1
    
    logger.info(f"✅ Removed watermarks for user {current_user.user_id}: {updated_count} documents, {pages_updated} pages")
    
    return {
        "success": True,
        "documents_updated": updated_count,
        "pages_updated": pages_updated,
        "message": f"Watermarks removed from {pages_updated} pages across {updated_count} documents"
    }


# ==================== DOCUMENT ENDPOINTS ====================

@api_router.post("/documents", response_model=Document)
async def create_document(
    doc_data: DocumentCreate,
    current_user: User = Depends(get_current_user)
):
    """Create a new document"""
    # Check scan limits for free users
    can_scan, limit_message = await check_scan_limits(current_user)
    if not can_scan:
        raise HTTPException(status_code=403, detail=limit_message)
    
    document_id = f"doc_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    
    # Check if user needs watermark (free users without premium or remove_ads)
    is_premium = current_user.subscription_type in ["premium", "trial"]
    if current_user.subscription_type == "trial" and current_user.trial_start_date:
        trial_end = current_user.trial_start_date + timedelta(days=TRIAL_DURATION_DAYS)
        if datetime.now(timezone.utc) >= trial_end.replace(tzinfo=timezone.utc):
            is_premium = False
    
    # Also check if user has removed ads (they shouldn't get watermarks either)
    user_data = await db.users.find_one({"user_id": current_user.user_id})
    has_removed_ads = user_data.get("has_removed_ads", False) if user_data else False
    
    # No watermark if premium OR has_removed_ads
    skip_watermark = is_premium or has_removed_ads
    
    # Process pages - create thumbnails, add watermark for free users, upload to S3
    processed_pages = []
    for i, page in enumerate(doc_data.pages):
        page_dict = page.dict()
        page_id = f"page_{uuid.uuid4().hex[:8]}"
        page_dict["page_id"] = page_id
        page_dict["order"] = i
        
        # Store original image and add watermark for free users (unless they removed ads)
        original_image_base64 = page.image_base64
        image_base64 = page.image_base64
        has_watermark = False
        
        if not skip_watermark:
            image_base64 = add_watermark(image_base64, "ScanUp")
            has_watermark = True
        
        # Create thumbnail (from watermarked image if applicable)
        thumbnail_base64 = create_thumbnail(image_base64)
        
        # Upload to S3 if configured
        if s3_client:
            # Upload main image (watermarked for free users)
            image_url = upload_to_s3(image_base64, current_user.user_id, document_id, page_id, "page")
            thumbnail_url = upload_to_s3(thumbnail_base64, current_user.user_id, document_id, page_id, "thumbnail")
            
            # Also upload original (non-watermarked) for free users
            original_image_url = None
            if has_watermark:
                original_image_url = upload_to_s3(original_image_base64, current_user.user_id, document_id, page_id, "original")
            
            if image_url and thumbnail_url:
                # Store URLs instead of base64
                page_dict["image_url"] = image_url
                page_dict["thumbnail_url"] = thumbnail_url
                page_dict["has_watermark"] = has_watermark
                if original_image_url:
                    page_dict["original_image_url"] = original_image_url
                page_dict.pop("image_base64", None)
                page_dict.pop("thumbnail_base64", None)
                logger.info(f"✅ Page {page_id} uploaded to S3 (watermark: {has_watermark})")
            else:
                # Fallback to base64 if S3 upload fails
                page_dict["image_base64"] = image_base64
                page_dict["thumbnail_base64"] = thumbnail_base64
                page_dict["has_watermark"] = has_watermark
                if has_watermark:
                    page_dict["original_image_base64"] = original_image_base64
                logger.warning(f"⚠️ S3 upload failed, storing base64 for page {page_id}")
        else:
            # No S3, store base64 in MongoDB
            page_dict["image_base64"] = image_base64
            page_dict["thumbnail_base64"] = thumbnail_base64
            page_dict["has_watermark"] = has_watermark
            if has_watermark:
                page_dict["original_image_base64"] = original_image_base64
        
        processed_pages.append(page_dict)
    
    document = {
        "document_id": document_id,
        "user_id": current_user.user_id,
        "name": doc_data.name,
        "folder_id": doc_data.folder_id,
        "tags": doc_data.tags,
        "pages": processed_pages,
        "ocr_full_text": None,
        "is_password_protected": False,
        "has_watermark": not skip_watermark,
        "storage_type": "s3" if s3_client else "mongodb",  # Track storage type
        "created_at": now,
        "updated_at": now
    }
    
    await db.documents.insert_one(document)
    
    # Increment scan count for this user
    await increment_scan_count(current_user.user_id)
    
    # Remove _id before returning
    document.pop("_id", None)
    return Document(**document)

@api_router.get("/documents", response_model=List[Document])
async def get_documents(
    folder_id: Optional[str] = None,
    search: Optional[str] = None,
    tag: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Get all documents for current user"""
    query = {"user_id": current_user.user_id}
    
    if folder_id:
        query["folder_id"] = folder_id
    
    if tag:
        query["tags"] = tag
    
    if search:
        query["$or"] = [
            {"name": {"$regex": search, "$options": "i"}},
            {"ocr_full_text": {"$regex": search, "$options": "i"}}
        ]
    
    documents = await db.documents.find(query, {"_id": 0}).sort("updated_at", -1).to_list(1000)
    return [Document(**doc) for doc in documents]

@api_router.get("/documents/{document_id}", response_model=Document)
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    """Get a specific document"""
    document = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return Document(**document)

@api_router.put("/documents/{document_id}", response_model=Document)
async def update_document(
    document_id: str,
    doc_update: DocumentUpdate,
    current_user: User = Depends(get_current_user)
):
    """Update a document"""
    document = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    update_data = {"updated_at": datetime.now(timezone.utc)}
    
    if doc_update.name is not None:
        update_data["name"] = doc_update.name
    if doc_update.folder_id is not None:
        update_data["folder_id"] = doc_update.folder_id
    if doc_update.tags is not None:
        update_data["tags"] = doc_update.tags
    if doc_update.pages is not None:
        processed_pages = []
        for i, page in enumerate(doc_update.pages):
            page_dict = page.dict()
            page_dict["order"] = i
            # Always regenerate thumbnail to ensure it matches the current image
            page_dict["thumbnail_base64"] = create_thumbnail(page.image_base64)
            processed_pages.append(page_dict)
        update_data["pages"] = processed_pages
        
        # Update document thumbnail to first page's thumbnail
        if processed_pages:
            update_data["thumbnail_base64"] = processed_pages[0].get("thumbnail_base64")
    if doc_update.is_password_protected is not None:
        update_data["is_password_protected"] = doc_update.is_password_protected
    if doc_update.password_hash is not None:
        update_data["password_hash"] = hash_password(doc_update.password_hash)
    
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": update_data}
    )
    
    updated_doc = await db.documents.find_one({"document_id": document_id}, {"_id": 0})
    return Document(**updated_doc)

@api_router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    """Delete a document and its S3 files"""
    # First get the document to check ownership
    document = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete files from S3 if S3 is configured
    if s3_client and AWS_S3_BUCKET_NAME:
        try:
            delete_from_s3(current_user.user_id, document_id)
            logger.info(f"Deleted S3 files for document {document_id}")
        except Exception as e:
            logger.error(f"Failed to delete S3 files for document {document_id}: {e}")
            # Continue with database deletion even if S3 fails
    
    # Delete from database
    result = await db.documents.delete_one(
        {"document_id": document_id, "user_id": current_user.user_id}
    )
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {"message": "Document deleted successfully"}

@api_router.post("/documents/{document_id}/pages", response_model=Document)
async def add_page_to_document(
    document_id: str,
    page: PageData,
    current_user: User = Depends(get_current_user)
):
    """Add a page to an existing document"""
    document = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    page_dict = page.dict()
    page_dict["order"] = len(document.get("pages", []))
    page_dict["thumbnail_base64"] = create_thumbnail(page.image_base64)
    
    await db.documents.update_one(
        {"document_id": document_id},
        {
            "$push": {"pages": page_dict},
            "$set": {"updated_at": datetime.now(timezone.utc)}
        }
    )
    
    updated_doc = await db.documents.find_one({"document_id": document_id}, {"_id": 0})
    return Document(**updated_doc)

# ==================== FOLDER ENDPOINTS ====================

@api_router.post("/folders", response_model=Folder)
async def create_folder(
    folder_data: FolderCreate,
    current_user: User = Depends(get_current_user)
):
    """Create a new folder"""
    folder_id = f"folder_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    
    folder = {
        "folder_id": folder_id,
        "user_id": current_user.user_id,
        "name": folder_data.name,
        "color": folder_data.color,
        "parent_id": folder_data.parent_id,
        "is_protected": False,
        "password_hash": None,
        "created_at": now
    }
    
    await db.folders.insert_one(folder)
    folder.pop("_id", None)
    return Folder(**folder)

@api_router.get("/folders")
async def get_folders(current_user: User = Depends(get_current_user)):
    """Get all folders for current user with document counts"""
    folders = await db.folders.find(
        {"user_id": current_user.user_id},
        {"_id": 0}
    ).to_list(1000)
    
    # Add document count to each folder
    result_folders = []
    for f in folders:
        folder_obj = Folder(**f)
        # Count documents in this folder
        doc_count = await db.documents.count_documents({
            "user_id": current_user.user_id,
            "folder_id": folder_obj.folder_id
        })
        # Convert to dict and add document_count
        folder_dict = folder_obj.model_dump()
        folder_dict["document_count"] = doc_count
        result_folders.append(folder_dict)
    
    return result_folders

class FolderUpdate(BaseModel):
    name: Optional[str] = None
    color: Optional[str] = None
    is_protected: Optional[bool] = None
    password_hash: Optional[str] = None

@api_router.put("/folders/{folder_id}", response_model=Folder)
async def update_folder(
    folder_id: str,
    folder_data: FolderUpdate,
    current_user: User = Depends(get_current_user)
):
    """Update a folder"""
    update_data = {}
    
    if folder_data.name is not None:
        update_data["name"] = folder_data.name
    if folder_data.color is not None:
        update_data["color"] = folder_data.color
    if folder_data.is_protected is not None:
        update_data["is_protected"] = folder_data.is_protected
    if folder_data.password_hash is not None:
        # Hash the password before storing
        update_data["password_hash"] = hash_password(folder_data.password_hash)
    elif not folder_data.is_protected:
        # If removing protection, also clear the password
        update_data["password_hash"] = None
    
    if not update_data:
        raise HTTPException(status_code=400, detail="No data to update")
    
    result = await db.folders.update_one(
        {"folder_id": folder_id, "user_id": current_user.user_id},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Folder not found")
    
    folder = await db.folders.find_one(
        {"folder_id": folder_id},
        {"_id": 0}
    )
    return Folder(**folder)

class FolderPasswordVerify(BaseModel):
    password: str

@api_router.post("/folders/{folder_id}/verify-password")
async def verify_folder_password(
    folder_id: str,
    password_data: FolderPasswordVerify,
    current_user: User = Depends(get_current_user)
):
    """Verify folder password"""
    folder = await db.folders.find_one(
        {"folder_id": folder_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")
    
    if not folder.get("is_protected"):
        return {"success": True, "message": "Folder is not protected"}
    
    stored_hash = folder.get("password_hash")
    if not stored_hash:
        return {"success": True, "message": "No password set"}
    
    # Verify password against stored hash
    if verify_password(password_data.password, stored_hash):
        return {"success": True, "message": "Password verified"}
    else:
        raise HTTPException(status_code=401, detail="Incorrect password")

@api_router.delete("/folders/{folder_id}")
async def delete_folder(
    folder_id: str,
    current_user: User = Depends(get_current_user)
):
    """Delete a folder"""
    result = await db.folders.delete_one(
        {"folder_id": folder_id, "user_id": current_user.user_id}
    )
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Folder not found")
    
    # Move documents to root
    await db.documents.update_many(
        {"folder_id": folder_id, "user_id": current_user.user_id},
        {"$set": {"folder_id": None}}
    )
    
    return {"message": "Folder deleted successfully"}

# ==================== IMAGE PROCESSING ENDPOINTS ====================

@api_router.post("/images/process", response_model=ImageProcessResponse)
async def process_image(
    request: ImageProcessRequest,
    current_user: User = Depends(get_current_user)
):
    """Process an image (crop, rotate, filter, perspective)"""
    result = request.image_base64
    
    if request.operation == "filter":
        filter_type = request.params.get("type", "original")
        brightness = request.params.get("brightness", 0)
        contrast = request.params.get("contrast", 0)
        saturation = request.params.get("saturation", 0)
        result = apply_image_filter(result, filter_type, brightness, contrast, saturation)
    elif request.operation == "rotate":
        degrees = request.params.get("degrees", 90)
        result = rotate_image(result, degrees)
    elif request.operation == "crop":
        x = request.params.get("x", 0)
        y = request.params.get("y", 0)
        width = request.params.get("width", 100)
        height = request.params.get("height", 100)
        result = crop_image(result, x, y, width, height)
    elif request.operation == "perspective_crop":
        corners = request.params.get("corners", [])
        if corners:
            result = perspective_crop(result, corners)
    
    return ImageProcessResponse(processed_image_base64=result)


@api_router.post("/images/process-public", response_model=ImageProcessResponse)
async def process_image_public(request: ImageProcessRequest):
    """Public endpoint to process an image (no auth required) - for guest users"""
    # DEBUG: Log incoming request details
    img_preview = (request.image_base64[:50] + "...") if request.image_base64 and len(request.image_base64) > 50 else request.image_base64
    logger.info(f"[process-public] operation={request.operation}, image_len={len(request.image_base64) if request.image_base64 else 0}, preview={img_preview}")
    result = request.image_base64
    
    if request.operation == "filter":
        filter_type = request.params.get("type", "original")
        brightness = request.params.get("brightness", 0)
        contrast = request.params.get("contrast", 0)
        saturation = request.params.get("saturation", 0)
        result = apply_image_filter(result, filter_type, brightness, contrast, saturation)
    elif request.operation == "rotate":
        degrees = request.params.get("degrees", 90)
        result = rotate_image(result, degrees)
    elif request.operation == "crop":
        x = request.params.get("x", 0)
        y = request.params.get("y", 0)
        width = request.params.get("width", 100)
        height = request.params.get("height", 100)
        result = crop_image(result, x, y, width, height)
    elif request.operation == "perspective_crop":
        corners = request.params.get("corners", [])
        if corners:
            result = perspective_crop(result, corners)
    
    return ImageProcessResponse(processed_image_base64=result)

@api_router.post("/images/auto-crop")
async def auto_crop_image(
    request: ImageProcessRequest,
    current_user: User = Depends(get_current_user)
):
    """Automatically detect and crop document from image"""
    document_type = request.params.get("document_type", "document")
    
    # First detect edges with improved algorithm
    edge_result = detect_document_edges(request.image_base64, document_type)
    
    if edge_result.get("detected") and edge_result.get("corners"):
        # Apply perspective transform
        cropped = perspective_crop(request.image_base64, edge_result["corners"])
        return {
            "success": True,
            "cropped_image_base64": cropped,
            "corners": edge_result["corners"],
            "confidence": edge_result.get("confidence", 0)
        }
    
    # Even if not detected, return default corners for manual adjustment
    return {
        "success": False,
        "cropped_image_base64": request.image_base64,
        "corners": edge_result.get("corners"),  # Default corners provided
        "message": edge_result.get("message", "Could not detect document edges automatically")
    }


@api_router.post("/images/auto-crop-public")
async def auto_crop_image_public(request: ImageProcessRequest):
    """Public endpoint to auto-crop image (no auth required) - for guest users"""
    document_type = request.params.get("document_type", "document")
    
    edge_result = detect_document_edges(request.image_base64, document_type)
    
    if edge_result.get("detected") and edge_result.get("corners"):
        cropped = perspective_crop(request.image_base64, edge_result["corners"])
        return {
            "success": True,
            "cropped_image_base64": cropped,
            "corners": edge_result["corners"],
            "confidence": edge_result.get("confidence", 0)
        }
    
    return {
        "success": False,
        "cropped_image_base64": request.image_base64,
        "corners": edge_result.get("corners"),
        "message": edge_result.get("message", "Could not detect document edges automatically")
    }

class ManualCropRequest(BaseModel):
    image_base64: str
    corners: List[Dict[str, float]]  # [{x: 0-1, y: 0-1}, ...]
    force_portrait: Optional[bool] = False  # Force rotation to portrait if image is landscape

@api_router.post("/images/perspective-crop")
async def manual_perspective_crop(
    request: ManualCropRequest,
    current_user: User = Depends(get_current_user)
):
    """
    Apply manual perspective crop with user-defined corners.
    
    The corners should be normalized (0-1 range) and will be converted to pixel coordinates.
    The function handles EXIF orientation and ensures corners are in correct order.
    """
    try:
        # Decode image to get dimensions
        if "," in request.image_base64:
            img_data = request.image_base64.split(",")[1]
        else:
            img_data = request.image_base64
            
        image_bytes = base64.b64decode(img_data)
        
        # Fix EXIF orientation FIRST using PIL (more reliable)
        corrected_bytes, was_rotated, orientation = fix_image_orientation_pil(image_bytes)
        logger.info(f"EXIF orientation: {orientation}, was_rotated: {was_rotated}")
        
        # Now decode the corrected image with OpenCV
        nparr = np.frombuffer(corrected_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return {"success": False, "cropped_image_base64": request.image_base64, "message": "Could not decode image"}
        
        height, width = img.shape[:2]
        
        logger.info(f"Image dimensions after EXIF fix: {width}x{height}")
        logger.info(f"Input corners (normalized): {request.corners}")
        
        # Convert normalized corners to pixel coordinates using FLOAT
        # Do NOT convert to int yet - let perspective_crop handle precision
        pixel_corners = []
        for corner in request.corners:
            # Use float multiplication for precision
            px = float(corner.get('x', 0)) * width
            py = float(corner.get('y', 0)) * height
            pixel_corners.append({'x': px, 'y': py})
        
        logger.info(f"Pixel corners: {pixel_corners}")
        
        # Re-encode the EXIF-corrected image for perspective_crop
        corrected_base64 = base64.b64encode(corrected_bytes).decode()
        
        # Apply perspective transform on the EXIF-corrected image
        cropped = perspective_crop(corrected_base64, pixel_corners)
        
        return {
            "success": True,
            "cropped_image_base64": cropped
        }
    except Exception as e:
        logger.error(f"Manual crop error: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {
            "success": False,
            "cropped_image_base64": request.image_base64,
            "message": str(e)
        }


# ==================== BOOK SCAN PAGE SPLITTING ====================

class BookSplitRequest(BaseModel):
    """Request model for book page splitting"""
    image_base64: str
    corners: Optional[List[Dict]] = None  # Optional outer corners for the entire book
    gutter_position: Optional[float] = None  # Optional manual gutter position (0-1), default is 0.5


def detect_book_gutter(img: np.ndarray) -> float:
    """
    Automatically detect the book gutter (center fold) position.
    Uses vertical line detection and intensity analysis.
    Returns normalized position (0-1) from left edge.
    """
    try:
        height, width = img.shape[:2]
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Focus on the middle portion where the gutter is most likely
        center_region = gray[:, int(width * 0.3):int(width * 0.7)]
        
        # Apply edge detection
        edges = cv2.Canny(center_region, 50, 150)
        
        # Look for vertical lines using Hough transform
        lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=50, minLineLength=height * 0.3, maxLineGap=20)
        
        vertical_lines = []
        if lines is not None:
            for line in lines:
                x1, y1, x2, y2 = line[0]
                # Check if line is mostly vertical (angle < 10 degrees)
                angle = abs(np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi)
                if angle > 80 or angle < 10:
                    avg_x = (x1 + x2) / 2 + int(width * 0.3)  # Adjust for region offset
                    vertical_lines.append(avg_x)
        
        if vertical_lines:
            # Use the median of detected vertical lines near center
            center_lines = [x for x in vertical_lines if abs(x - width/2) < width * 0.15]
            if center_lines:
                gutter_x = np.median(center_lines)
                return gutter_x / width
        
        # Fallback: analyze brightness pattern (gutter often has shadow)
        # Calculate column-wise mean intensity in center region
        col_intensity = np.mean(gray[:, int(width*0.35):int(width*0.65)], axis=0)
        
        # Find minimum intensity (shadow in gutter)
        min_idx = np.argmin(col_intensity)
        gutter_x = min_idx + int(width * 0.35)
        
        return gutter_x / width
        
    except Exception as e:
        logger.warning(f"Gutter detection failed: {e}, using center")
        return 0.5


def split_book_pages(img: np.ndarray, gutter_pos: float = 0.5) -> Tuple[np.ndarray, np.ndarray]:
    """
    Split a book image into left and right pages at the gutter position.
    
    Args:
        img: OpenCV image (BGR)
        gutter_pos: Normalized position (0-1) of the gutter from left edge
        
    Returns:
        Tuple of (left_page, right_page) as numpy arrays
    """
    height, width = img.shape[:2]
    gutter_x = int(width * gutter_pos)
    
    # Add small overlap at gutter to ensure no content is lost
    overlap = int(width * 0.01)  # 1% overlap
    
    left_page = img[:, :min(gutter_x + overlap, width)]
    right_page = img[:, max(gutter_x - overlap, 0):]
    
    return left_page, right_page


def perspective_correct_page(img: np.ndarray) -> np.ndarray:
    """
    Apply automatic perspective correction to a single page.
    Detects document edges and corrects keystoning.
    """
    try:
        height, width = img.shape[:2]
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        
        # Edge detection
        edges = cv2.Canny(blurred, 50, 150)
        
        # Find contours
        contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        if not contours:
            return img
        
        # Find the largest contour
        largest_contour = max(contours, key=cv2.contourArea)
        
        # Get the minimum area rectangle
        rect = cv2.minAreaRect(largest_contour)
        box = cv2.boxPoints(rect)
        box = np.int32(box)
        
        # Order points: top-left, top-right, bottom-right, bottom-left
        def order_points(pts):
            rect = np.zeros((4, 2), dtype="float32")
            s = pts.sum(axis=1)
            rect[0] = pts[np.argmin(s)]  # top-left
            rect[2] = pts[np.argmax(s)]  # bottom-right
            diff = np.diff(pts, axis=1)
            rect[1] = pts[np.argmin(diff)]  # top-right
            rect[3] = pts[np.argmax(diff)]  # bottom-left
            return rect
        
        ordered = order_points(box.astype("float32"))
        
        # Check if the detected area is reasonable (at least 30% of image)
        contour_area = cv2.contourArea(largest_contour)
        if contour_area < (height * width * 0.3):
            # Not enough content detected, return original
            return img
        
        # Calculate output dimensions
        width_top = np.linalg.norm(ordered[0] - ordered[1])
        width_bottom = np.linalg.norm(ordered[2] - ordered[3])
        output_width = max(int(width_top), int(width_bottom))
        
        height_left = np.linalg.norm(ordered[0] - ordered[3])
        height_right = np.linalg.norm(ordered[1] - ordered[2])
        output_height = max(int(height_left), int(height_right))
        
        # Ensure minimum dimensions
        output_width = max(output_width, 100)
        output_height = max(output_height, 100)
        
        # Destination points
        dst = np.array([
            [0, 0],
            [output_width - 1, 0],
            [output_width - 1, output_height - 1],
            [0, output_height - 1]
        ], dtype="float32")
        
        # Apply perspective transform
        matrix = cv2.getPerspectiveTransform(ordered, dst)
        corrected = cv2.warpPerspective(img, matrix, (output_width, output_height), 
                                         flags=cv2.INTER_LINEAR, 
                                         borderMode=cv2.BORDER_REPLICATE)
        
        return corrected
        
    except Exception as e:
        logger.warning(f"Perspective correction failed: {e}")
        return img


@api_router.post("/images/split-book-pages")
async def split_book_pages_endpoint(request: BookSplitRequest):
    """
    Split a book scan into two separate pages with automatic perspective correction.
    
    This endpoint:
    1. Optionally applies perspective correction to the whole book image first
    2. Detects or uses the provided gutter position
    3. Splits the image into left and right pages
    4. Applies independent perspective correction to each page
    5. Returns both pages as separate base64 images
    
    Returns:
        {
            "success": bool,
            "left_page_base64": str,  # Page 1
            "right_page_base64": str,  # Page 2
            "gutter_position": float,  # Detected or provided gutter position
            "message": str
        }
    """
    try:
        # Decode image
        img_data = request.image_base64
        if "," in img_data:
            img_data = img_data.split(",")[1]
        
        image_bytes = base64.b64decode(img_data)
        
        # Fix EXIF orientation
        corrected_bytes, was_rotated, orientation = fix_image_orientation_pil(image_bytes)
        logger.info(f"[BookSplit] EXIF orientation: {orientation}, was_rotated: {was_rotated}")
        
        # Decode with OpenCV
        nparr = np.frombuffer(corrected_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return {
                "success": False,
                "message": "Could not decode image"
            }
        
        height, width = img.shape[:2]
        logger.info(f"[BookSplit] Image dimensions: {width}x{height}")
        
        # Apply perspective correction to whole book if corners provided
        if request.corners and len(request.corners) == 4:
            logger.info("[BookSplit] Applying perspective correction with provided corners")
            # Convert normalized corners to pixels
            pixel_corners = []
            for corner in request.corners:
                px = float(corner.get('x', 0)) * width
                py = float(corner.get('y', 0)) * height
                pixel_corners.append({'x': px, 'y': py})
            
            # Apply perspective transform
            corrected_base64 = base64.b64encode(corrected_bytes).decode()
            cropped_base64 = perspective_crop(corrected_base64, pixel_corners)
            
            # Decode the cropped result
            cropped_bytes = base64.b64decode(cropped_base64)
            nparr = np.frombuffer(cropped_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            height, width = img.shape[:2]
        
        # Detect or use provided gutter position
        if request.gutter_position is not None:
            gutter_pos = request.gutter_position
            logger.info(f"[BookSplit] Using provided gutter position: {gutter_pos}")
        else:
            gutter_pos = detect_book_gutter(img)
            logger.info(f"[BookSplit] Detected gutter position: {gutter_pos}")
        
        # Split into two pages
        left_page, right_page = split_book_pages(img, gutter_pos)
        logger.info(f"[BookSplit] Left page: {left_page.shape}, Right page: {right_page.shape}")
        
        # Apply perspective correction to each page independently
        left_corrected = perspective_correct_page(left_page)
        right_corrected = perspective_correct_page(right_page)
        logger.info(f"[BookSplit] After correction - Left: {left_corrected.shape}, Right: {right_corrected.shape}")
        
        # Encode both pages to base64
        _, left_buffer = cv2.imencode('.jpg', left_corrected, [cv2.IMWRITE_JPEG_QUALITY, 95])
        _, right_buffer = cv2.imencode('.jpg', right_corrected, [cv2.IMWRITE_JPEG_QUALITY, 95])
        
        left_base64 = base64.b64encode(left_buffer.tobytes()).decode()
        right_base64 = base64.b64encode(right_buffer.tobytes()).decode()
        
        return {
            "success": True,
            "left_page_base64": left_base64,
            "right_page_base64": right_base64,
            "gutter_position": gutter_pos,
            "message": "Book pages split successfully"
        }
        
    except Exception as e:
        logger.error(f"[BookSplit] Error: {e}")
        return {
            "success": False,
            "message": str(e)
        }


# ==================== 6-POINT BOOK PERSPECTIVE CORRECTION ====================

class BookSixPointRequest(BaseModel):
    """Request model for 6-point book perspective correction"""
    image_base64: str
    points: List[Dict]  # 6 points: [TL, GT, TR, BR, GB, BL] - all normalized 0-1
    # TL = Top Left, GT = Gutter Top, TR = Top Right
    # BR = Bottom Right, GB = Gutter Bottom, BL = Bottom Left


def perspective_transform_page(img: np.ndarray, src_points: List[List[float]], is_portrait: bool = True) -> np.ndarray:
    """
    Apply perspective transform to a single page using 4 source points.
    
    Args:
        img: OpenCV image (BGR)
        src_points: List of 4 [x, y] points in order: TL, TR, BR, BL (pixel coordinates)
        is_portrait: Whether the output should be portrait orientation
        
    Returns:
        Perspective-corrected image
    """
    src_pts = np.float32(src_points)
    
    # Calculate output dimensions from source points
    width_top = np.linalg.norm(src_pts[0] - src_pts[1])
    width_bottom = np.linalg.norm(src_pts[3] - src_pts[2])
    height_left = np.linalg.norm(src_pts[0] - src_pts[3])
    height_right = np.linalg.norm(src_pts[1] - src_pts[2])
    
    output_width = int(max(width_top, width_bottom))
    output_height = int(max(height_left, height_right))
    
    # Ensure minimum dimensions
    output_width = max(output_width, 100)
    output_height = max(output_height, 100)
    
    # For book pages, we want portrait output (height > width typically)
    # If needed, we can adjust aspect ratio
    
    # Destination points - perfect rectangle
    dst_pts = np.float32([
        [0, 0],                                # TL
        [output_width - 1, 0],                 # TR
        [output_width - 1, output_height - 1], # BR
        [0, output_height - 1]                 # BL
    ])
    
    # Get perspective transform matrix
    matrix = cv2.getPerspectiveTransform(src_pts, dst_pts)
    
    # Apply transform with high-quality interpolation
    warped = cv2.warpPerspective(
        img, 
        matrix, 
        (output_width, output_height),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_CONSTANT,
        borderValue=(255, 255, 255)
    )
    
    return warped


@api_router.post("/images/book-6point-crop")
async def book_six_point_crop(request: BookSixPointRequest):
    """
    Apply 6-point perspective correction for book scanning.
    
    This provides precise control over each page's perspective by using:
    - 4 outer corners (TL, TR, BR, BL)
    - 2 gutter points (GT, GB) that define the book spine
    
    Each page gets its own 4-point perspective transform:
    - Left page: TL, GT, GB, BL
    - Right page: GT, TR, BR, GB
    
    Point order expected: [TL, GT, TR, BR, GB, BL]
    All coordinates should be normalized (0-1).
    
    Returns:
        {
            "success": bool,
            "left_page_base64": str,   # Page 1 (left side of book)
            "right_page_base64": str,  # Page 2 (right side of book)
            "message": str
        }
    """
    try:
        # Validate we have exactly 6 points
        if not request.points or len(request.points) != 6:
            return {
                "success": False,
                "message": f"Expected 6 points, got {len(request.points) if request.points else 0}"
            }
        
        # Decode image
        img_data = request.image_base64
        if "," in img_data:
            img_data = img_data.split(",")[1]
        
        image_bytes = base64.b64decode(img_data)
        
        # Fix EXIF orientation
        corrected_bytes, was_rotated, orientation = fix_image_orientation_pil(image_bytes)
        logger.info(f"[Book6Point] EXIF orientation: {orientation}, was_rotated: {was_rotated}")
        
        # Decode with OpenCV
        nparr = np.frombuffer(corrected_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return {
                "success": False,
                "message": "Could not decode image"
            }
        
        height, width = img.shape[:2]
        logger.info(f"[Book6Point] Image dimensions: {width}x{height}")
        
        # Extract and convert normalized points to pixel coordinates
        # Expected order: [TL, GT, TR, BR, GB, BL]
        points = request.points
        
        TL = [float(points[0].get('x', 0)) * width, float(points[0].get('y', 0)) * height]
        GT = [float(points[1].get('x', 0.5)) * width, float(points[1].get('y', 0)) * height]  # Gutter Top
        TR = [float(points[2].get('x', 1)) * width, float(points[2].get('y', 0)) * height]
        BR = [float(points[3].get('x', 1)) * width, float(points[3].get('y', 1)) * height]
        GB = [float(points[4].get('x', 0.5)) * width, float(points[4].get('y', 1)) * height]  # Gutter Bottom
        BL = [float(points[5].get('x', 0)) * width, float(points[5].get('y', 1)) * height]
        
        logger.info(f"[Book6Point] Points - TL:{TL}, GT:{GT}, TR:{TR}, BR:{BR}, GB:{GB}, BL:{BL}")
        
        # Left page: TL -> GT -> GB -> BL (clockwise from top-left)
        left_src = [TL, GT, GB, BL]
        left_page = perspective_transform_page(img, left_src)
        logger.info(f"[Book6Point] Left page shape: {left_page.shape}")
        
        # Right page: GT -> TR -> BR -> GB (clockwise from top-left of right page)
        right_src = [GT, TR, BR, GB]
        right_page = perspective_transform_page(img, right_src)
        logger.info(f"[Book6Point] Right page shape: {right_page.shape}")
        
        # Encode both pages to base64
        _, left_buffer = cv2.imencode('.jpg', left_page, [cv2.IMWRITE_JPEG_QUALITY, 95])
        _, right_buffer = cv2.imencode('.jpg', right_page, [cv2.IMWRITE_JPEG_QUALITY, 95])
        
        left_base64 = base64.b64encode(left_buffer.tobytes()).decode()
        right_base64 = base64.b64encode(right_buffer.tobytes()).decode()
        
        return {
            "success": True,
            "left_page_base64": left_base64,
            "right_page_base64": right_base64,
            "message": "Book pages perspective-corrected successfully"
        }
        
    except Exception as e:
        logger.error(f"[Book6Point] Error: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {
            "success": False,
            "message": str(e)
        }


# ==================== REAL-TIME EDGE DETECTION ====================

class EdgeDetectionRequest(BaseModel):
    """Request for real-time edge detection"""
    image_base64: str
    mode: str = "document"  # document, book, id_card


def detect_document_edges_cv(img: np.ndarray) -> Optional[List[Dict]]:
    """
    Detect document edges using contour detection (OpenCV based).
    Returns 4 corner points normalized to 0-1 range, or None if not found.
    Uses multiple detection passes with different parameters for better detection.
    """
    height, width = img.shape[:2]
    
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Try multiple detection strategies
    detection_params = [
        # Strategy 1: Standard Canny
        {'blur': (5, 5), 'canny_low': 50, 'canny_high': 150, 'dilate_iter': 1},
        # Strategy 2: Lower thresholds for faded documents
        {'blur': (5, 5), 'canny_low': 30, 'canny_high': 100, 'dilate_iter': 2},
        # Strategy 3: Higher thresholds for noisy backgrounds
        {'blur': (7, 7), 'canny_low': 75, 'canny_high': 200, 'dilate_iter': 1},
        # Strategy 4: Adaptive threshold approach
        {'blur': (5, 5), 'canny_low': 20, 'canny_high': 80, 'dilate_iter': 3},
    ]
    
    for params in detection_params:
        # Apply Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, params['blur'], 0)
        
        # Edge detection using Canny
        edges = cv2.Canny(blurred, params['canny_low'], params['canny_high'])
        
        # Dilate edges to close gaps
        kernel = np.ones((3, 3), np.uint8)
        edges = cv2.dilate(edges, kernel, iterations=params['dilate_iter'])
        
        # Find contours
        contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        if not contours:
            continue
        
        # Sort by area and get largest contours
        contours = sorted(contours, key=cv2.contourArea, reverse=True)[:10]
        
        for contour in contours:
            # Approximate contour with varying precision
            peri = cv2.arcLength(contour, True)
            
            for epsilon_factor in [0.02, 0.03, 0.04, 0.05]:
                approx = cv2.approxPolyDP(contour, epsilon_factor * peri, True)
                
                # If we found a quadrilateral
                if len(approx) == 4:
                    # Check if area is significant (at least 5% of image)
                    area = cv2.contourArea(approx)
                    if area < (height * width * 0.05):
                        continue
                    
                    # Extract points and order them
                    points = approx.reshape(4, 2)
                    
                    # Order points: TL, TR, BR, BL
                    # Sort by y first to get top and bottom pairs
                    sorted_by_y = points[np.argsort(points[:, 1])]
                    top_points = sorted_by_y[:2]
                    bottom_points = sorted_by_y[2:]
                    
                    # Sort top points by x (left to right)
                    top_points = top_points[np.argsort(top_points[:, 0])]
                    # Sort bottom points by x (left to right)  
                    bottom_points = bottom_points[np.argsort(bottom_points[:, 0])]
                    
                    # TL, TR, BR, BL
                    ordered = [
                        {'x': float(top_points[0][0]) / width, 'y': float(top_points[0][1]) / height},
                        {'x': float(top_points[1][0]) / width, 'y': float(top_points[1][1]) / height},
                        {'x': float(bottom_points[1][0]) / width, 'y': float(bottom_points[1][1]) / height},
                        {'x': float(bottom_points[0][0]) / width, 'y': float(bottom_points[0][1]) / height},
                    ]
                    
                    logger.info(f"[EdgeDetect] Found document edges with params: {params}")
                    return ordered
    
    logger.info("[EdgeDetect] No quadrilateral found, using default margins")
    return None


def detect_book_edges_cv(img: np.ndarray) -> Optional[List[Dict]]:
    """
    Detect book edges - returns 6 points for two-page layout.
    Points: [TL, GT, TR, BR, GB, BL]
    """
    height, width = img.shape[:2]
    
    # First detect outer edges
    outer_corners = detect_document_edges_cv(img)
    
    if not outer_corners:
        return None
    
    # Now detect the gutter (center line)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Focus on the center region for gutter detection
    center_start = int(width * 0.35)
    center_end = int(width * 0.65)
    center_region = gray[:, center_start:center_end]
    
    # Apply edge detection to find vertical lines
    edges = cv2.Canny(center_region, 50, 150)
    
    # Use Hough transform to find vertical lines
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=50, minLineLength=height * 0.3, maxLineGap=20)
    
    gutter_x = width * 0.5  # Default to center
    
    if lines is not None:
        vertical_lines = []
        for line in lines:
            x1, y1, x2, y2 = line[0]
            # Check if line is mostly vertical
            angle = abs(np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi)
            if angle > 80 or angle < 10:
                avg_x = (x1 + x2) / 2 + center_start
                vertical_lines.append(avg_x)
        
        if vertical_lines:
            gutter_x = np.median(vertical_lines)
    else:
        # Fallback: use brightness analysis
        col_intensity = np.mean(gray[:, center_start:center_end], axis=0)
        min_idx = np.argmin(col_intensity)
        gutter_x = min_idx + center_start
    
    # Calculate gutter top and bottom y positions
    # Use the detected outer corners y values
    tl_y = outer_corners[0]['y']
    bl_y = outer_corners[3]['y']
    tr_y = outer_corners[1]['y']
    br_y = outer_corners[2]['y']
    
    # Gutter top is average of top corners y, gutter bottom is average of bottom corners y
    gt_y = (tl_y + tr_y) / 2
    gb_y = (bl_y + br_y) / 2
    
    # Normalize gutter_x
    gutter_x_norm = gutter_x / width
    
    # Return 6 points: [TL, GT, TR, BR, GB, BL]
    return [
        outer_corners[0],  # TL
        {'x': gutter_x_norm, 'y': gt_y},  # GT (Gutter Top)
        outer_corners[1],  # TR
        outer_corners[2],  # BR  
        {'x': gutter_x_norm, 'y': gb_y},  # GB (Gutter Bottom)
        outer_corners[3],  # BL
    ]


@api_router.post("/images/detect-edges")
async def detect_edges_endpoint(request: EdgeDetectionRequest):
    """
    Real-time edge detection for document scanning.
    
    Returns detected corner points for:
    - document mode: 4 points [TL, TR, BR, BL]
    - book mode: 6 points [TL, GT, TR, BR, GB, BL]
    - id_card mode: 4 points [TL, TR, BR, BL]
    
    All points are normalized to 0-1 range.
    If no edges detected, returns default frame with 5% margins.
    """
    try:
        # Decode image
        img_data = request.image_base64
        if "," in img_data:
            img_data = img_data.split(",")[1]
        
        image_bytes = base64.b64decode(img_data)
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return {"success": False, "message": "Could not decode image"}
        
        height, width = img.shape[:2]
        
        if request.mode == "book":
            points = detect_book_edges_cv(img)
            point_count = 6
            # Default book points with gutter in middle
            default_points = [
                {"x": 0.05, "y": 0.05},  # TL
                {"x": 0.5, "y": 0.05},   # GT (Gutter Top)
                {"x": 0.95, "y": 0.05},  # TR
                {"x": 0.95, "y": 0.95},  # BR
                {"x": 0.5, "y": 0.95},   # GB (Gutter Bottom)
                {"x": 0.05, "y": 0.95},  # BL
            ]
        else:
            points = detect_document_edges_cv(img)
            point_count = 4
            # Default document points with 5% margins
            default_points = [
                {"x": 0.05, "y": 0.05},  # TL
                {"x": 0.95, "y": 0.05},  # TR
                {"x": 0.95, "y": 0.95},  # BR
                {"x": 0.05, "y": 0.95},  # BL
            ]
        
        if points and len(points) == point_count:
            logger.info(f"[EdgeDetect] Successfully detected {point_count} points for {request.mode} mode")
            return {
                "success": True,
                "points": points,
                "image_size": {"width": width, "height": height},
                "auto_detected": True
            }
        else:
            logger.info(f"[EdgeDetect] No edges detected for {request.mode} mode, returning default frame")
            return {
                "success": True,  # Return success with default points so UI can use them
                "points": default_points,
                "image_size": {"width": width, "height": height},
                "auto_detected": False,
                "message": "Document edges not clearly detected, using default frame"
            }
            
    except Exception as e:
        logger.error(f"[EdgeDetect] Error: {e}")
        return {
            "success": False,
            "message": str(e)
        }


@api_router.post("/images/perspective-crop-public")
async def public_perspective_crop(request: ManualCropRequest):
    """
    Public endpoint for perspective crop - no authentication required.
    Used for guest mode scanning.
    Handles EXIF orientation and applies perspective transform.
    """
    try:
        # Decode image
        if "," in request.image_base64:
            img_data = request.image_base64.split(",")[1]
        else:
            img_data = request.image_base64
            
        image_bytes = base64.b64decode(img_data)
        
        # Fix EXIF orientation FIRST using PIL
        corrected_bytes, was_rotated, orientation = fix_image_orientation_pil(image_bytes)
        logger.info(f"[Public] EXIF orientation: {orientation}, was_rotated: {was_rotated}")
        
        # Decode the corrected image with OpenCV
        nparr = np.frombuffer(corrected_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return {"success": False, "cropped_image_base64": request.image_base64, "message": "Could not decode image"}
        
        height, width = img.shape[:2]
        logger.info(f"[Public] Image dimensions after EXIF fix: {width}x{height}")
        
        # Handle Android cameras that don't set EXIF properly
        # If image is landscape (width > height) and force_portrait is True, rotate it
        if request.force_portrait and width > height:
            logger.info(f"[Public] Force portrait: rotating landscape image 90° CCW")
            img = cv2.rotate(img, cv2.ROTATE_90_COUNTERCLOCKWISE)
            height, width = img.shape[:2]
            # Re-encode the rotated image
            _, corrected_buffer = cv2.imencode('.jpg', img, [cv2.IMWRITE_JPEG_QUALITY, 95])
            corrected_bytes = corrected_buffer.tobytes()
            logger.info(f"[Public] New dimensions after forced portrait: {width}x{height}")
        
        # Convert normalized corners to pixel coordinates
        pixel_corners = []
        for corner in request.corners:
            px = float(corner.get('x', 0)) * width
            py = float(corner.get('y', 0)) * height
            pixel_corners.append({'x': px, 'y': py})
        
        # Re-encode and apply perspective crop
        corrected_base64 = base64.b64encode(corrected_bytes).decode()
        cropped = perspective_crop(corrected_base64, pixel_corners)
        
        return {
            "success": True,
            "cropped_image_base64": cropped
        }
    except Exception as e:
        logger.error(f"[Public] Manual crop error: {e}")
        return {
            "success": False,
            "cropped_image_base64": request.image_base64,
            "message": str(e)
        }

# ==================== OCR ENDPOINTS ====================

@api_router.post("/ocr/extract", response_model=OCRResponse)
async def extract_text(
    ocr_request: OCRRequest,
    current_user: User = Depends(get_current_user)
):
    """Extract text from image using OpenAI Vision OCR"""
    # Check OCR usage limits for free users
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    is_premium = current_user.subscription_type == "premium"
    
    if not is_premium:
        if current_user.ocr_usage_date == today and current_user.ocr_usage_today >= 5:
            raise HTTPException(
                status_code=429,
                detail="Daily OCR limit reached. Upgrade to Premium for unlimited OCR."
            )
    
    # Update OCR usage
    if current_user.ocr_usage_date == today:
        await db.users.update_one(
            {"user_id": current_user.user_id},
            {"$inc": {"ocr_usage_today": 1}}
        )
    else:
        await db.users.update_one(
            {"user_id": current_user.user_id},
            {"$set": {"ocr_usage_today": 1, "ocr_usage_date": today}}
        )
    
    # Perform actual OCR using Emergent LLM
    extracted_text = await perform_ocr_with_openai(ocr_request.image_base64)
    
    return OCRResponse(
        text=extracted_text,
        confidence=0.95
    )

@api_router.post("/documents/{document_id}/ocr")
async def run_document_ocr(
    document_id: str,
    ocr_text: str,
    page_index: int = 0,
    current_user: User = Depends(get_current_user)
):
    """Save OCR text for a document page (text extracted on device)"""
    document = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    pages = document.get("pages", [])
    if page_index >= len(pages):
        raise HTTPException(status_code=400, detail="Invalid page index")
    
    # Update page OCR text
    pages[page_index]["ocr_text"] = ocr_text
    
    # Update full document OCR text
    full_text = " ".join([p.get("ocr_text", "") for p in pages if p.get("ocr_text")])
    
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {
            "pages": pages,
            "ocr_full_text": full_text,
            "updated_at": datetime.now(timezone.utc)
        }}
    )
    
    return {"message": "OCR text saved successfully", "full_text": full_text}

# ==================== SIGNATURE ENDPOINTS ====================

class SignatureOverlayRequest(BaseModel):
    image_base64: str
    signature_base64: str
    position_x: float  # Normalized 0-1 position from left
    position_y: float  # Normalized 0-1 position from top
    scale: float = 0.3  # Signature scale relative to image width

@api_router.post("/images/add-signature")
async def add_signature_to_image(
    request: SignatureOverlayRequest,
    current_user: User = Depends(get_current_user)
):
    """Overlay a signature on an image at the specified position"""
    try:
        from PIL import Image
        import io
        
        # Decode base image
        img_data = request.image_base64
        if ',' in img_data:
            img_data = img_data.split(',')[1]
        
        image_bytes = base64.b64decode(img_data)
        base_image = Image.open(io.BytesIO(image_bytes)).convert('RGBA')
        
        # Decode signature image
        sig_data = request.signature_base64
        if ',' in sig_data:
            sig_data = sig_data.split(',')[1]
        
        sig_bytes = base64.b64decode(sig_data)
        
        # Check if it's an SVG (starts with XML declaration or svg tag)
        try:
            sig_str = sig_bytes.decode('utf-8', errors='ignore')
            if sig_str.strip().startswith('<') or '<?xml' in sig_str[:100] or '<svg' in sig_str[:500]:
                # It's an SVG, try to convert to PNG using cairosvg
                try:
                    import cairosvg
                    png_bytes = cairosvg.svg2png(bytestring=sig_bytes, output_width=400)
                    signature = Image.open(io.BytesIO(png_bytes)).convert('RGBA')
                    logger.info("Converted SVG signature to PNG")
                except ImportError:
                    # cairosvg not available, try svglib
                    try:
                        from svglib.svglib import svg2rlg
                        from reportlab.graphics import renderPM
                        drawing = svg2rlg(io.BytesIO(sig_bytes))
                        png_bytes = io.BytesIO()
                        renderPM.drawToFile(drawing, png_bytes, fmt="PNG")
                        png_bytes.seek(0)
                        signature = Image.open(png_bytes).convert('RGBA')
                        logger.info("Converted SVG signature to PNG using svglib")
                    except Exception as svg_error:
                        logger.error(f"SVG conversion failed: {svg_error}")
                        # Create a simple placeholder signature
                        signature = Image.new('RGBA', (300, 150), (0, 0, 0, 0))
                        from PIL import ImageDraw, ImageFont
                        draw = ImageDraw.Draw(signature)
                        draw.text((50, 50), "Signature", fill=(0, 0, 0, 255))
            else:
                # Regular image (PNG/JPEG)
                signature = Image.open(io.BytesIO(sig_bytes)).convert('RGBA')
        except Exception as decode_error:
            # Try as regular image
            logger.warning(f"Decoding error: {decode_error}, trying as image")
            signature = Image.open(io.BytesIO(sig_bytes)).convert('RGBA')
        
        # Calculate signature size
        sig_width = int(base_image.width * request.scale)
        sig_height = int(signature.height * (sig_width / signature.width))
        signature = signature.resize((sig_width, sig_height), Image.Resampling.LANCZOS)
        
        # Calculate position
        pos_x = int(request.position_x * base_image.width - sig_width / 2)
        pos_y = int(request.position_y * base_image.height - sig_height / 2)
        
        # Clamp to bounds
        pos_x = max(0, min(base_image.width - sig_width, pos_x))
        pos_y = max(0, min(base_image.height - sig_height, pos_y))
        
        # Create composite
        composite = base_image.copy()
        composite.paste(signature, (pos_x, pos_y), signature)
        
        # Convert back to RGB for JPEG
        if composite.mode == 'RGBA':
            background = Image.new('RGB', composite.size, (255, 255, 255))
            background.paste(composite, mask=composite.split()[3])
            composite = background
        
        # Encode result
        buffer = io.BytesIO()
        composite.save(buffer, format='JPEG', quality=95)
        result_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        logger.info(f"Signature added at ({pos_x}, {pos_y}) with scale {request.scale}")
        
        return {
            "success": True,
            "signed_image_base64": result_base64,
            "message": "Signature added successfully"
        }
        
    except Exception as e:
        logger.error(f"Signature overlay error: {str(e)}")
        return {
            "success": False,
            "message": f"Failed to add signature: {str(e)}"
        }


# ==================== ANNOTATION ENDPOINT ====================

class AnnotationItem(BaseModel):
    id: str
    type: str  # freehand, text, arrow, rectangle, circle, highlight
    color: str
    strokeWidth: float
    x: Optional[float] = 0  # Optional for freehand/highlight that use points array
    y: Optional[float] = 0
    width: Optional[float] = None
    height: Optional[float] = None
    endX: Optional[float] = None
    endY: Optional[float] = None
    text: Optional[str] = None
    points: Optional[List[Dict[str, float]]] = None

class ApplyAnnotationsRequest(BaseModel):
    image_base64: str
    annotations: List[AnnotationItem]
    display_width: Optional[float] = None  # Width of the image as displayed on screen
    display_height: Optional[float] = None  # Height of the image as displayed on screen

@api_router.post("/images/apply-annotations")
async def apply_annotations_to_image(request: ApplyAnnotationsRequest):
    """Apply SVG-like annotations to an image"""
    try:
        from PIL import Image, ImageDraw, ImageFont
        import io
        import math
        
        # Decode base image
        img_data = request.image_base64
        if ',' in img_data:
            img_data = img_data.split(',')[1]
        
        image_bytes = base64.b64decode(img_data)
        base_image = Image.open(io.BytesIO(image_bytes)).convert('RGBA')
        
        # Create a transparent overlay for annotations
        overlay = Image.new('RGBA', base_image.size, (0, 0, 0, 0))
        draw = ImageDraw.Draw(overlay)
        
        # Parse color from hex/rgba
        def parse_color(color_str, alpha=255):
            """Parse color string to RGBA tuple"""
            color_str = color_str.strip()
            if color_str.startswith('#'):
                hex_color = color_str.lstrip('#')
                if len(hex_color) == 8:  # Has alpha
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    a = int(hex_color[6:8], 16)
                    return (r, g, b, a)
                elif len(hex_color) == 6:
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    return (r, g, b, alpha)
                elif len(hex_color) == 3:
                    r = int(hex_color[0], 16) * 17
                    g = int(hex_color[1], 16) * 17
                    b = int(hex_color[2], 16) * 17
                    return (r, g, b, alpha)
            elif color_str.startswith('rgba'):
                # Parse rgba(r,g,b,a)
                import re
                match = re.match(r'rgba?\((\d+),\s*(\d+),\s*(\d+)(?:,\s*([\d.]+))?\)', color_str)
                if match:
                    r = int(match.group(1))
                    g = int(match.group(2))
                    b = int(match.group(3))
                    a = int(float(match.group(4) or 1) * 255)
                    return (r, g, b, a)
            return (0, 0, 0, alpha)  # Default to black
        
        # The annotations are in screen coordinates, we need to scale to image coordinates
        # The frontend displays the image in a container, so we need to scale based on display size
        img_width, img_height = base_image.size
        
        # Calculate scale factor if display dimensions are provided
        scale_x = 1.0
        scale_y = 1.0
        if request.display_width and request.display_height and request.display_width > 0 and request.display_height > 0:
            scale_x = img_width / request.display_width
            scale_y = img_height / request.display_height
            logger.info(f"Scaling annotations: display({request.display_width}x{request.display_height}) -> image({img_width}x{img_height}), scale({scale_x:.2f}x{scale_y:.2f})")
        
        # Helper function to scale coordinates
        def scale_point(x, y):
            return (int(x * scale_x), int(y * scale_y))
        
        # Process each annotation
        for annotation in request.annotations:
            color = parse_color(annotation.color)
            stroke_width = max(1, int(annotation.strokeWidth))
            
            if annotation.type in ['freehand', 'highlight']:
                if annotation.points and len(annotation.points) >= 2:
                    # Scale all points
                    points = [scale_point(p['x'], p['y']) for p in annotation.points]
                    draw.line(points, fill=color, width=stroke_width, joint='curve')
                    
            elif annotation.type == 'arrow':
                if annotation.endX is not None and annotation.endY is not None:
                    start = scale_point(annotation.x, annotation.y)
                    end = scale_point(annotation.endX, annotation.endY)
                    
                    # Draw the main line
                    draw.line([start, end], fill=color, width=stroke_width)
                    
                    # Draw arrowhead
                    angle = math.atan2(end[1] - start[1], end[0] - start[0])
                    arrow_length = 15
                    arrow_angle = math.pi / 6
                    
                    point1 = (
                        end[0] - arrow_length * math.cos(angle - arrow_angle),
                        end[1] - arrow_length * math.sin(angle - arrow_angle)
                    )
                    point2 = (
                        end[0] - arrow_length * math.cos(angle + arrow_angle),
                        end[1] - arrow_length * math.sin(angle + arrow_angle)
                    )
                    
                    draw.line([end, point1], fill=color, width=stroke_width)
                    draw.line([end, point2], fill=color, width=stroke_width)
                    
            elif annotation.type == 'rectangle':
                if annotation.width and annotation.height:
                    x1, y1 = scale_point(min(annotation.x, annotation.endX or annotation.x), 
                                         min(annotation.y, annotation.endY or annotation.y))
                    x2 = x1 + int(annotation.width * scale_x)
                    y2 = y1 + int(annotation.height * scale_y)
                    draw.rectangle([x1, y1, x2, y2], outline=color, width=stroke_width)
                    
            elif annotation.type == 'circle':
                if annotation.width and annotation.height:
                    x1, y1 = scale_point(min(annotation.x, annotation.endX or annotation.x), 
                                         min(annotation.y, annotation.endY or annotation.y))
                    x2 = x1 + int(annotation.width * scale_x)
                    y2 = y1 + int(annotation.height * scale_y)
                    # Draw ellipse that fits the bounding box
                    draw.ellipse([x1, y1, x2, y2], outline=color, width=stroke_width)
                    
            elif annotation.type == 'text':
                if annotation.text:
                    try:
                        # Try to load a font; fallback to default
                        font_size = max(12, int(annotation.strokeWidth * max(scale_x, scale_y)))
                        try:
                            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
                        except:
                            font = ImageFont.load_default()
                        text_pos = scale_point(annotation.x, annotation.y)
                        draw.text(text_pos, annotation.text, fill=color, font=font)
                    except Exception as text_error:
                        logger.warning(f"Failed to draw text: {text_error}")
        
        # Composite the overlay onto the base image
        result = Image.alpha_composite(base_image, overlay)
        
        # Convert back to RGB for JPEG
        if result.mode == 'RGBA':
            background = Image.new('RGB', result.size, (255, 255, 255))
            background.paste(result, mask=result.split()[3])
            result = background
        
        # Encode result
        buffer = io.BytesIO()
        result.save(buffer, format='JPEG', quality=95)
        result_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        logger.info(f"Applied {len(request.annotations)} annotations to image")
        
        return {
            "success": True,
            "annotated_image_base64": result_base64,
            "message": f"Applied {len(request.annotations)} annotations successfully"
        }
        
    except Exception as e:
        logger.error(f"Annotation error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"Failed to apply annotations: {str(e)}"
        }


class ApplySignatureRequest(BaseModel):
    image_base64: str
    signature_base64: str
    position_x: float  # 0-1 normalized position
    position_y: float  # 0-1 normalized position
    scale: float = 0.3  # Size relative to image width

@api_router.post("/images/apply-signature")
async def apply_signature_to_image(request: ApplySignatureRequest):
    """Burn a signature image onto the document image"""
    try:
        from PIL import Image
        import io
        
        # Decode base image
        img_data = request.image_base64
        if ',' in img_data:
            img_data = img_data.split(',')[1]
        
        image_bytes = base64.b64decode(img_data)
        base_image = Image.open(io.BytesIO(image_bytes)).convert('RGBA')
        img_width, img_height = base_image.size
        
        # Decode signature image
        sig_data = request.signature_base64
        if ',' in sig_data:
            sig_data = sig_data.split(',')[1]
        
        sig_bytes = base64.b64decode(sig_data)
        
        # Check if it's an SVG (starts with XML declaration or svg tag)
        try:
            sig_str = sig_bytes.decode('utf-8', errors='ignore')
            if sig_str.strip().startswith('<') or '<?xml' in sig_str[:100] or '<svg' in sig_str[:500]:
                # It's an SVG, try to convert to PNG using cairosvg
                try:
                    import cairosvg
                    png_bytes = cairosvg.svg2png(bytestring=sig_bytes, output_width=400)
                    signature = Image.open(io.BytesIO(png_bytes)).convert('RGBA')
                    logger.info("Converted SVG signature to PNG")
                except ImportError:
                    # cairosvg not available, try svglib
                    try:
                        from svglib.svglib import svg2rlg
                        from reportlab.graphics import renderPM
                        drawing = svg2rlg(io.BytesIO(sig_bytes))
                        png_bytes = io.BytesIO()
                        renderPM.drawToFile(drawing, png_bytes, fmt="PNG")
                        png_bytes.seek(0)
                        signature = Image.open(png_bytes).convert('RGBA')
                        logger.info("Converted SVG signature to PNG using svglib")
                    except Exception as svg_error:
                        logger.error(f"SVG conversion failed: {svg_error}")
                        # Create a simple placeholder signature
                        signature = Image.new('RGBA', (300, 150), (0, 0, 0, 0))
                        from PIL import ImageDraw, ImageFont
                        draw = ImageDraw.Draw(signature)
                        draw.text((50, 50), "Signature", fill=(0, 0, 0, 255))
            else:
                # Regular image (PNG/JPEG)
                signature = Image.open(io.BytesIO(sig_bytes)).convert('RGBA')
        except Exception as decode_error:
            # Try as regular image
            logger.warning(f"Decoding error: {decode_error}, trying as image")
            signature = Image.open(io.BytesIO(sig_bytes)).convert('RGBA')
        
        # Log original signature dimensions
        logger.info(f"Original signature size: {signature.width}x{signature.height}")
        
        # Calculate signature size based on scale
        if signature.width <= 0 or signature.height <= 0:
            return {
                "success": False,
                "message": "Invalid signature image dimensions"
            }
        
        sig_width = max(int(img_width * request.scale), 10)  # Minimum 10px
        sig_height = max(int(sig_width * (signature.height / signature.width)), 10)  # Minimum 10px
        
        logger.info(f"Resizing signature to: {sig_width}x{sig_height} (scale: {request.scale})")
        signature = signature.resize((sig_width, sig_height), Image.Resampling.LANCZOS)
        
        # Calculate position (position is center point in 0-1 normalized coords)
        x = int(request.position_x * img_width - sig_width / 2)
        y = int(request.position_y * img_height - sig_height / 2)
        
        # Clamp to image bounds
        x = max(0, min(x, img_width - sig_width))
        y = max(0, min(y, img_height - sig_height))
        
        # Paste signature onto image
        base_image.paste(signature, (x, y), signature)
        
        # Convert back to JPEG
        rgb_image = base_image.convert('RGB')
        buffer = io.BytesIO()
        rgb_image.save(buffer, format='JPEG', quality=95)
        result_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        logger.info(f"Applied signature at ({request.position_x:.2f}, {request.position_y:.2f})")
        
        return {
            "success": True,
            "image_base64": result_base64,
            "message": "Signature applied successfully"
        }
        
    except Exception as e:
        logger.error(f"Signature error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"Failed to apply signature: {str(e)}"
        }


class ApplyFilterRequest(BaseModel):
    image_base64: str
    filter_type: str  # original, grayscale, bw, enhanced, document
    brightness: float = 0  # -50 to 50
    contrast: float = 0    # -50 to 50
    saturation: float = 0  # -50 to 50

@api_router.post("/images/apply-filter")
async def apply_filter_to_image(request: ApplyFilterRequest):
    """Apply filter and adjustments to an image"""
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        import io
        
        # Decode base image
        img_data = request.image_base64
        if ',' in img_data:
            img_data = img_data.split(',')[1]
        
        image_bytes = base64.b64decode(img_data)
        image = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        
        # Apply filter based on type
        if request.filter_type == 'grayscale':
            image = image.convert('L').convert('RGB')
        elif request.filter_type == 'bw':
            # Black and white (high contrast grayscale)
            image = image.convert('L')
            # Apply threshold for B&W effect
            threshold = 128
            image = image.point(lambda x: 255 if x > threshold else 0, mode='1')
            image = image.convert('RGB')
        elif request.filter_type == 'enhanced':
            # Auto enhance
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.2)
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.3)
        elif request.filter_type == 'document':
            # Document mode - increase contrast and brightness
            image = image.convert('L')  # Grayscale first
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(1.1)
            image = image.convert('RGB')
        
        # Apply brightness adjustment (-50 to 50 -> 0.5 to 1.5)
        if request.brightness != 0:
            factor = 1 + (request.brightness / 50)  # -50 -> 0, 0 -> 1, 50 -> 2
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(factor)
        
        # Apply contrast adjustment (-50 to 50 -> 0.5 to 1.5)
        if request.contrast != 0:
            factor = 1 + (request.contrast / 50)
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(factor)
        
        # Apply saturation adjustment (-50 to 50 -> 0 to 2)
        if request.saturation != 0:
            factor = 1 + (request.saturation / 50)
            enhancer = ImageEnhance.Color(image)
            image = enhancer.enhance(factor)
        
        # Convert back to base64
        buffer = io.BytesIO()
        image.save(buffer, format='JPEG', quality=90)
        result_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        logger.info(f"Applied filter '{request.filter_type}' with adjustments B:{request.brightness} C:{request.contrast} S:{request.saturation}")
        
        return {
            "success": True,
            "image_base64": result_base64,
            "message": f"Filter '{request.filter_type}' applied successfully"
        }
        
    except Exception as e:
        logger.error(f"Filter error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"Failed to apply filter: {str(e)}"
        }


# ==================== PDF PASSWORD PROTECTION ====================

class PDFPasswordRequest(BaseModel):
    pdf_base64: str
    password: str

@api_router.post("/pdf/protect")
async def protect_pdf_with_password(request: PDFPasswordRequest):
    """Add password protection to a PDF using pypdf"""
    try:
        from pypdf import PdfReader, PdfWriter
        import io
        
        # Decode base64 PDF
        pdf_data = request.pdf_base64
        if ',' in pdf_data:
            pdf_data = pdf_data.split(',')[1]
        
        pdf_bytes = base64.b64decode(pdf_data)
        
        # Read the PDF
        reader = PdfReader(io.BytesIO(pdf_bytes))
        writer = PdfWriter()
        
        # Copy all pages to writer
        for page in reader.pages:
            writer.add_page(page)
        
        # Add password protection (AES-256)
        writer.encrypt(
            user_password=request.password,
            owner_password=request.password,
            permissions_flag=0b11110100,  # Allow printing, filling forms
        )
        
        # Save to bytes
        output = io.BytesIO()
        writer.write(output)
        
        # Convert to base64
        result_base64 = base64.b64encode(output.getvalue()).decode('utf-8')
        
        logger.info("PDF encrypted with password successfully")
        
        return {
            "success": True,
            "pdf_base64": result_base64,
            "message": "PDF protected with password successfully"
        }
        
    except Exception as e:
        logger.error(f"PDF protection error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"Failed to protect PDF: {str(e)}"
        }


# ==================== EXPORT ENDPOINTS ====================

class ExportRequest(BaseModel):
    format: str  # pdf, jpeg, docx, xlsx
    page_indices: Optional[List[int]] = None  # None means all pages
    include_ocr: bool = False

class ExportResponse(BaseModel):
    file_base64: str
    filename: str
    mime_type: str

class PublicExportRequest(BaseModel):
    """Export request for guest/local documents - includes image data directly"""
    format: str = "pdf"
    images_base64: List[str]
    document_name: str = "Document"

@api_router.post("/export/public")
async def export_public_pdf(request: PublicExportRequest):
    """
    Public endpoint for PDF export - no authentication required.
    Used for guest/local documents where we send image data directly.
    """
    try:
        if not request.images_base64:
            return {"success": False, "message": "No images provided"}
        
        if request.format.lower() == "pdf":
            pdf_bytes = create_pdf_from_images(request.images_base64, None)
            
            return {
                "success": True,
                "file_base64": base64.b64encode(pdf_bytes).decode(),
                "filename": f"{request.document_name.replace(' ', '_')}.pdf",
                "mime_type": "application/pdf"
            }
        else:
            # For single image export
            img_base64 = request.images_base64[0]
            if "," in img_base64:
                img_base64 = img_base64.split(",")[1]
            
            return {
                "success": True,
                "file_base64": img_base64,
                "filename": f"{request.document_name.replace(' ', '_')}.jpg",
                "mime_type": "image/jpeg"
            }
            
    except Exception as e:
        logger.error(f"Public export error: {e}")
        return {
            "success": False,
            "message": f"Export failed: {str(e)}"
        }

def create_pdf_from_images(images_base64: List[str], include_text: List[str] = None) -> bytes:
    """Create a PDF from base64 images"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    for i, img_base64 in enumerate(images_base64):
        if "," in img_base64:
            img_base64 = img_base64.split(",")[1]
        
        img_data = base64.b64decode(img_base64)
        img = Image.open(BytesIO(img_data))
        
        # Calculate scaling to fit A4
        img_width, img_height = img.size
        scale = min(width / img_width, height / img_height) * 0.9  # 90% of page
        new_width = img_width * scale
        new_height = img_height * scale
        
        # Center image on page
        x = (width - new_width) / 2
        y = (height - new_height) / 2
        
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        img_buffer = BytesIO()
        img.save(img_buffer, format='JPEG', quality=90)
        img_buffer.seek(0)
        
        c.drawImage(ImageReader(img_buffer), x, y, new_width, new_height)
        
        # Add OCR text if available
        if include_text and i < len(include_text) and include_text[i]:
            c.setFont("Helvetica", 8)
            text_y = 30
            for line in include_text[i].split('\n')[:5]:  # First 5 lines as footer
                c.drawString(40, text_y, line[:80])  # Truncate long lines
                text_y -= 10
        
        if i < len(images_base64) - 1:
            c.showPage()
    
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_docx_from_text(texts: List[str], title: str) -> bytes:
    """Create a Word document from OCR text"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    
    doc = DocxDocument()
    doc.add_heading(title, 0)
    
    for i, text in enumerate(texts):
        if i > 0:
            doc.add_page_break()
        
        doc.add_heading(f'Page {i + 1}', level=1)
        
        if text:
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.style.font.size = Pt(11)
        else:
            doc.add_paragraph('(No text extracted for this page)')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_xlsx_from_text(texts: List[str], title: str) -> bytes:
    """Create an Excel document from OCR text"""
    from openpyxl import Workbook
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Extracted Text"
    
    # Header
    ws['A1'] = 'Page'
    ws['B1'] = 'Extracted Text'
    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 100
    
    row = 2
    for i, text in enumerate(texts):
        ws.cell(row=row, column=1, value=f'Page {i + 1}')
        ws.cell(row=row, column=2, value=text or '(No text)')
        row += 1
    
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@api_router.post("/documents/{document_id}/export", response_model=ExportResponse)
async def export_document(
    document_id: str,
    export_request: ExportRequest,
    current_user: User = Depends(get_current_user)
):
    """Export document in various formats"""
    document = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    pages = document.get("pages", [])
    page_indices = export_request.page_indices or list(range(len(pages)))
    
    # Filter pages
    selected_pages = [pages[i] for i in page_indices if i < len(pages)]
    
    if not selected_pages:
        raise HTTPException(status_code=400, detail="No pages to export")
    
    doc_name = document.get("name", "document").replace(" ", "_")
    format_type = export_request.format.lower()
    
    # Helper function to get image data from page (either base64 or download from S3)
    async def get_image_data(page: dict) -> str:
        """Get image data as base64, downloading from S3 if needed"""
        # First check if we have base64 data
        if page.get("image_base64"):
            img_data = page["image_base64"]
            if "," in img_data:
                img_data = img_data.split(",")[1]
            # Validate it's actual base64
            if len(img_data) > 100:  # Should be more than 100 chars for a real image
                return img_data
        
        # Otherwise download from S3 URL
        if page.get("image_url") and s3_client:
            try:
                import urllib.parse
                import httpx
                
                url = page["image_url"]
                logger.info(f"Downloading image from S3 URL: {url[:100]}...")
                
                # Try direct HTTP download first (simpler and works with pre-signed URLs)
                async with httpx.AsyncClient() as client:
                    response = await client.get(url, timeout=30.0)
                    if response.status_code == 200:
                        image_data = response.content
                        logger.info(f"Downloaded {len(image_data)} bytes from S3")
                        return base64.b64encode(image_data).decode('utf-8')
                    else:
                        logger.error(f"HTTP download failed: {response.status_code}")
                
                # Fallback to S3 SDK
                parsed = urllib.parse.urlparse(url)
                key = parsed.path.lstrip('/')
                
                # Download from S3
                response = s3_client.get_object(Bucket=AWS_S3_BUCKET_NAME, Key=key)
                image_data = response['Body'].read()
                logger.info(f"Downloaded {len(image_data)} bytes from S3 SDK")
                return base64.b64encode(image_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to download image from S3: {e}")
                import traceback
                traceback.print_exc()
        
        logger.error(f"No image data found for page. Has base64: {bool(page.get('image_base64'))}, Has URL: {bool(page.get('image_url'))}")
        return ""
    
    try:
        if format_type == "pdf":
            # Get images, downloading from S3 if needed
            images = []
            for idx, p in enumerate(selected_pages):
                img_data = await get_image_data(p)
                if not img_data:
                    raise HTTPException(
                        status_code=520, 
                        detail=f"Page {idx + 1} has no image data. Has URL: {bool(p.get('image_url'))}, Has base64: {bool(p.get('image_base64'))}"
                    )
                images.append(img_data)
            
            texts = [p.get("ocr_text", "") for p in selected_pages] if export_request.include_ocr else None
            
            pdf_bytes = create_pdf_from_images(images, texts)
            
            return ExportResponse(
                file_base64=base64.b64encode(pdf_bytes).decode(),
                filename=f"{doc_name}.pdf",
                mime_type="application/pdf"
            )
        
        elif format_type == "jpeg":
            # Export single page or first page as JPEG
            page = selected_pages[0]
            img_base64 = await get_image_data(page)
            
            return ExportResponse(
                file_base64=img_base64,
                filename=f"{doc_name}_page1.jpg",
                mime_type="image/jpeg"
            )
        
        elif format_type == "docx":
            if not export_request.include_ocr:
                raise HTTPException(status_code=400, detail="DOCX export requires OCR text. Set include_ocr=true")
            
            texts = [p.get("ocr_text", "") for p in selected_pages]
            docx_bytes = create_docx_from_text(texts, document.get("name", "Document"))
            
            return ExportResponse(
                file_base64=base64.b64encode(docx_bytes).decode(),
                filename=f"{doc_name}.docx",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        elif format_type == "xlsx":
            if not export_request.include_ocr:
                raise HTTPException(status_code=400, detail="XLSX export requires OCR text. Set include_ocr=true")
            
            texts = [p.get("ocr_text", "") for p in selected_pages]
            xlsx_bytes = create_xlsx_from_text(texts, document.get("name", "Document"))
            
            return ExportResponse(
                file_base64=base64.b64encode(xlsx_bytes).decode(),
                filename=f"{doc_name}.xlsx",
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format_type}")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# ==================== BASIC ENDPOINTS ====================

@api_router.get("/")
async def root():
    return {"message": "Document Scanner API", "version": "1.0.0"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy"}

# Note: app.include_router is called at the end of the file after all routes are defined

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_db_client():
    """Initialize MongoDB collections on startup"""
    try:
        # Test connection with shorter timeout
        await client.admin.command('ping')
        logger.info(f"✅ Connected to MongoDB Atlas: {os.environ.get('DB_NAME', 'scanup')}")
        
        # Get list of existing collections
        existing_collections = await db.list_collection_names()
        logger.info(f"📋 Existing collections: {existing_collections}")
        
        # Define required collections with their indexes
        required_collections = {
            'users': [
                ('user_id', True),  # (field, unique)
                ('email', True),
            ],
            'documents': [
                ('document_id', True),
                ('user_id', False),
            ],
            'folders': [
                ('folder_id', True),
                ('user_id', False),
            ],
            'signatures': [
                ('signature_id', True),
                ('user_id', False),
            ],
        }
        
        for collection_name, indexes in required_collections.items():
            if collection_name not in existing_collections:
                # Create collection
                await db.create_collection(collection_name)
                logger.info(f"✅ Created collection: {collection_name}")
            
            # Create indexes
            collection = db[collection_name]
            for field, unique in indexes:
                try:
                    await collection.create_index(field, unique=unique)
                    logger.info(f"  📌 Index on {collection_name}.{field} (unique={unique})")
                except Exception as idx_err:
                    logger.warning(f"  ⚠️ Index {field} might already exist: {idx_err}")
        
        logger.info("✅ MongoDB collections initialized successfully")
        
        # Initialize content management collections (languages, translations, legal pages)
        await init_content_collections()
        
    except Exception as e:
        logger.warning(f"⚠️ MongoDB Atlas connection issue (will retry on demand): {e}")
        # Don't raise - let app start and handle DB errors at request time

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


# ==================== CONTENT MANAGEMENT & TRANSLATIONS ====================

# Default supported languages
DEFAULT_LANGUAGES = [
    {"code": "en", "name": "English", "native_name": "English", "is_default": True},
    {"code": "de", "name": "German", "native_name": "Deutsch", "is_default": False},
    {"code": "fr", "name": "French", "native_name": "Français", "is_default": False},
    {"code": "es", "name": "Spanish", "native_name": "Español", "is_default": False},
]

# Default English translations (base language)
DEFAULT_TRANSLATIONS = {
    # Common
    "app_name": "ScanUp",
    "loading": "Loading...",
    "error": "Error",
    "success": "Success",
    "cancel": "Cancel",
    "save": "Save",
    "delete": "Delete",
    "edit": "Edit",
    "done": "Done",
    "ok": "OK",
    "yes": "Yes",
    "no": "No",
    "confirm": "Confirm",
    "close": "Close",
    "back": "Back",
    "next": "Next",
    "skip": "Skip",
    "retry": "Retry",
    "search": "Search",
    "settings": "Settings",
    "share": "Share",
    "print": "Print",
    
    # Onboarding
    "onboarding_scan_title": "Scan",
    "onboarding_scan_desc": "Use your phone camera easily to scan your documents",
    "onboarding_save_title": "Save",
    "onboarding_save_desc": "Save your document as PDF, create and share encrypted files",
    "get_started": "Get started",
    
    # Auth
    "sign_in": "Sign In",
    "sign_up": "Sign Up",
    "sign_out": "Sign Out",
    "email": "Email",
    "password": "Password",
    "forgot_password": "Forgot Password?",
    "continue_with_google": "Continue with Google",
    "continue_with_apple": "Continue with Apple",
    "continue_as_guest": "Continue as Guest",
    "create_account": "Create Account",
    "already_have_account": "Already have an account?",
    "dont_have_account": "Don't have an account?",
    
    # Home Screen
    "your_documents": "Your Documents",
    "documents": "Documents",
    "folders": "Folders",
    "latest": "Latest",
    "all": "All",
    "no_documents": "No Documents Yet",
    "no_documents_desc": "Scan your first document using the camera button below",
    "no_folders": "No Folders Yet",
    "no_folders_desc": "Create folders to organize your documents",
    "sort_by": "Sort by",
    "sort_az": "A-Z",
    "sort_za": "Z-A",
    "sort_newest": "Newest First",
    "sort_oldest": "Oldest First",
    "view_list": "List View",
    "view_grid": "Grid View",
    
    # Document Actions
    "rename": "Name",
    "move_to_folder": "Move to Folder",
    "set_password": "Password",
    "remove_password": "Remove Password",
    "document_protected": "Protected Document",
    "enter_password": "Enter Password",
    "wrong_password": "Wrong password. Please try again.",
    
    # Scanner
    "scan_document": "Scan Document",
    "take_photo": "Take Photo",
    "choose_from_gallery": "Choose from Gallery",
    "flash_on": "Flash On",
    "flash_off": "Flash Off",
    "auto_detect": "Auto Detect",
    "manual_crop": "Manual Crop",
    "retake": "Retake",
    "use_photo": "Use Photo",
    "add_page": "Add Page",
    "pages": "Pages",
    
    # Editor
    "edit_document": "Edit Document",
    "filters": "Filters",
    "crop": "Crop",
    "rotate": "Rotate",
    "sign": "Sign",
    "text": "Text",
    "annotations": "Annotations",
    "original": "Original",
    "grayscale": "Grayscale",
    "black_white": "Black & White",
    "enhanced": "Enhanced",
    "brightness": "Brightness",
    "contrast": "Contrast",
    "saturation": "Saturation",
    
    # Signature
    "signature": "Signature",
    "add_signature": "Add Signature",
    "draw_signature": "Draw Signature",
    "saved_signatures": "Saved Signatures",
    "apply_all_pages": "Apply all pages",
    "manage_signatures": "Manage Signatures",
    "delete_signature": "Delete Signature",
    "no_signatures": "No Saved Signatures",
    "create_first_signature": "Create your first signature",
    
    # Folders
    "new_folder": "New Folder",
    "folder_name": "Folder name",
    "folder_color": "Color",
    "create_folder": "Create Folder",
    "delete_folder": "Delete Folder",
    "folder_password": "Folder Password",
    
    # Settings
    "account": "Account",
    "appearance": "Appearance",
    "dark_mode": "Dark Mode",
    "language": "Language",
    "notifications": "Notifications",
    "storage": "Storage",
    "premium": "Premium",
    "restore_purchases": "Restore Purchases",
    "rate_app": "Rate App",
    "share_app": "Share App",
    "help_support": "Help & Support",
    "privacy_policy": "Privacy Policy",
    "terms_conditions": "Terms & Conditions",
    "about": "About",
    "version": "Version",
    
    # Premium
    "go_premium": "Go Premium",
    "premium_title": "Unlock Premium",
    "premium_subtitle": "Get access to all features",
    "premium_feature_1": "Remove watermark",
    "premium_feature_2": "Unlimited scans",
    "premium_feature_3": "Cloud backup",
    "premium_feature_4": "Priority support",
    "monthly": "Monthly",
    "yearly": "Yearly",
    "lifetime": "Lifetime",
    "start_free_trial": "Start Free Trial",
    "already_premium": "You're already Premium!",
    
    # Alerts & Messages
    "document_saved": "Document saved successfully",
    "document_deleted": "Document deleted",
    "folder_created": "Folder created successfully",
    "folder_deleted": "Folder deleted",
    "signature_saved": "Signature saved",
    "signature_applied": "Signature applied",
    "changes_saved": "Changes saved",
    "delete_confirm": "Are you sure you want to delete this?",
    "delete_document_confirm": "Delete this document?",
    "delete_folder_confirm": "Delete this folder and all its contents?",
    "unsaved_changes": "You have unsaved changes",
    "discard_changes": "Discard changes?",
    "network_error": "Network error. Please check your connection.",
    "something_went_wrong": "Something went wrong. Please try again.",
    "session_expired": "Session expired. Please sign in again.",
    
    # Time
    "just_now": "Just now",
    "minutes_ago": "{count} minutes ago",
    "hours_ago": "{count} hours ago",
    "days_ago": "{count} days ago",
    "today": "Today",
    "yesterday": "Yesterday",
    
    # Settings Page Additional
    "settings": "Settings",
    "scan_settings": "Scan Settings",
    "default_scan_quality": "Default Scan Quality",
    "default_filter": "Default Filter",
    "auto_enhance": "Auto Enhance",
    "auto_enhance_desc": "Automatically enhance scanned images",
    "feedback": "Feedback",
    "sound_effects": "Sound Effects",
    "haptic_feedback": "Haptic Feedback",
    "storage_data": "Storage & Data",
    "clear_cache": "Clear Cache",
    "clear_cache_desc": "Free up storage space",
    "reset_settings": "Reset Settings",
    "reset_settings_desc": "Restore default settings",
    "select_language": "Select Language",
    
    # Additional Keys for UI
    "deleting": "Deleting...",
    "selected": "selected",
    "untitled_document": "Untitled Document",
    "create": "Create",
    "name_placeholder": "Name",
    "color": "Color",
    "list_view": "List",
    "grid_view": "Grid",
    "view_mode": "View",
    "profile": "Profile",
    "guest_user": "Guest User",
    "sign_in_to_sync": "Sign in to sync your documents",
    
    # New keys for complete translation
    "newest": "Newest",
    "oldest": "Oldest",
    "name": "Name",
    "documents_migrated": "Documents Migrated",
    "document": "document",
    "have_been_added": "have been added",
    "has_been_added": "has been added",
    "to_your_account": "to your account",
    "delete_documents": "Delete Documents",
    "delete_confirm_multiple": "Are you sure you want to delete",
    "document_s": "document(s)",
    "failed_to_delete_documents": "Failed to delete some documents",
    "moved": "Moved",
    "to": "to",
    "folder": "folder",
    "main_library": "main library",
    "failed_to_move_document": "Failed to move document",
    "failed_to_create_folder": "Failed to create folder",
    "enter_folder_name": "Please enter a folder name",
    "failed_to_create_folder_retry": "Failed to create folder. Please try again.",
    "protected_document": "Protected Document",
    "document_password_protected_share": "This document is password protected. Enter the password to share.",
    "document_renamed": "Document renamed successfully",
    "failed_to_rename_document": "Failed to rename document",
    "no_images_to_print": "No images available to print",
    "failed_to_print_document": "Failed to print document",
    "password_set": "Password set successfully",
    "password_removed": "Password removed",
    "failed_to_update_password": "Failed to update password",
    "failed_to_delete_document": "Failed to delete document",
    "folder_renamed": "Folder renamed successfully",
    "failed_to_rename_folder": "Failed to rename folder",
    "failed_to_set_folder_password": "Failed to set folder password",
    "delete_folder_confirm": "Are you sure you want to delete",
    "documents_will_be_moved": "Documents inside will be moved to the main library.",
    "failed_to_delete_folder": "Failed to delete folder",
    "no_documents_yet": "No Documents Yet",
    "no_folders_yet": "No Folders Yet",
    "tap_scan_first_document": "Tap the scan button to scan your first document",
    "create_folders_organize": "Create folders to organize your documents",
    "newest_first": "Newest First",
    "oldest_first": "Oldest First",
    "list": "List",
    "grid": "Grid",
    "edit_name": "Edit Name",
    "creating": "Creating",
    "rename_document": "Rename Document",
    "enter_new_name": "Enter new name",
    "rename_folder": "Rename Folder",
    "set_folder_password": "Set Folder Password",
    "enter_password_protect_folder": "Enter a password to protect this folder. Leave empty to remove password.",
    "enter_password_optional": "Enter password (optional)",
    "enter_password_protect_document": "Enter a password to protect this document. Leave empty to remove password.",
    "enter_password_to_view": "Enter the password to view this document.",
    "unlock": "Unlock",
    "incorrect_password": "Incorrect Password",
    "please_try_again": "Please try again.",
    
    # Signatures
    "signatures": "Signatures",
    "delete_signature_confirm": "Are you sure you want to delete this signature?",
    "delete_selected": "Delete Selected",
    "signature_s": "signature(s)",
    "select": "Select",
    "no_signatures_yet": "No Signatures Yet",
    "create_first_signature": "Create your first signature to use when signing documents",
    "add_new": "Add new",
    "please_draw_signature": "Please draw a signature first",
    "sign_here": "Sign here",
    "draw_signature_above": "Draw your signature above",
    
    # Auth
    "please_fill_all_fields": "Please fill in all fields",
    "login_failed": "Login Failed",
    "check_credentials": "Please check your credentials",
    "google_session_failed": "Failed to get session from Google. Please try again.",
    "google_signin_failed": "Google sign-in failed. Please try again.",
    "welcome_back": "Welcome Back",
    "sign_in_to_continue": "Sign in to continue",
    "enter_your_email": "Enter your email",
    "enter_your_password": "Enter your password",
    "or_continue_with": "or continue with",
    "passwords_do_not_match": "Passwords do not match",
    "password_min_length": "Password must be at least 6 characters",
    "registration_failed": "Registration Failed",
    "login_cancelled": "Login Cancelled",
    "please_try_signing_in_again": "Please try signing in again",
    "sign_up_to_get_started": "Sign up to get started",
    "full_name": "Full Name",
    "enter_your_name": "Enter your name",
    "create_a_password": "Create a password",
    "confirm_password": "Confirm Password",
    "confirm_your_password": "Confirm your password",
    "filter_by": "Filter by",
    
    # Share Modal
    "file_format": "File format",
    "password_protect": "Password protect",
    "enter_password_for_pdf": "Enter password for PDF",
    "share_document": "Share document",
    "mail": "Mail",
    "page": "Page",
    
    # Tab Bar
    "home": "Home",
    
    # Settings - Clear Cache & Reset
    "clear_cache_warning": "This will clear ALL local documents and cached data. This cannot be undone!",
    "clear_all": "Clear All",
    "cache_cleared_restart": "All local data has been cleared. Please restart the app.",
    "failed_to_clear_cache": "Failed to clear cache",
    "reset_settings_confirm": "Are you sure you want to reset all settings to default?",
    "reset": "Reset",
    "settings_reset_success": "Settings have been reset",
    "failed_to_reset_settings": "Failed to reset settings",
}

# Default German translations
DEFAULT_TRANSLATIONS_DE = {
    "app_name": "ScanUp",
    "loading": "Laden...",
    "error": "Fehler",
    "success": "Erfolg",
    "cancel": "Abbrechen",
    "save": "Speichern",
    "delete": "Löschen",
    "edit": "Bearbeiten",
    "done": "Fertig",
    "ok": "OK",
    "yes": "Ja",
    "no": "Nein",
    "confirm": "Bestätigen",
    "close": "Schließen",
    "back": "Zurück",
    "next": "Weiter",
    "skip": "Überspringen",
    "retry": "Erneut versuchen",
    "search": "Suchen",
    "settings": "Einstellungen",
    "share": "Teilen",
    "print": "Drucken",
    
    "onboarding_scan_title": "Scannen",
    "onboarding_scan_desc": "Verwenden Sie Ihre Handykamera, um Dokumente einfach zu scannen",
    "onboarding_save_title": "Speichern",
    "onboarding_save_desc": "Speichern Sie Ihr Dokument als PDF, erstellen und teilen Sie verschlüsselte Dateien",
    "get_started": "Loslegen",
    
    "sign_in": "Anmelden",
    "sign_up": "Registrieren",
    "sign_out": "Abmelden",
    "email": "E-Mail",
    "password": "Passwort",
    "forgot_password": "Passwort vergessen?",
    "continue_with_google": "Mit Google fortfahren",
    "continue_with_apple": "Mit Apple fortfahren",
    "continue_as_guest": "Als Gast fortfahren",
    "create_account": "Konto erstellen",
    "already_have_account": "Haben Sie bereits ein Konto?",
    "dont_have_account": "Haben Sie noch kein Konto?",
    
    "your_documents": "Ihre Dokumente",
    "documents": "Dokumente",
    "folders": "Ordner",
    "latest": "Neueste",
    "all": "Alle",
    "no_documents": "Noch keine Dokumente",
    "no_documents_desc": "Scannen Sie Ihr erstes Dokument mit der Kamera-Taste unten",
    "no_folders": "Noch keine Ordner",
    "no_folders_desc": "Erstellen Sie Ordner, um Ihre Dokumente zu organisieren",
    "sort_by": "Sortieren nach",
    "sort_az": "A-Z",
    "sort_za": "Z-A",
    "sort_newest": "Neueste zuerst",
    "sort_oldest": "Älteste zuerst",
    "view_list": "Listenansicht",
    "view_grid": "Rasteransicht",
    
    "rename": "Umbenennen",
    "move_to_folder": "In Ordner verschieben",
    "set_password": "Passwort",
    "remove_password": "Passwort entfernen",
    "document_protected": "Geschütztes Dokument",
    "enter_password": "Passwort eingeben",
    "wrong_password": "Falsches Passwort. Bitte versuchen Sie es erneut.",
    
    "scan_document": "Dokument scannen",
    "take_photo": "Foto aufnehmen",
    "choose_from_gallery": "Aus Galerie wählen",
    "flash_on": "Blitz an",
    "flash_off": "Blitz aus",
    "retake": "Neu aufnehmen",
    "use_photo": "Foto verwenden",
    "add_page": "Seite hinzufügen",
    "pages": "Seiten",
    
    "edit_document": "Dokument bearbeiten",
    "filters": "Filter",
    "crop": "Zuschneiden",
    "rotate": "Drehen",
    "sign": "Unterschreiben",
    "text": "Text",
    "annotations": "Anmerkungen",
    "original": "Original",
    "grayscale": "Graustufen",
    "black_white": "Schwarz-Weiß",
    "enhanced": "Verbessert",
    "brightness": "Helligkeit",
    "contrast": "Kontrast",
    "saturation": "Sättigung",
    
    "signature": "Unterschrift",
    "add_signature": "Unterschrift hinzufügen",
    "draw_signature": "Unterschrift zeichnen",
    "saved_signatures": "Gespeicherte Unterschriften",
    "apply_all_pages": "Auf alle Seiten anwenden",
    "manage_signatures": "Unterschriften verwalten",
    "delete_signature": "Unterschrift löschen",
    "no_signatures": "Keine gespeicherten Unterschriften",
    "create_first_signature": "Erstellen Sie Ihre erste Unterschrift",
    
    "new_folder": "Neuer Ordner",
    "folder_name": "Ordnername",
    "folder_color": "Farbe",
    "create_folder": "Ordner erstellen",
    "delete_folder": "Ordner löschen",
    "folder_password": "Ordner-Passwort",
    
    "account": "Konto",
    "appearance": "Erscheinungsbild",
    "dark_mode": "Dunkelmodus",
    "language": "Sprache",
    "notifications": "Benachrichtigungen",
    "storage": "Speicher",
    "premium": "Premium",
    "restore_purchases": "Käufe wiederherstellen",
    "rate_app": "App bewerten",
    "share_app": "App teilen",
    "help_support": "Hilfe & Support",
    "privacy_policy": "Datenschutzrichtlinie",
    "terms_conditions": "Allgemeine Geschäftsbedingungen",
    "about": "Über",
    "version": "Version",
    
    "go_premium": "Premium werden",
    "premium_title": "Premium freischalten",
    "premium_subtitle": "Zugang zu allen Funktionen erhalten",
    "premium_feature_1": "Wasserzeichen entfernen",
    "premium_feature_2": "Unbegrenzte Scans",
    "premium_feature_3": "Cloud-Backup",
    "premium_feature_4": "Prioritäts-Support",
    "monthly": "Monatlich",
    "yearly": "Jährlich",
    "lifetime": "Lebenslang",
    "start_free_trial": "Kostenlose Testversion starten",
    "already_premium": "Sie sind bereits Premium!",
    
    "document_saved": "Dokument erfolgreich gespeichert",
    "document_deleted": "Dokument gelöscht",
    "folder_created": "Ordner erfolgreich erstellt",
    "folder_deleted": "Ordner gelöscht",
    "signature_saved": "Unterschrift gespeichert",
    "signature_applied": "Unterschrift angewendet",
    "changes_saved": "Änderungen gespeichert",
    "delete_confirm": "Sind Sie sicher, dass Sie dies löschen möchten?",
    "delete_document_confirm": "Dieses Dokument löschen?",
    "delete_folder_confirm": "Diesen Ordner und seinen gesamten Inhalt löschen?",
    "unsaved_changes": "Sie haben ungespeicherte Änderungen",
    "discard_changes": "Änderungen verwerfen?",
    "network_error": "Netzwerkfehler. Bitte überprüfen Sie Ihre Verbindung.",
    "something_went_wrong": "Etwas ist schief gelaufen. Bitte versuchen Sie es erneut.",
    "session_expired": "Sitzung abgelaufen. Bitte melden Sie sich erneut an.",
    
    "just_now": "Gerade eben",
    "minutes_ago": "vor {count} Minuten",
    "hours_ago": "vor {count} Stunden",
    "days_ago": "vor {count} Tagen",
    "today": "Heute",
    "yesterday": "Gestern",
    
    # Settings Page Additional
    "settings": "Einstellungen",
    "scan_settings": "Scan-Einstellungen",
    "default_scan_quality": "Standard-Scanqualität",
    "default_filter": "Standardfilter",
    "auto_enhance": "Automatische Verbesserung",
    "auto_enhance_desc": "Gescannte Bilder automatisch verbessern",
    "feedback": "Rückmeldung",
    "sound_effects": "Soundeffekte",
    "haptic_feedback": "Haptisches Feedback",
    "storage_data": "Speicher & Daten",
    "clear_cache": "Cache leeren",
    "clear_cache_desc": "Speicherplatz freigeben",
    "reset_settings": "Einstellungen zurücksetzen",
    "reset_settings_desc": "Standardeinstellungen wiederherstellen",
    "select_language": "Sprache auswählen",
    
    # Additional Keys for UI
    "deleting": "Löschen...",
    "selected": "ausgewählt",
    "untitled_document": "Unbenanntes Dokument",
    "create": "Erstellen",
    "name_placeholder": "Name",
    "color": "Farbe",
    "list_view": "Liste",
    "grid_view": "Raster",
    "view_mode": "Ansicht",
    "profile": "Profil",
    "guest_user": "Gastbenutzer",
    "sign_in_to_sync": "Melden Sie sich an, um Ihre Dokumente zu synchronisieren",
}

# Default Legal Pages Content (Markdown format)
DEFAULT_LEGAL_PAGES = {
    "terms": {
        "en": """# Terms and Conditions

**Last updated: January 2025**

## 1. Acceptance of Terms

By accessing and using ScanUp ("the App"), you accept and agree to be bound by the terms and conditions of this agreement.

## 2. Description of Service

ScanUp is a document scanning and management application that allows users to:
- Scan documents using their device camera
- Edit and enhance scanned documents
- Organize documents in folders
- Add signatures to documents
- Export documents as PDF

## 3. User Accounts

- You are responsible for maintaining the confidentiality of your account
- You agree to accept responsibility for all activities under your account
- You must provide accurate and complete information when creating an account

## 4. Privacy

Your use of ScanUp is also governed by our Privacy Policy. Please review our Privacy Policy to understand our practices.

## 5. Intellectual Property

The App and its original content, features, and functionality are owned by ScanUp and are protected by international copyright, trademark, and other intellectual property laws.

## 6. User Content

- You retain ownership of documents you create or upload
- You grant us a license to store and process your content to provide the service
- We do not access your documents except as necessary to provide the service

## 7. Premium Subscription

- Premium features require a paid subscription
- Subscriptions automatically renew unless cancelled
- Refunds are handled according to the app store policies

## 8. Limitation of Liability

ScanUp shall not be liable for any indirect, incidental, special, consequential, or punitive damages resulting from your use of the service.

## 9. Changes to Terms

We reserve the right to modify these terms at any time. We will notify users of any material changes.

## 10. Contact Us

If you have any questions about these Terms, please contact us at support@scanup.app
""",
        "de": """# Allgemeine Geschäftsbedingungen

**Zuletzt aktualisiert: Januar 2025**

## 1. Annahme der Bedingungen

Durch den Zugriff auf und die Nutzung von ScanUp ("die App") akzeptieren Sie die Bedingungen dieser Vereinbarung.

## 2. Beschreibung des Dienstes

ScanUp ist eine Anwendung zum Scannen und Verwalten von Dokumenten, mit der Benutzer:
- Dokumente mit der Gerätekamera scannen
- Gescannte Dokumente bearbeiten und verbessern
- Dokumente in Ordnern organisieren
- Unterschriften zu Dokumenten hinzufügen
- Dokumente als PDF exportieren

## 3. Benutzerkonten

- Sie sind für die Vertraulichkeit Ihres Kontos verantwortlich
- Sie erklären sich damit einverstanden, die Verantwortung für alle Aktivitäten unter Ihrem Konto zu übernehmen
- Bei der Kontoerstellung müssen Sie genaue und vollständige Informationen angeben

## 4. Datenschutz

Ihre Nutzung von ScanUp unterliegt auch unserer Datenschutzrichtlinie.

## 5. Geistiges Eigentum

Die App und ihr ursprünglicher Inhalt sind Eigentum von ScanUp und durch internationale Urheberrechts- und Markengesetze geschützt.

## 6. Benutzerinhalte

- Sie behalten das Eigentum an den von Ihnen erstellten oder hochgeladenen Dokumenten
- Sie gewähren uns eine Lizenz zur Speicherung und Verarbeitung Ihrer Inhalte
- Wir greifen nicht auf Ihre Dokumente zu, außer wenn dies zur Bereitstellung des Dienstes erforderlich ist

## 7. Premium-Abonnement

- Premium-Funktionen erfordern ein kostenpflichtiges Abonnement
- Abonnements werden automatisch verlängert, sofern sie nicht gekündigt werden
- Rückerstattungen werden gemäß den App-Store-Richtlinien abgewickelt

## 8. Haftungsbeschränkung

ScanUp haftet nicht für indirekte, zufällige oder Folgeschäden, die sich aus Ihrer Nutzung des Dienstes ergeben.

## 9. Änderungen der Bedingungen

Wir behalten uns das Recht vor, diese Bedingungen jederzeit zu ändern.

## 10. Kontakt

Bei Fragen zu diesen Bedingungen kontaktieren Sie uns bitte unter support@scanup.app
"""
    },
    "privacy": {
        "en": """# Privacy Policy

**Last updated: January 2025**

## 1. Introduction

ScanUp ("we", "our", or "us") is committed to protecting your privacy. This Privacy Policy explains how we collect, use, and share information about you.

## 2. Information We Collect

### Information You Provide
- Account information (email, name)
- Documents you scan and upload
- Signatures you create

### Information Collected Automatically
- Device information (model, OS version)
- Usage data (features used, scan count)
- Crash reports and diagnostics

## 3. How We Use Your Information

We use your information to:
- Provide and maintain the service
- Process and store your documents
- Send important notifications
- Improve our service
- Provide customer support

## 4. Data Storage

- Your documents are stored securely on our servers
- Premium users can enable cloud backup
- We use industry-standard encryption to protect your data

## 5. Data Sharing

We do not sell your personal information. We may share data with:
- Service providers who help us operate the app
- Law enforcement when required by law
- Other parties with your consent

## 6. Your Rights

You have the right to:
- Access your personal data
- Delete your account and data
- Export your documents
- Opt out of marketing communications

## 7. Data Security

We implement appropriate security measures to protect your information, including:
- Encryption in transit and at rest
- Secure authentication
- Regular security audits

## 8. Children's Privacy

Our service is not intended for children under 13. We do not knowingly collect information from children.

## 9. Changes to This Policy

We may update this Privacy Policy from time to time. We will notify you of any changes by posting the new policy on this page.

## 10. Contact Us

If you have questions about this Privacy Policy, please contact us at:
- Email: privacy@scanup.app
""",
        "de": """# Datenschutzrichtlinie

**Zuletzt aktualisiert: Januar 2025**

## 1. Einführung

ScanUp ("wir" oder "uns") verpflichtet sich zum Schutz Ihrer Privatsphäre. Diese Datenschutzrichtlinie erklärt, wie wir Informationen über Sie sammeln, verwenden und weitergeben.

## 2. Informationen, die wir sammeln

### Von Ihnen bereitgestellte Informationen
- Kontoinformationen (E-Mail, Name)
- Dokumente, die Sie scannen und hochladen
- Von Ihnen erstellte Unterschriften

### Automatisch gesammelte Informationen
- Geräteinformationen (Modell, Betriebssystemversion)
- Nutzungsdaten (verwendete Funktionen, Scan-Anzahl)
- Absturzberichte und Diagnosen

## 3. Wie wir Ihre Informationen verwenden

Wir verwenden Ihre Informationen, um:
- Den Dienst bereitzustellen und zu warten
- Ihre Dokumente zu verarbeiten und zu speichern
- Wichtige Benachrichtigungen zu senden
- Unseren Service zu verbessern
- Kundensupport zu bieten

## 4. Datenspeicherung

- Ihre Dokumente werden sicher auf unseren Servern gespeichert
- Premium-Benutzer können Cloud-Backup aktivieren
- Wir verwenden branchenübliche Verschlüsselung zum Schutz Ihrer Daten

## 5. Datenweitergabe

Wir verkaufen Ihre persönlichen Daten nicht. Wir können Daten teilen mit:
- Dienstleistern, die uns beim Betrieb der App helfen
- Strafverfolgungsbehörden, wenn gesetzlich vorgeschrieben
- Anderen Parteien mit Ihrer Zustimmung

## 6. Ihre Rechte

Sie haben das Recht:
- Auf Ihre persönlichen Daten zuzugreifen
- Ihr Konto und Ihre Daten zu löschen
- Ihre Dokumente zu exportieren
- Marketing-Kommunikation abzulehnen

## 7. Datensicherheit

Wir implementieren angemessene Sicherheitsmaßnahmen zum Schutz Ihrer Informationen, einschließlich:
- Verschlüsselung bei der Übertragung und Speicherung
- Sichere Authentifizierung
- Regelmäßige Sicherheitsaudits

## 8. Datenschutz für Kinder

Unser Service ist nicht für Kinder unter 13 Jahren bestimmt.

## 9. Änderungen dieser Richtlinie

Wir können diese Datenschutzrichtlinie von Zeit zu Zeit aktualisieren.

## 10. Kontakt

Bei Fragen zu dieser Datenschutzrichtlinie kontaktieren Sie uns bitte unter:
- E-Mail: privacy@scanup.app
"""
    },
    "support": {
        "en": """# Help & Support

## Frequently Asked Questions

### How do I scan a document?
1. Tap the blue scan button at the bottom of the screen
2. Point your camera at the document
3. The app will automatically detect the edges
4. Tap the capture button to take the photo
5. Adjust the crop if needed and tap "Use Photo"

### How do I add multiple pages?
After scanning your first page, tap "Add Page" to scan additional pages. All pages will be saved as a single document.

### How do I edit a scanned document?
1. Open the document from your document list
2. Tap "Edit" to access editing tools
3. You can apply filters, adjust brightness/contrast, rotate, crop, and add signatures
4. Tap "Save" when finished

### How do I add a signature?
1. Open a document and tap "Edit"
2. Select "Sign" from the toolbar
3. Choose an existing signature or draw a new one
4. Position and resize the signature on the document
5. Tap "Save" to apply

### How do I share a document?
Tap the share icon on any document to export it as a PDF. You can then share via email, messaging apps, or save to cloud services.

### How do I protect a document with a password?
1. Tap the "..." menu on a document
2. Select "Password"
3. Enter your desired password
4. The document will now require the password to open

## Contact Support

If you need additional help, please contact us:

**Email:** support@scanup.app

**Response Time:** We typically respond within 24-48 hours

## Report a Bug

Found a bug? Please email us at bugs@scanup.app with:
- Your device model and OS version
- Steps to reproduce the issue
- Screenshots if possible
""",
        "de": """# Hilfe & Support

## Häufig gestellte Fragen

### Wie scanne ich ein Dokument?
1. Tippen Sie auf die blaue Scan-Taste am unteren Bildschirmrand
2. Richten Sie Ihre Kamera auf das Dokument
3. Die App erkennt automatisch die Kanten
4. Tippen Sie auf die Aufnahmetaste
5. Passen Sie den Zuschnitt bei Bedarf an und tippen Sie auf "Foto verwenden"

### Wie füge ich mehrere Seiten hinzu?
Nach dem Scannen der ersten Seite tippen Sie auf "Seite hinzufügen", um weitere Seiten zu scannen.

### Wie bearbeite ich ein gescanntes Dokument?
1. Öffnen Sie das Dokument aus Ihrer Dokumentenliste
2. Tippen Sie auf "Bearbeiten"
3. Sie können Filter anwenden, Helligkeit/Kontrast anpassen, drehen, zuschneiden und Unterschriften hinzufügen
4. Tippen Sie auf "Speichern", wenn Sie fertig sind

### Wie füge ich eine Unterschrift hinzu?
1. Öffnen Sie ein Dokument und tippen Sie auf "Bearbeiten"
2. Wählen Sie "Unterschreiben" aus der Werkzeugleiste
3. Wählen Sie eine bestehende Unterschrift oder zeichnen Sie eine neue
4. Positionieren und skalieren Sie die Unterschrift
5. Tippen Sie auf "Speichern"

### Wie teile ich ein Dokument?
Tippen Sie auf das Teilen-Symbol, um es als PDF zu exportieren.

### Wie schütze ich ein Dokument mit einem Passwort?
1. Tippen Sie auf das "..."-Menü eines Dokuments
2. Wählen Sie "Passwort"
3. Geben Sie Ihr gewünschtes Passwort ein

## Support kontaktieren

**E-Mail:** support@scanup.app

**Antwortzeit:** Wir antworten in der Regel innerhalb von 24-48 Stunden

## Fehler melden

Fehler gefunden? Senden Sie uns eine E-Mail an bugs@scanup.app
"""
    }
}


# Pydantic models for content management
class LanguageModel(BaseModel):
    code: str
    name: str
    native_name: str
    is_default: bool = False

class TranslationUpdate(BaseModel):
    translations: Dict[str, str]

class LegalPageUpdate(BaseModel):
    content: str

class TranslationKeyUpdate(BaseModel):
    key: str
    value: str


# Initialize default content on startup
async def init_content_collections():
    """Initialize languages, translations, and legal pages with defaults"""
    try:
        # Initialize languages
        languages_count = await db.languages.count_documents({})
        if languages_count == 0:
            for lang in DEFAULT_LANGUAGES:
                await db.languages.update_one(
                    {"code": lang["code"]},
                    {"$set": lang},
                    upsert=True
                )
            logger.info("✅ Default languages initialized")
        
        # Initialize English translations
        en_trans = await db.translations.find_one({"language_code": "en"})
        if not en_trans:
            await db.translations.insert_one({
                "language_code": "en",
                "translations": DEFAULT_TRANSLATIONS,
                "updated_at": datetime.now(timezone.utc)
            })
            logger.info("✅ English translations initialized")
        
        # Initialize German translations
        de_trans = await db.translations.find_one({"language_code": "de"})
        if not de_trans:
            await db.translations.insert_one({
                "language_code": "de",
                "translations": DEFAULT_TRANSLATIONS_DE,
                "updated_at": datetime.now(timezone.utc)
            })
            logger.info("✅ German translations initialized")
        
        # Initialize legal pages
        for page_type, content in DEFAULT_LEGAL_PAGES.items():
            for lang_code, page_content in content.items():
                existing = await db.legal_pages.find_one({
                    "page_type": page_type,
                    "language_code": lang_code
                })
                if not existing:
                    await db.legal_pages.insert_one({
                        "page_type": page_type,
                        "language_code": lang_code,
                        "content": page_content,
                        "updated_at": datetime.now(timezone.utc)
                    })
            logger.info(f"✅ Legal page '{page_type}' initialized")
        
    except Exception as e:
        logger.error(f"Error initializing content: {e}")


# Public API endpoints for content

@api_router.get("/content/languages")
async def get_languages():
    """Get all supported languages"""
    try:
        languages = await db.languages.find({}, {"_id": 0}).to_list(100)
        if not languages:
            return DEFAULT_LANGUAGES
        return languages
    except Exception as e:
        logger.error(f"Error fetching languages: {e}")
        return DEFAULT_LANGUAGES


@api_router.get("/content/translations/{language_code}")
async def get_translations(language_code: str):
    """Get all translations for a language"""
    try:
        trans = await db.translations.find_one(
            {"language_code": language_code},
            {"_id": 0}
        )
        
        if trans:
            return {
                "language_code": language_code,
                "translations": trans.get("translations", {}),
                "updated_at": trans.get("updated_at")
            }
        
        # Fallback to English if language not found
        en_trans = await db.translations.find_one(
            {"language_code": "en"},
            {"_id": 0}
        )
        
        if en_trans:
            # Return English translations as fallback with the requested language code
            return {
                "language_code": language_code,
                "translations": en_trans.get("translations", DEFAULT_TRANSLATIONS),
                "updated_at": en_trans.get("updated_at"),
                "is_fallback": True
            }
        
        # Ultimate fallback to hardcoded defaults
        return {
            "language_code": language_code,
            "translations": DEFAULT_TRANSLATIONS,
            "is_fallback": True
        }
        
    except Exception as e:
        logger.error(f"Error fetching translations: {e}")
        return {
            "language_code": language_code,
            "translations": DEFAULT_TRANSLATIONS,
            "is_fallback": True
        }


@api_router.get("/content/legal/{page_type}")
async def get_legal_page(page_type: str, language_code: str = "en"):
    """Get a legal page (terms, privacy, support)"""
    if page_type not in ["terms", "privacy", "support"]:
        raise HTTPException(status_code=400, detail="Invalid page type")
    
    try:
        page = await db.legal_pages.find_one(
            {"page_type": page_type, "language_code": language_code},
            {"_id": 0}
        )
        
        if page:
            return {
                "page_type": page_type,
                "language_code": language_code,
                "content": page.get("content", ""),
                "updated_at": page.get("updated_at")
            }
        
        # Fallback to English
        en_page = await db.legal_pages.find_one(
            {"page_type": page_type, "language_code": "en"},
            {"_id": 0}
        )
        
        if en_page:
            return {
                "page_type": page_type,
                "language_code": language_code,
                "content": en_page.get("content", ""),
                "updated_at": en_page.get("updated_at"),
                "is_fallback": True
            }
        
        # Ultimate fallback
        default_content = DEFAULT_LEGAL_PAGES.get(page_type, {}).get("en", "")
        return {
            "page_type": page_type,
            "language_code": language_code,
            "content": default_content,
            "is_fallback": True
        }
        
    except Exception as e:
        logger.error(f"Error fetching legal page: {e}")
        default_content = DEFAULT_LEGAL_PAGES.get(page_type, {}).get("en", "")
        return {
            "page_type": page_type,
            "language_code": language_code,
            "content": default_content,
            "is_fallback": True
        }


# ==================== ADMIN DASHBOARD API ====================

ADMIN_EMAIL = os.environ.get("ADMIN_EMAIL", "admin@scanup.com")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "admin123")

class AdminLoginRequest(BaseModel):
    email: str
    password: str

@api_router.post("/admin/login")
async def admin_login(request: AdminLoginRequest):
    """Admin login endpoint"""
    if request.email == ADMIN_EMAIL and request.password == ADMIN_PASSWORD:
        # Create admin token
        token_data = {
            "sub": "admin",
            "email": request.email,
            "is_admin": True,
            "exp": datetime.now(timezone.utc) + timedelta(days=7)
        }
        token = jwt.encode(token_data, JWT_SECRET, algorithm="HS256")
        return {
            "token": token,
            "user": {"email": request.email, "is_admin": True}
        }
    raise HTTPException(status_code=401, detail="Invalid credentials")

async def get_admin_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify admin token"""
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
        if not payload.get("is_admin"):
            raise HTTPException(status_code=403, detail="Admin access required")
        return payload
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid admin token")

# ==================== PUSH NOTIFICATIONS API ====================

class PushTokenRequest(BaseModel):
    push_token: str
    device_type: str = "unknown"  # ios, android
    device_id: Optional[str] = None

class SendNotificationRequest(BaseModel):
    user_ids: Optional[List[str]] = None  # If None, send to all
    title: str
    body: str
    data: Optional[Dict[str, Any]] = None

@api_router.post("/notifications/register-token")
async def register_push_token(request: PushTokenRequest, user: dict = Depends(get_current_user)):
    """Register push token for a user"""
    try:
        # Update user with push token
        await db.users.update_one(
            {"user_id": user["user_id"]},
            {
                "$set": {
                    "push_token": request.push_token,
                    "device_type": request.device_type,
                    "device_id": request.device_id,
                    "push_token_updated_at": datetime.now(timezone.utc).isoformat()
                }
            }
        )
        
        logger.info(f"Push token registered for user {user['user_id']}")
        return {"success": True, "message": "Push token registered"}
    except Exception as e:
        logger.error(f"Failed to register push token: {e}")
        raise HTTPException(status_code=500, detail="Failed to register push token")

@api_router.delete("/notifications/unregister-token")
async def unregister_push_token(user: dict = Depends(get_current_user)):
    """Remove push token for a user (on logout)"""
    try:
        await db.users.update_one(
            {"user_id": user["user_id"]},
            {
                "$unset": {
                    "push_token": "",
                    "device_type": "",
                    "device_id": ""
                }
            }
        )
        
        logger.info(f"Push token removed for user {user['user_id']}")
        return {"success": True, "message": "Push token removed"}
    except Exception as e:
        logger.error(f"Failed to unregister push token: {e}")
        raise HTTPException(status_code=500, detail="Failed to unregister push token")

@api_router.post("/admin/send-notification")
async def send_push_notification(request: SendNotificationRequest, admin: dict = Depends(get_admin_user)):
    """Send push notification to users (admin only)"""
    try:
        # Build query
        query = {"push_token": {"$exists": True, "$ne": None}}
        if request.user_ids:
            query["user_id"] = {"$in": request.user_ids}
        
        # Get users with push tokens
        users = await db.users.find(query, {"push_token": 1, "user_id": 1}).to_list(1000)
        
        if not users:
            return {"success": True, "sent_count": 0, "message": "No users with push tokens found"}
        
        # Prepare push tokens
        push_tokens = [u["push_token"] for u in users if u.get("push_token")]
        
        # Send via Expo Push Service
        sent_count = 0
        failed_count = 0
        
        # Expo Push API endpoint
        expo_push_url = "https://exp.host/--/api/v2/push/send"
        
        # Send in batches of 100
        async with httpx.AsyncClient() as client:
            for i in range(0, len(push_tokens), 100):
                batch = push_tokens[i:i+100]
                messages = [
                    {
                        "to": token,
                        "title": request.title,
                        "body": request.body,
                        "data": request.data or {},
                        "sound": "default",
                        "priority": "high"
                    }
                    for token in batch
                ]
                
                try:
                    response = await client.post(
                        expo_push_url,
                        json=messages,
                        headers={"Content-Type": "application/json"}
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        for ticket in result.get("data", []):
                            if ticket.get("status") == "ok":
                                sent_count += 1
                            else:
                                failed_count += 1
                    else:
                        failed_count += len(batch)
                        logger.error(f"Expo push failed: {response.text}")
                except Exception as e:
                    failed_count += len(batch)
                    logger.error(f"Expo push error: {e}")
        
        return {
            "success": True,
            "sent_count": sent_count,
            "failed_count": failed_count,
            "total_tokens": len(push_tokens)
        }
    except Exception as e:
        logger.error(f"Failed to send notifications: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to send notifications: {str(e)}")

@api_router.get("/admin/me")
async def admin_me(admin: dict = Depends(get_admin_user)):
    """Get current admin user"""
    return {"email": admin.get("email"), "is_admin": True}

@api_router.get("/admin/stats")
async def get_admin_stats(admin: dict = Depends(get_admin_user)):
    """Get dashboard statistics"""
    try:
        # Get counts
        total_users = await db.users.count_documents({})
        total_documents = await db.documents.count_documents({})
        
        # Get users from last 7 days
        week_ago = datetime.now(timezone.utc) - timedelta(days=7)
        new_users_week = await db.users.count_documents({"created_at": {"$gte": week_ago.isoformat()}})
        
        # Calculate storage (approximate)
        pipeline = [
            {"$unwind": "$pages"},
            {"$group": {"_id": None, "total_pages": {"$sum": 1}}}
        ]
        page_count_result = await db.documents.aggregate(pipeline).to_list(1)
        total_pages = page_count_result[0]["total_pages"] if page_count_result else 0
        storage_mb = total_pages * 0.5  # Estimate ~500KB per page
        
        # User growth data (last 30 days)
        user_growth = []
        for i in range(30, -1, -1):
            day = datetime.now(timezone.utc) - timedelta(days=i)
            day_start = day.replace(hour=0, minute=0, second=0, microsecond=0)
            count = await db.users.count_documents({"created_at": {"$lte": day_start.isoformat()}})
            user_growth.append({
                "date": day.strftime("%m/%d"),
                "users": count
            })
        
        # Daily scans (last 14 days)
        daily_scans = []
        for i in range(14, -1, -1):
            day = datetime.now(timezone.utc) - timedelta(days=i)
            day_start = day.replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = day_start + timedelta(days=1)
            count = await db.documents.count_documents({
                "created_at": {"$gte": day_start.isoformat(), "$lt": day_end.isoformat()}
            })
            daily_scans.append({
                "date": day.strftime("%m/%d"),
                "scans": count
            })
        
        return {
            "total_users": total_users,
            "total_documents": total_documents,
            "storage_used": f"{storage_mb:.1f} MB",
            "new_users_week": new_users_week,
            "users_change": f"+{new_users_week}" if new_users_week > 0 else "0",
            "docs_change": "+0%",
            "storage_change": "+0%",
            "new_users_change": "+0%",
            "user_growth": user_growth,
            "daily_scans": daily_scans,
        }
    except Exception as e:
        logger.error(f"Admin stats error: {e}")
        return {
            "total_users": 0,
            "total_documents": 0,
            "storage_used": "0 MB",
            "new_users_week": 0,
        }

@api_router.get("/admin/users")
async def get_admin_users(
    admin: dict = Depends(get_admin_user),
    search: str = None,
    filter: str = "all",
    limit: int = 50
):
    """Get all users for admin"""
    try:
        query = {}
        
        if search:
            query["$or"] = [
                {"email": {"$regex": search, "$options": "i"}},
                {"name": {"$regex": search, "$options": "i"}}
            ]
        
        if filter == "premium":
            query["subscription_type"] = "premium"
        elif filter == "trial":
            query["subscription_type"] = "trial"
        elif filter == "ad-free":
            query["has_removed_ads"] = True
        elif filter == "free":
            query["$and"] = [
                {"subscription_type": {"$nin": ["premium", "trial"]}},
                {"has_removed_ads": {"$ne": True}}
            ]
        elif filter == "google":
            query["auth_provider"] = "google"
        elif filter == "recent":
            week_ago = datetime.now(timezone.utc) - timedelta(days=7)
            query["created_at"] = {"$gte": week_ago.isoformat()}
        
        users = await db.users.find(query, {"_id": 0, "password_hash": 0}).sort("created_at", -1).limit(limit).to_list(limit)
        
        # Add document count and scan count for each user
        for user in users:
            user["document_count"] = await db.documents.count_documents({"user_id": user.get("user_id")})
            # Check if is_premium based on subscription_type for backward compatibility
            if not user.get("is_premium"):
                user["is_premium"] = user.get("subscription_type") == "premium"
        
        return {"users": users}
    except Exception as e:
        logger.error(f"Admin users error: {e}")
        return {"users": []}


@api_router.get("/admin/users/{user_id}")
async def get_admin_user_detail(user_id: str, admin: dict = Depends(get_admin_user)):
    """Get detailed user info for admin"""
    try:
        user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
        
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Add document and usage counts
        user["document_count"] = await db.documents.count_documents({"user_id": user_id})
        
        # Check if user has password set (for OAuth users)
        full_user = await db.users.find_one({"user_id": user_id})
        user["has_password"] = bool(full_user.get("password_hash"))
        
        # Backward compatibility for is_premium
        if not user.get("is_premium"):
            user["is_premium"] = user.get("subscription_type") == "premium"
        
        return user
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Admin user detail error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.delete("/admin/users/{user_id}")
async def delete_admin_user(user_id: str, admin: dict = Depends(get_admin_user)):
    """Delete a user and all their data"""
    try:
        # Delete user's documents from S3
        if s3_client:
            try:
                delete_from_s3(user_id, "all")
            except Exception as s3_err:
                logger.warning(f"S3 cleanup failed: {s3_err}")
        
        # Delete documents
        await db.documents.delete_many({"user_id": user_id})
        
        # Delete folders
        await db.folders.delete_many({"user_id": user_id})
        
        # Delete signatures
        await db.signatures.delete_many({"user_id": user_id})
        
        # Delete user
        result = await db.users.delete_one({"user_id": user_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="User not found")
        
        return {"message": "User deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete user error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class AdminUserUpdate(BaseModel):
    """Schema for admin user update"""
    subscription_type: str = None  # "free", "premium", "trial"
    is_premium: bool = None
    has_removed_ads: bool = None
    has_removed_watermark: bool = None
    subscription_expires_at: str = None  # ISO date string or None
    notes: str = None  # Admin notes


@api_router.put("/admin/users/{user_id}")
async def update_admin_user(user_id: str, update: AdminUserUpdate, admin: dict = Depends(get_admin_user)):
    """Update user subscription status (admin only)"""
    try:
        # Find user
        user = await db.users.find_one({"user_id": user_id})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Build update dict
        update_data = {"updated_at": datetime.now(timezone.utc)}
        
        if update.subscription_type is not None:
            update_data["subscription_type"] = update.subscription_type
            # Auto-set is_premium based on subscription_type
            update_data["is_premium"] = update.subscription_type == "premium"
            update_data["is_trial"] = update.subscription_type == "trial"
        
        if update.is_premium is not None:
            update_data["is_premium"] = update.is_premium
            if update.is_premium:
                update_data["subscription_type"] = "premium"
            elif update.subscription_type is None:
                update_data["subscription_type"] = "free"
        
        if update.has_removed_ads is not None:
            update_data["has_removed_ads"] = update.has_removed_ads
        
        if update.has_removed_watermark is not None:
            update_data["has_removed_watermark"] = update.has_removed_watermark
        
        if update.subscription_expires_at is not None:
            if update.subscription_expires_at == "":
                update_data["subscription_expires_at"] = None
            else:
                update_data["subscription_expires_at"] = update.subscription_expires_at
        
        if update.notes is not None:
            update_data["admin_notes"] = update.notes
        
        # Update user
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": update_data}
        )
        
        # Fetch updated user
        updated_user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
        updated_user["document_count"] = await db.documents.count_documents({"user_id": user_id})
        
        logger.info(f"Admin updated user {user_id}: {update_data}")
        
        return updated_user
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Admin user update error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.get("/admin/documents")
async def get_admin_documents(
    admin: dict = Depends(get_admin_user),
    search: str = None,
    limit: int = 50
):
    """Get all documents for admin"""
    try:
        query = {}
        
        if search:
            query["name"] = {"$regex": search, "$options": "i"}
        
        documents = await db.documents.find(query, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)
        
        # Add user email and page count
        for doc in documents:
            user = await db.users.find_one({"user_id": doc.get("user_id")}, {"email": 1})
            doc["user_email"] = user.get("email") if user else "Unknown"
            doc["page_count"] = len(doc.get("pages", []))
            # Get thumbnail from first page
            if doc.get("pages") and len(doc["pages"]) > 0:
                first_page = doc["pages"][0]
                doc["thumbnail_url"] = first_page.get("thumbnail_url") or first_page.get("image_url")
            # Remove full page data to reduce response size
            doc.pop("pages", None)
        
        return {"documents": documents}
    except Exception as e:
        logger.error(f"Admin documents error: {e}")
        return {"documents": []}

@api_router.delete("/admin/documents/{document_id}")
async def delete_admin_document(document_id: str, admin: dict = Depends(get_admin_user)):
    """Delete a document"""
    try:
        doc = await db.documents.find_one({"document_id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete from S3
        if s3_client:
            try:
                delete_from_s3(doc.get("user_id"), document_id)
            except Exception as s3_err:
                logger.warning(f"S3 cleanup failed: {s3_err}")
        
        # Delete from DB
        await db.documents.delete_one({"document_id": document_id})
        
        return {"message": "Document deleted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/admin/analytics")
async def get_admin_analytics(
    admin: dict = Depends(get_admin_user),
    range: str = "7d"
):
    """Get analytics data"""
    try:
        days = {"7d": 7, "30d": 30, "90d": 90, "all": 365}.get(range, 7)
        start_date = datetime.now(timezone.utc) - timedelta(days=days)
        
        # Get active users
        active_users = await db.documents.distinct("user_id", {
            "created_at": {"$gte": start_date.isoformat()}
        })
        
        # Get total scans in period
        total_scans = await db.documents.count_documents({
            "created_at": {"$gte": start_date.isoformat()}
        })
        
        # Document types distribution
        doc_types = await db.documents.aggregate([
            {"$match": {"created_at": {"$gte": start_date.isoformat()}}},
            {"$group": {"_id": "$document_type", "count": {"$sum": 1}}}
        ]).to_list(10)
        
        document_types = [
            {"name": dt.get("_id") or "document", "value": dt.get("count", 0)}
            for dt in doc_types
        ]
        
        # Scans over time
        scans_over_time = []
        for i in range(min(days, 30), -1, -1):
            day = datetime.now(timezone.utc) - timedelta(days=i)
            day_start = day.replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = day_start + timedelta(days=1)
            
            scan_count = await db.documents.count_documents({
                "created_at": {"$gte": day_start.isoformat(), "$lt": day_end.isoformat()}
            })
            user_count = len(await db.documents.distinct("user_id", {
                "created_at": {"$gte": day_start.isoformat(), "$lt": day_end.isoformat()}
            }))
            
            scans_over_time.append({
                "date": day.strftime("%m/%d"),
                "scans": scan_count,
                "users": user_count
            })
        
        # Platform distribution (mock data for now)
        platforms = [
            {"name": "iOS", "users": int(len(active_users) * 0.6)},
            {"name": "Android", "users": int(len(active_users) * 0.35)},
            {"name": "Web", "users": int(len(active_users) * 0.05)},
        ]
        
        return {
            "total_scans": total_scans,
            "active_users": len(active_users),
            "avg_session": "4m 32s",
            "retention_rate": "68%",
            "scans_over_time": scans_over_time,
            "document_types": document_types if document_types else [{"name": "Document", "value": 1}],
            "platforms": platforms,
        }
    except Exception as e:
        logger.error(f"Analytics error: {e}")
        return {
            "total_scans": 0,
            "active_users": 0,
            "avg_session": "0m",
            "retention_rate": "0%",
        }

# ============== LOCALIZATION/TRANSLATIONS ==============

# Default English translations - all app strings
DEFAULT_TRANSLATIONS = {
    "en": {
        # General
        "app_name": "ScanUp",
        "loading": "Loading...",
        "please_wait": "Please wait",
        "cancel": "Cancel",
        "save": "Save",
        "done": "Done",
        "apply": "Apply",
        "delete": "Delete",
        "edit": "Edit",
        "create": "Create",
        "search": "Search",
        "close": "Close",
        "back": "Back",
        "next": "Next",
        "skip": "Skip",
        "retry": "Retry",
        "yes": "Yes",
        "no": "No",
        "ok": "OK",
        "error": "Error",
        "success": "Success",
        "warning": "Warning",
        
        # Auth
        "sign_in": "Sign In",
        "sign_up": "Sign Up",
        "sign_out": "Sign Out",
        "logout": "Logout",
        "create_account": "Create Account",
        "welcome_back": "Welcome Back",
        "sign_in_to_continue": "Sign in to continue",
        "sign_up_to_get_started": "Sign up to get started",
        "already_have_account": "Already have an account?",
        "dont_have_account": "Don't have an account?",
        "continue_with_google": "Continue with Google",
        "email": "Email",
        "password": "Password",
        "full_name": "Full Name",
        "confirm_password": "Confirm Password",
        "enter_email": "Enter your email",
        "enter_password": "Enter your password",
        "enter_name": "Enter your name",
        "create_password": "Create a password",
        "confirm_your_password": "Confirm your password",
        "forgot_password": "Forgot Password?",
        "reset_password": "Reset Password",
        "set_password": "Set Password",
        "sign_in_to_sync": "Sign In to Sync",
        "guest_mode": "Guest Mode",
        "continue_as_guest": "Continue as Guest",
        
        # Onboarding
        "get_started": "Get Started",
        "smart_scanning": "Smart Scanning",
        "auto_detect_boundaries": "Auto-detect document boundaries",
        "smart_enhancement": "Smart Enhancement",
        "auto_correct_beautify": "Auto-correct & beautify scans",
        "instant_ocr": "Instant OCR",
        "extract_text_instantly": "Extract text instantly",
        "cloud_sync": "Cloud Sync",
        "access_from_anywhere": "Access from anywhere",
        
        # Home / Documents
        "documents": "Documents",
        "my_documents": "My Documents",
        "no_documents": "No Documents",
        "no_documents_yet": "No Documents Yet",
        "add_documents": "Add Documents",
        "search_documents": "Search documents, tags, or text...",
        "loading_documents": "Loading documents...",
        "loading_document": "Loading document...",
        "document_name": "Document name",
        "rename_document": "Rename Document",
        "delete_document": "Delete Document",
        "export_document": "Export Document",
        "share_document": "Share Document",
        "view_mode": "View Mode",
        
        # Folders
        "folders": "Folders",
        "no_folders_yet": "No Folders Yet",
        "new_folder": "New Folder",
        "folder_name": "Folder name",
        "delete_folder": "Delete Folder",
        "move_to_folder": "Move to Folder",
        "remove_from_folder": "Remove from folder",
        "folder_protected": "Folder Protected",
        "set_folder_password": "Set Folder Password",
        "remove_password": "Remove Password",
        "enter_folder_password": "Enter folder password",
        "unlock": "Unlock",
        
        # Scanner
        "scan": "Scan",
        "camera": "Camera",
        "gallery": "Gallery",
        "auto_detect": "Auto Detect",
        "capturing": "Capturing...",
        "camera_permission_required": "Camera Permission Required",
        "grant_permission": "Grant Permission",
        "live_preview": "Live Preview",
        "live_edge_detection": "Live Edge Detection",
        "show_grid_overlay": "Show Grid Overlay",
        "display_grid_on_camera": "Display alignment grid on camera",
        "flash": "Flash",
        "flash_on": "Flash On",
        "flash_off": "Flash Off",
        "flash_auto": "Flash Auto",
        
        # Document Modes
        "single_page": "Single Page",
        "multi_page": "Multi Page",
        "book_mode": "Book Mode",
        "id_card": "ID Card",
        "batch_scan": "Batch Scan",
        "left_page": "Left Page",
        "right_page": "Right Page",
        
        # Document Types/Templates
        "select_document_type": "Select Document Type",
        "select_document_template": "Select Document Template",
        "template": "Template",
        "general_document": "General Document",
        "receipt": "Receipt",
        "business_card": "Business Card",
        "whiteboard": "Whiteboard",
        "passport": "Passport",
        "form": "Form",
        
        # Edit / Crop
        "edit": "Edit",
        "crop": "Crop",
        "adjust_crop": "Adjust Crop",
        "auto_crop": "Auto Crop",
        "rotate": "Rotate",
        "adjust_filter": "Adjust & Filter",
        "revert": "Revert",
        "revert_to_original": "Revert to Original",
        "reset": "Reset",
        
        # Filters
        "filters": "Filters",
        "original": "Original",
        "enhanced": "Enhanced",
        "grayscale": "Grayscale",
        "black_white": "Black & White",
        "color": "Color",
        "magic": "Magic",
        "default_filter": "Default Filter",
        "select_default_filter": "Select Default Filter",
        
        # Adjustments
        "brightness": "Brightness",
        "contrast": "Contrast",
        "saturation": "Saturation",
        "auto_enhance": "Auto Enhance",
        "automatically_enhance": "Automatically enhance scanned images",
        
        # Pages
        "pages": "Pages",
        "page": "Page",
        "add_page": "Add Page",
        "add_more": "Add More",
        "del_page": "Del Page",
        "delete_page": "Delete Page",
        "move_up": "Move Up",
        "move_down": "Move Down",
        "add_to_document": "Add to Document",
        "edit_page_screen": "Edit Page Screen",
        
        # Annotations
        "annotate": "Annotate",
        "annotations": "Annotations",
        "draw": "Draw",
        "text": "Text",
        "shapes": "Shapes",
        "arrow": "Arrow",
        "rectangle": "Rectangle",
        "circle": "Circle",
        "line": "Line",
        "highlight": "Highlight",
        "thickness": "Thickness",
        "enter_text": "Enter text...",
        "clear": "Clear",
        "undo": "Undo",
        
        # Signature
        "sign": "Sign",
        "signature": "Signature",
        "draw_signature": "Draw Signature",
        "sign_here": "Sign here",
        "position_signature": "Position Signature",
        "signature_hint": "Hold & drag signature to move • Pinch to resize",
        "re_sign": "Re-Sign",
        "saved_signatures": "Saved Signatures",
        
        # OCR
        "ocr": "OCR",
        "extract_text": "Extract Text",
        "extracted_text": "Extracted Text",
        "copy_text": "Copy Text",
        "text_copied": "Text copied to clipboard",
        "no_text_found": "No text found",
        "ocr_processing": "Processing OCR...",
        "ocr_left": "OCR Left",
        
        # Export
        "export": "Export",
        "share": "Share",
        "export_as_pdf": "Export as PDF",
        "export_as_image": "Export as Image",
        "export_as_jpeg": "Export as JPEG",
        "export_as_png": "Export as PNG",
        "pdf": "PDF",
        "jpeg": "JPEG",
        "png": "PNG",
        "high_quality": "High Quality",
        "medium_quality": "Medium Quality",
        "low_quality": "Low Quality",
        
        # Settings
        "settings": "Settings",
        "all_settings": "All Settings",
        "preferences": "Preferences",
        "general": "General",
        "appearance": "Appearance",
        "dark_mode": "Dark Mode",
        "language": "Language",
        "select_language": "Select Language",
        "default_scan_quality": "Default Scan Quality",
        "select_scan_quality": "Select Scan Quality",
        "sound_effects": "Sound Effects",
        "haptic_feedback": "Haptic Feedback",
        "auto_backup": "Auto Backup",
        "clear_cache": "Clear Cache",
        "free_up_storage": "Free up storage space",
        "reset_settings": "Reset Settings",
        "restore_defaults": "Restore default settings",
        "advanced_features_coming": "Advanced editing features coming soon",
        
        # Profile
        "profile": "Profile",
        "account": "Account",
        "plan": "Plan",
        "pro": "Pro",
        "free": "Free",
        "trial": "Trial",
        "free_plan": "Free Plan",
        "premium_member": "Premium Member",
        "try_premium_free": "Try Premium Free",
        "seven_day_trial": "7-day trial, no credit card needed",
        "start": "Start",
        "start_trial": "Start Trial",
        "scans_left": "Scans Left",
        "monthly_scans_remaining": "Monthly scans remaining:",
        "daily_usage": "Daily Usage",
        "daily": "daily",
        "help_support": "Help & Support",
        "privacy_policy": "Privacy Policy",
        "terms_of_service": "Terms of Service",
        "version": "Version",
        "go_back": "Go Back",
        "guest_user": "Guest User",
        "not_signed_in": "Not signed in",
        "sign_in_to_sync": "Sign In to Sync",
        "manage_subscription": "Manage Subscription",
        "upgrade_to_pro": "Upgrade to Pro",
        "all_features_unlocked": "All features unlocked",
        "full_access_during_trial": "Full access during trial",
        "limited_features": "Limited features",
        "days_left": "days left",
        "exit_guest_mode": "Exit Guest Mode",
        "exit_guest_confirm": "Are you sure you want to exit?",
        "logout_confirm": "Are you sure you want to logout?",
        "exit": "Exit",
        "sign_in_required": "Sign In Required",
        "sign_in_to_start_trial": "Please sign in to start your free trial.",
        "sign_in_to_upgrade": "Please sign in to upgrade to premium.",
        "start_7_day_trial": "Start 7-Day Free Trial",
        "unlock_premium_features": "Unlock all premium features for 7 days:",
        "unlimited_scans": "Unlimited scans",
        "unlimited_ocr": "Unlimited OCR",
        "no_watermarks": "No watermarks",
        "all_premium_features": "All premium features",
        "trial_started": "Trial Started!",
        "enjoy_premium": "Enjoy 7 days of premium features!",
        "cancel_premium": "Cancel Premium",
        "cancel_premium_confirm": "Are you sure you want to cancel your premium subscription?",
        "yes_cancel": "Yes, Cancel",
        "subscription_cancelled": "Your premium subscription has been cancelled.",
        "premium_activated": "Premium Activated!",
        "enjoy_premium_features": "Enjoy all premium features!",
        "auto_backup": "Auto Backup",
        "all_settings": "All Settings",
        "signatures": "Signatures",
        "create_manage_signatures": "Create and manage your saved signatures",
        "manage_signatures": "Manage Signatures",
        "create_edit_delete_signatures": "Create, edit, and delete signatures",
        "no_signatures_yet": "No Signatures Yet",
        "create_first_signature": "Create your first signature to use when signing documents",
        "delete_signature": "Delete Signature",
        "delete_signature_confirm": "Are you sure you want to delete this signature?",
        "delete_selected": "Delete Selected",
        "delete_confirm_multiple": "Are you sure you want to delete",
        "signature_s": "signature(s)",
        "new_signature": "New Signature",
        "create_new_signature": "Create New Signature",
        "long_press_to_delete": "Long press to delete",
        "upgrade_now": "Upgrade Now",
        "continue_with_apple": "Continue with Apple",
        
        # Errors & Messages
        "something_went_wrong": "Something went wrong",
        "network_error": "Network error. Please try again.",
        "invalid_credentials": "Invalid email or password",
        "email_already_exists": "Email already exists",
        "password_too_short": "Password must be at least 6 characters",
        "passwords_dont_match": "Passwords don't match",
        "document_saved": "Document saved successfully",
        "document_deleted": "Document deleted",
        "changes_saved": "Changes saved",
        "cache_cleared": "Cache cleared",
        "settings_reset": "Settings reset to defaults",
        
        # Misc
        "add": "Add",
        "remove": "Remove",
        "select": "Select",
        "selected": "Selected",
        "all": "All",
        "none": "None",
        "recent": "Recent",
        "favorites": "Favorites",
        "tags": "Tags",
        "date": "Date",
        "size": "Size",
        "name": "Name",
        "type": "Type",
        "sort_by": "Sort by",
        "filter_by": "Filter by",
    }
}

@api_router.get("/translations/{lang}")
async def get_translations(lang: str):
    """Get translations for a specific language (public endpoint)"""
    try:
        # First check if we have custom translations in the database
        data = await db.settings.find_one({"key": "localization"}, {"_id": 0})
        
        if data and "value" in data:
            translations = data["value"].get("translations", {})
            if lang in translations:
                # Merge with defaults (custom overrides defaults)
                merged = {**DEFAULT_TRANSLATIONS.get("en", {}), **DEFAULT_TRANSLATIONS.get(lang, {}), **translations.get(lang, {})}
                return {"lang": lang, "translations": merged}
        
        # Return default translations
        if lang in DEFAULT_TRANSLATIONS:
            return {"lang": lang, "translations": DEFAULT_TRANSLATIONS[lang]}
        
        # Fallback to English
        return {"lang": "en", "translations": DEFAULT_TRANSLATIONS["en"]}
    except Exception as e:
        logger.error(f"Get translations error: {e}")
        return {"lang": "en", "translations": DEFAULT_TRANSLATIONS["en"]}

@api_router.get("/translations")
async def get_all_translations():
    """Get all available translations (public endpoint)"""
    try:
        data = await db.settings.find_one({"key": "localization"}, {"_id": 0})
        
        if data and "value" in data:
            # Merge custom translations with defaults
            custom = data["value"].get("translations", {})
            languages = list(set(list(DEFAULT_TRANSLATIONS.keys()) + list(custom.keys())))
            return {
                "languages": languages,
                "default_language": "en",
                "translations": {**DEFAULT_TRANSLATIONS, **custom}
            }
        
        return {
            "languages": list(DEFAULT_TRANSLATIONS.keys()),
            "default_language": "en",
            "translations": DEFAULT_TRANSLATIONS
        }
    except Exception as e:
        logger.error(f"Get all translations error: {e}")
        return {
            "languages": ["en"],
            "default_language": "en",
            "translations": DEFAULT_TRANSLATIONS
        }

@api_router.get("/admin/localization")
async def get_localization(admin: dict = Depends(get_admin_user)):
    """Get all localization data (languages and translations)"""
    try:
        # Get all languages
        languages = await db.languages.find({}, {"_id": 0}).to_list(100)
        if not languages:
            languages = DEFAULT_LANGUAGES
        
        # Get all translations
        translations_cursor = db.translations.find({}, {"_id": 0})
        translations_list = await translations_cursor.to_list(100)
        
        translations = {}
        for t in translations_list:
            translations[t["language_code"]] = t.get("translations", {})
        
        # Ensure English defaults exist
        if "en" not in translations:
            translations["en"] = DEFAULT_TRANSLATIONS
        
        return {
            "languages": languages,
            "translations": translations,
            "default_keys": list(DEFAULT_TRANSLATIONS.keys())
        }
    except Exception as e:
        logger.error(f"Get localization error: {e}")
        return {
            "languages": DEFAULT_LANGUAGES,
            "translations": {"en": DEFAULT_TRANSLATIONS},
            "default_keys": list(DEFAULT_TRANSLATIONS.keys())
        }


@api_router.post("/admin/languages")
async def add_language(
    language: LanguageModel,
    admin: dict = Depends(get_admin_user)
):
    """Add a new language"""
    try:
        # Check if language already exists
        existing = await db.languages.find_one({"code": language.code})
        if existing:
            raise HTTPException(status_code=400, detail="Language already exists")
        
        # Add language
        await db.languages.insert_one(language.dict())
        
        # Initialize with English translations as default
        en_trans = await db.translations.find_one({"language_code": "en"})
        default_trans = en_trans.get("translations", DEFAULT_TRANSLATIONS) if en_trans else DEFAULT_TRANSLATIONS
        
        await db.translations.insert_one({
            "language_code": language.code,
            "translations": default_trans,  # Copy English as starting point
            "updated_at": datetime.now(timezone.utc)
        })
        
        return {"message": f"Language '{language.name}' added successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Add language error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.put("/admin/languages/{language_code}")
async def update_language(
    language_code: str,
    language: LanguageModel,
    admin: dict = Depends(get_admin_user)
):
    """Update a language"""
    try:
        result = await db.languages.update_one(
            {"code": language_code},
            {"$set": {
                "name": language.name,
                "native_name": language.native_name,
                "is_default": language.is_default
            }}
        )
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Language not found")
        return {"message": "Language updated"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Update language error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.delete("/admin/languages/{language_code}")
async def delete_language(
    language_code: str,
    admin: dict = Depends(get_admin_user)
):
    """Delete a language (cannot delete English)"""
    try:
        if language_code == "en":
            raise HTTPException(status_code=400, detail="Cannot delete default language")
        
        await db.languages.delete_one({"code": language_code})
        await db.translations.delete_one({"language_code": language_code})
        await db.legal_pages.delete_many({"language_code": language_code})
        
        return {"message": f"Language '{language_code}' deleted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete language error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.put("/admin/translations/{language_code}")
async def update_translations(
    language_code: str,
    data: TranslationUpdate,
    admin: dict = Depends(get_admin_user)
):
    """Update translations for a language"""
    try:
        result = await db.translations.update_one(
            {"language_code": language_code},
            {
                "$set": {
                    "translations": data.translations,
                    "updated_at": datetime.now(timezone.utc)
                }
            },
            upsert=True
        )
        return {"message": f"Translations for '{language_code}' updated"}
    except Exception as e:
        logger.error(f"Update translations error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.put("/admin/translations/{language_code}/key")
async def update_translation_key(
    language_code: str,
    data: TranslationKeyUpdate,
    admin: dict = Depends(get_admin_user)
):
    """Update a single translation key"""
    try:
        result = await db.translations.update_one(
            {"language_code": language_code},
            {
                "$set": {
                    f"translations.{data.key}": data.value,
                    "updated_at": datetime.now(timezone.utc)
                }
            },
            upsert=True
        )
        return {"message": f"Translation key '{data.key}' updated"}
    except Exception as e:
        logger.error(f"Update translation key error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Admin Legal Pages endpoints
@api_router.get("/admin/legal-pages")
async def get_all_legal_pages(admin: dict = Depends(get_admin_user)):
    """Get all legal pages"""
    try:
        pages = await db.legal_pages.find({}, {"_id": 0}).to_list(100)
        
        # Group by page type
        grouped = {}
        for page in pages:
            page_type = page.get("page_type")
            if page_type not in grouped:
                grouped[page_type] = {}
            grouped[page_type][page.get("language_code")] = {
                "content": page.get("content", ""),
                "updated_at": page.get("updated_at")
            }
        
        # Add defaults for missing pages
        for page_type in ["terms", "privacy", "support"]:
            if page_type not in grouped:
                grouped[page_type] = {}
            if "en" not in grouped[page_type]:
                grouped[page_type]["en"] = {
                    "content": DEFAULT_LEGAL_PAGES.get(page_type, {}).get("en", ""),
                    "updated_at": None
                }
        
        return grouped
    except Exception as e:
        logger.error(f"Get legal pages error: {e}")
        return {}


@api_router.put("/admin/legal-pages/{page_type}/{language_code}")
async def update_legal_page(
    page_type: str,
    language_code: str,
    data: LegalPageUpdate,
    admin: dict = Depends(get_admin_user)
):
    """Update a legal page"""
    if page_type not in ["terms", "privacy", "support"]:
        raise HTTPException(status_code=400, detail="Invalid page type")
    
    try:
        await db.legal_pages.update_one(
            {"page_type": page_type, "language_code": language_code},
            {
                "$set": {
                    "content": data.content,
                    "updated_at": datetime.now(timezone.utc)
                }
            },
            upsert=True
        )
        return {"message": f"Legal page '{page_type}' ({language_code}) updated"}
    except Exception as e:
        logger.error(f"Update legal page error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/admin/localization")
async def save_localization(
    data: dict,
    admin: dict = Depends(get_admin_user)
):
    """Save localization data (legacy endpoint)"""
    try:
        await db.settings.update_one(
            {"key": "localization"},
            {"$set": {"key": "localization", "value": data}},
            upsert=True
        )
        return {"message": "Localization saved"}
    except Exception as e:
        logger.error(f"Save localization error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/admin/localization/refresh")
async def refresh_translations_from_defaults(
    admin: dict = Depends(get_admin_user)
):
    """Refresh translations by merging new default keys (preserves existing values)"""
    try:
        # Update English translations - merge with defaults
        en_trans = await db.translations.find_one({"language_code": "en"})
        if en_trans:
            existing = en_trans.get("translations", {})
            # Merge: existing values take priority, new keys from defaults added
            merged_en = {**DEFAULT_TRANSLATIONS, **existing}
            await db.translations.update_one(
                {"language_code": "en"},
                {"$set": {"translations": merged_en, "updated_at": datetime.now(timezone.utc)}}
            )
        else:
            await db.translations.insert_one({
                "language_code": "en",
                "translations": DEFAULT_TRANSLATIONS,
                "updated_at": datetime.now(timezone.utc)
            })
        
        # Update German translations
        de_trans = await db.translations.find_one({"language_code": "de"})
        if de_trans:
            existing = de_trans.get("translations", {})
            merged_de = {**DEFAULT_TRANSLATIONS_DE, **existing}
            await db.translations.update_one(
                {"language_code": "de"},
                {"$set": {"translations": merged_de, "updated_at": datetime.now(timezone.utc)}}
            )
        else:
            await db.translations.insert_one({
                "language_code": "de",
                "translations": DEFAULT_TRANSLATIONS_DE,
                "updated_at": datetime.now(timezone.utc)
            })
        
        # For other languages, add any missing keys from English defaults
        other_langs = await db.translations.find({
            "language_code": {"$nin": ["en", "de"]}
        }).to_list(100)
        
        for lang in other_langs:
            existing = lang.get("translations", {})
            # Add missing keys from English (untranslated)
            merged = {**DEFAULT_TRANSLATIONS, **existing}
            await db.translations.update_one(
                {"language_code": lang["language_code"]},
                {"$set": {"translations": merged, "updated_at": datetime.now(timezone.utc)}}
            )
        
        return {"message": "Translations refreshed with new default keys"}
    except Exception as e:
        logger.error(f"Refresh translations error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.get("/admin/settings")
async def get_admin_settings(admin: dict = Depends(get_admin_user)):
    """Get app settings"""
    try:
        data = await db.settings.find_one({"key": "app_settings"}, {"_id": 0})
        if data:
            return data.get("value", {})
        return {}
    except Exception as e:
        logger.error(f"Get settings error: {e}")
        return {}

@api_router.post("/admin/settings")
async def save_admin_settings(
    settings: dict,
    admin: dict = Depends(get_admin_user)
):
    """Save app settings"""
    try:
        await db.settings.update_one(
            {"key": "app_settings"},
            {"$set": {"key": "app_settings", "value": settings}},
            upsert=True
        )
        return {"message": "Settings saved"}
    except Exception as e:
        logger.error(f"Save settings error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/admin/system-status")
async def get_system_status(admin: dict = Depends(get_admin_user)):
    """Get system status"""
    status = {
        "database": "disconnected",
        "storage": "disconnected",
        "api": "connected"
    }
    
    try:
        # Check database
        await db.command("ping")
        status["database"] = "connected"
    except:
        pass
    
    try:
        # Check S3
        if s3_client:
            s3_client.head_bucket(Bucket=AWS_S3_BUCKET_NAME)
            status["storage"] = "connected"
    except:
        pass
    
    return status


# Include the router in the main app (must be after all routes are defined)
app.include_router(api_router)

# ============== ADMIN DASHBOARD STATIC FILES ==============
# Serve the admin dashboard from /admin path
import os as os_module

admin_static_path = os_module.path.join(os_module.path.dirname(__file__), "admin-static")
if os_module.path.exists(admin_static_path):
    # Mount static assets (JS, CSS)
    app.mount("/admin/assets", StaticFiles(directory=os_module.path.join(admin_static_path, "assets")), name="admin-assets")
    
    # Serve admin dashboard index.html for all /admin routes (SPA support)
    @app.get("/admin")
    @app.get("/admin/")
    @app.get("/admin/{full_path:path}")
    async def serve_admin_dashboard(full_path: str = ""):
        index_path = os_module.path.join(admin_static_path, "index.html")
        if os_module.path.exists(index_path):
            return FileResponse(index_path)
        raise HTTPException(status_code=404, detail="Admin dashboard not found")
    
    logger.info("✅ Admin dashboard mounted at /admin")

# Serve the landing page from /landing path
landing_page_path = os_module.path.join(os_module.path.dirname(__file__), "landing-page")
if os_module.path.exists(landing_page_path):
    @app.get("/landing")
    @app.get("/landing/")
    async def serve_landing_page():
        index_path = os_module.path.join(landing_page_path, "index.html")
        if os_module.path.exists(index_path):
            return FileResponse(index_path, media_type="text/html")
        raise HTTPException(status_code=404, detail="Landing page not found")
    
    @app.get("/landing/terms")
    async def serve_terms_page():
        terms_path = os_module.path.join(landing_page_path, "terms.html")
        if os_module.path.exists(terms_path):
            return FileResponse(terms_path, media_type="text/html")
        raise HTTPException(status_code=404, detail="Terms page not found")
    
    @app.get("/landing/privacy")
    async def serve_privacy_page():
        privacy_path = os_module.path.join(landing_page_path, "privacy.html")
        if os_module.path.exists(privacy_path):
            return FileResponse(privacy_path, media_type="text/html")
        raise HTTPException(status_code=404, detail="Privacy page not found")
    
    @app.get("/landing/support")
    async def serve_support_page():
        support_path = os_module.path.join(landing_page_path, "support.html")
        if os_module.path.exists(support_path):
            return FileResponse(support_path, media_type="text/html")
        raise HTTPException(status_code=404, detail="Support page not found")
    
    @app.get("/landing/dashboard")
    async def serve_dashboard_page():
        dashboard_path = os_module.path.join(landing_page_path, "dashboard.html")
        if os_module.path.exists(dashboard_path):
            return FileResponse(dashboard_path, media_type="text/html")
        raise HTTPException(status_code=404, detail="Dashboard page not found")
    
    logger.info("✅ Landing page mounted at /landing")
